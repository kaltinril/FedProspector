"""Generate FedProspect Data Dictionary as a Word document."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = r"c:\git\fedProspect\docs\database"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "FedProspect-Data-Dictionary.docx")


def add_table_section(doc, name, purpose, ownership, data_source, columns,
                      relationships, change_detection, growth):
    """Add a complete table section to the document."""
    doc.add_heading(name, level=1)

    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(purpose)

    doc.add_heading("Schema Ownership", level=2)
    doc.add_paragraph(ownership)

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph(data_source)

    doc.add_heading("Key Columns", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Column Name", "Type", "Description"]):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for col_name, col_type, col_desc in columns:
        row = tbl.add_row().cells
        row[0].text = col_name
        row[1].text = col_type
        row[2].text = col_desc

    doc.add_heading("Relationships", level=2)
    doc.add_paragraph(relationships)

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(change_detection)

    doc.add_heading("Growth", level=2)
    doc.add_paragraph(growth)

    doc.add_page_break()


def build_document():
    doc = Document()

    # -- Title page --
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("FedProspect Data Dictionary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Federal Contract Prospecting System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    date_p = doc.add_paragraph("Generated: March 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    db_p = doc.add_paragraph("Database: fed_contracts (MySQL 8.4)")
    db_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # -- Table of Contents placeholder --
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "Tables are organized by domain. Use navigation pane or Ctrl+F to jump to a table."
    )
    toc_items = [
        "1. Opportunity Tables",
        "   opportunity, opportunity_history, opportunity_relationship, opportunity_attachment,",
        "   opportunity_attachment_intel, opportunity_intel_source, opportunity_poc, contracting_officer",
        "2. Award / Contract Tables",
        "   fpds_contract, usaspending_award, usaspending_transaction, sam_subaward, gsa_labor_rate",
        "3. Entity Tables",
        "   entity, entity_address, entity_naics, entity_psc, entity_business_type,",
        "   entity_sba_certification, entity_poc, entity_disaster_response, entity_history, sam_exclusion",
        "4. Reference Tables",
        "   ref_naics_code, ref_sba_size_standard, ref_naics_footnote, ref_psc_code,",
        "   ref_set_aside_type, ref_sba_type, ref_business_type, ref_entity_structure,",
        "   ref_country_code, ref_state_code, ref_fips_county",
        "5. ETL / Operations Tables",
        "   etl_load_log, etl_load_error, etl_data_quality_rule, etl_rate_limit,",
        "   etl_health_snapshot, data_load_request, usaspending_load_checkpoint, ai_usage_log",
        "6. Staging Tables",
        "   stg_entity_raw, stg_opportunity_raw, stg_fpds_award_raw, stg_usaspending_raw,",
        "   stg_exclusion_raw, stg_fedhier_raw, stg_subaward_raw",
        "7. Application / Multi-Tenant Tables",
        "   organization, app_user, app_session, prospect, prospect_note, prospect_team_member,",
        "   saved_search, notification, activity_log, organization_invite, organization_entity,",
        "   organization_certification, organization_naics, organization_past_performance,",
        "   proposal, proposal_document, proposal_milestone",
    ]
    for line in toc_items:
        doc.add_paragraph(line)
    doc.add_page_break()

    # =========================================================================
    # DOMAIN 1: OPPORTUNITY TABLES
    # =========================================================================
    doc.add_heading("1. Opportunity Tables", level=0)
    doc.add_paragraph(
        "Contract opportunities sourced from the SAM.gov Opportunities API. "
        "These are the primary objects users search and prospect against."
    )

    add_table_section(doc,
        name="opportunity",
        purpose="Core table storing federal contract opportunities (solicitations, presolicitations, awards, etc.) from SAM.gov. Each row is one notice_id representing a single procurement action.",
        ownership="Python DDL (tables/30_opportunity.sql)",
        data_source="SAM.gov Opportunities API v2 via opportunity_loader.py. Loaded by CLI command: python main.py load opportunities",
        columns=[
            ("notice_id", "VARCHAR(100) PK", "SAM.gov unique opportunity identifier (e.g., '3fa85f6457174562b3fc2c963f66afa6')"),
            ("title", "VARCHAR(500)", "Human-readable title of the contract opportunity"),
            ("solicitation_number", "VARCHAR(100)", "Government-assigned solicitation/RFP number (e.g., 'W911QY-24-R-0042')"),
            ("department_name", "VARCHAR(200)", "Top-level department (e.g., 'Department of Defense')"),
            ("sub_tier", "VARCHAR(200)", "Sub-tier agency within the department (e.g., 'Army Contracting Command')"),
            ("office", "VARCHAR(200)", "Contracting office name"),
            ("posted_date", "DATE", "Date the opportunity was posted on SAM.gov"),
            ("response_deadline", "DATETIME", "Deadline for proposal/response submission"),
            ("archive_date", "DATE", "Date the opportunity will be archived from SAM.gov"),
            ("type", "VARCHAR(50)", "Opportunity type: Solicitation, Presolicitation, Award Notice, etc."),
            ("base_type", "VARCHAR(50)", "Base opportunity type before modifications"),
            ("set_aside_code", "VARCHAR(20)", "Set-aside designation code (e.g., 'SBA' for small business, 'WOSB')"),
            ("set_aside_description", "VARCHAR(200)", "Full text description of the set-aside type"),
            ("classification_code", "VARCHAR(10)", "Product Service Code (PSC) for this opportunity"),
            ("naics_code", "VARCHAR(6)", "NAICS industry code determining small business size standards"),
            ("pop_state", "VARCHAR(6)", "Place of performance state (ISO 3166-2 subdivision code)"),
            ("pop_zip", "VARCHAR(20)", "Place of performance ZIP code (supports military APO/FPO)"),
            ("pop_country", "VARCHAR(3)", "Place of performance country (ISO 3166 alpha-3)"),
            ("pop_city", "VARCHAR(100)", "Place of performance city name"),
            ("active", "CHAR(1)", "Whether opportunity is currently active ('Y') or archived ('N')"),
            ("award_number", "VARCHAR(200)", "Contract award number if opportunity has been awarded"),
            ("award_date", "DATE", "Date the contract was awarded"),
            ("award_amount", "DECIMAL(15,2)", "Dollar amount of the award"),
            ("awardee_uei", "VARCHAR(12)", "Unique Entity Identifier of the award recipient"),
            ("awardee_name", "VARCHAR(200)", "Legal business name of the award recipient"),
            ("description_url", "VARCHAR(500)", "SAM.gov API URL to fetch the full description text"),
            ("description_text", "LONGTEXT", "Cached full description fetched from description_url"),
            ("link", "VARCHAR(500)", "Direct URL to the opportunity on SAM.gov"),
            ("resource_links", "JSON", "Array of URLs to additional resources/documents"),
            ("record_hash", "CHAR(64)", "SHA-256 hash of key fields for change detection"),
            ("first_loaded_at", "DATETIME", "Timestamp of initial load into the database"),
            ("last_loaded_at", "DATETIME", "Timestamp of most recent update"),
            ("last_load_id", "INT", "FK to etl_load_log identifying the load batch"),
        ],
        relationships="Referenced by: prospect (notice_id), opportunity_history (notice_id), opportunity_relationship (parent/child_notice_id), opportunity_attachment (notice_id), opportunity_attachment_intel (notice_id), opportunity_poc (notice_id). References: ref_set_aside_type (set_aside_code), ref_naics_code (naics_code), ref_psc_code (classification_code).",
        change_detection="SHA-256 record_hash on key fields. Change detector compares incoming hash with stored hash; only updates rows where hash differs. Field-level changes logged to opportunity_history.",
        growth="50,000-200,000 rows. ~500-2,000 new opportunities per weekly load; old opportunities archived but retained."
    )

    add_table_section(doc,
        name="opportunity_history",
        purpose="Audit trail of field-level changes to opportunity records. Each row records one field change detected during an ETL load.",
        ownership="Python DDL (tables/30_opportunity.sql)",
        data_source="Populated automatically by opportunity_loader.py when a record_hash change is detected and field diffs are computed.",
        columns=[
            ("id", "BIGINT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("notice_id", "VARCHAR(100)", "The opportunity that changed"),
            ("field_name", "VARCHAR(100)", "Name of the column that changed (e.g., 'response_deadline', 'type')"),
            ("old_value", "TEXT", "Previous value of the field"),
            ("new_value", "TEXT", "New value of the field"),
            ("changed_at", "DATETIME", "When the change was detected"),
            ("load_id", "INT", "FK to etl_load_log identifying which load detected the change"),
        ],
        relationships="References: opportunity (notice_id), etl_load_log (load_id).",
        change_detection="None (append-only audit log).",
        growth="500,000-2,000,000+ rows. Grows proportionally to opportunity updates; ~5-10 field changes per updated opportunity."
    )

    add_table_section(doc,
        name="opportunity_relationship",
        purpose="Links related opportunities together, such as a presolicitation linked to its subsequent solicitation, or a solicitation linked to its award notice.",
        ownership="Python DDL (tables/30_opportunity.sql)",
        data_source="Created by users or automated matching in the application. Types: RFI_TO_RFP, PRESOL_TO_SOL, SOL_TO_AWARD.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("parent_notice_id", "VARCHAR(100)", "The earlier/parent opportunity"),
            ("child_notice_id", "VARCHAR(100)", "The later/child opportunity"),
            ("relationship_type", "VARCHAR(30)", "Type of link: RFI_TO_RFP, PRESOL_TO_SOL, SOL_TO_AWARD"),
            ("created_by", "INT", "User who created the link (NULL if system-generated)"),
            ("created_at", "DATETIME", "When the relationship was created"),
            ("notes", "TEXT", "Optional notes explaining the relationship"),
        ],
        relationships="References: opportunity (parent_notice_id, child_notice_id), app_user (created_by).",
        change_detection="None.",
        growth="1,000-10,000 rows. Grows slowly as users link related opportunities."
    )

    add_table_section(doc,
        name="opportunity_attachment",
        purpose="Tracks downloadable files (PDFs, Excel, Word docs) attached to SAM.gov opportunities. Manages the download and text extraction pipeline.",
        ownership="Python DDL (tables/36_attachment.sql)",
        data_source="URLs discovered by opportunity_loader.py from resource_links. Files downloaded by attachment_downloader.py, text extracted by attachment_text_extractor.py.",
        columns=[
            ("attachment_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("notice_id", "VARCHAR(100)", "The opportunity this attachment belongs to"),
            ("url", "VARCHAR(500)", "Download URL for the attachment"),
            ("filename", "VARCHAR(500)", "Original filename from the URL or Content-Disposition header"),
            ("content_type", "VARCHAR(100)", "MIME type (e.g., 'application/pdf', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')"),
            ("file_size_bytes", "BIGINT", "File size in bytes after download"),
            ("file_path", "VARCHAR(500)", "Local filesystem path where downloaded file is stored (under ATTACHMENT_DIR)"),
            ("extracted_text", "LONGTEXT", "Full text content extracted from the document"),
            ("page_count", "INT", "Number of pages (PDFs only)"),
            ("is_scanned", "TINYINT", "Whether the PDF is a scanned image requiring OCR (0=no, 1=yes)"),
            ("ocr_quality", "ENUM", "OCR quality assessment: good, fair, poor"),
            ("download_status", "ENUM", "Pipeline state: pending, downloaded, failed, skipped"),
            ("extraction_status", "ENUM", "Text extraction state: pending, extracted, failed, unsupported"),
            ("content_hash", "CHAR(64)", "SHA-256 hash of the file content for deduplication"),
            ("text_hash", "CHAR(64)", "SHA-256 hash of the extracted text for change detection"),
            ("downloaded_at", "DATETIME", "When the file was successfully downloaded"),
            ("extracted_at", "DATETIME", "When text extraction completed"),
            ("last_load_id", "INT", "FK to etl_load_log"),
            ("created_at", "DATETIME", "Row creation timestamp"),
        ],
        relationships="References: opportunity (notice_id). Referenced by: opportunity_attachment_intel (attachment_id), opportunity_intel_source (attachment_id).",
        change_detection="content_hash for file dedup; text_hash for re-extraction detection.",
        growth="100,000-500,000 rows. ~2-5 attachments per opportunity."
    )

    add_table_section(doc,
        name="opportunity_attachment_intel",
        purpose="Intelligence extracted from opportunity attachments: security clearance requirements, evaluation methods, contract vehicles, incumbent/recompete signals, and scope summaries.",
        ownership="Python DDL (tables/36_attachment.sql)",
        data_source="Populated by attachment_intel_extractor.py using keyword matching, heuristics, or AI models (Haiku/Sonnet).",
        columns=[
            ("intel_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("notice_id", "VARCHAR(100)", "The opportunity this intel relates to"),
            ("attachment_id", "INT", "The specific attachment the intel was extracted from (NULL for merged intel)"),
            ("extraction_method", "ENUM", "How intel was extracted: keyword, heuristic, ai_haiku, ai_sonnet"),
            ("source_text_hash", "CHAR(64)", "Hash of source text to detect when re-extraction is needed"),
            ("clearance_required", "CHAR(1)", "Whether security clearance is required (Y/N/NULL)"),
            ("clearance_level", "VARCHAR(50)", "Required clearance level (e.g., 'SECRET', 'TOP SECRET/SCI')"),
            ("clearance_scope", "VARCHAR(50)", "Scope of clearance requirement (e.g., 'all personnel', 'key personnel')"),
            ("clearance_details", "TEXT", "Full text details about clearance requirements"),
            ("eval_method", "VARCHAR(50)", "Evaluation methodology (e.g., 'LPTA', 'Best Value', 'Trade-Off')"),
            ("eval_details", "TEXT", "Details about evaluation criteria and weighting"),
            ("vehicle_type", "VARCHAR(100)", "Contract vehicle type (e.g., 'GSA Schedule', 'GWAC', 'BPA')"),
            ("vehicle_details", "TEXT", "Details about the contract vehicle"),
            ("is_recompete", "CHAR(1)", "Whether this is a recompete of an existing contract (Y/N/NULL)"),
            ("incumbent_name", "VARCHAR(200)", "Name of the current incumbent contractor"),
            ("recompete_details", "TEXT", "Details about the recompete situation"),
            ("scope_summary", "TEXT", "AI or heuristic-generated summary of the contract scope"),
            ("period_of_performance", "VARCHAR(200)", "Duration/period of performance extracted from documents"),
            ("labor_categories", "JSON", "Array of labor categories and qualifications mentioned"),
            ("key_requirements", "JSON", "Array of key technical/functional requirements"),
            ("overall_confidence", "ENUM", "Confidence in the extraction: high, medium, low"),
            ("confidence_details", "JSON", "Per-field confidence scores and rationale"),
            ("last_load_id", "INT", "FK to etl_load_log"),
            ("extracted_at", "DATETIME", "When the intel was extracted"),
        ],
        relationships="References: opportunity (notice_id), opportunity_attachment (attachment_id). Referenced by: opportunity_intel_source (intel_id).",
        change_detection="source_text_hash compared to detect when source document has changed and re-extraction is needed.",
        growth="50,000-200,000 rows. One row per attachment per extraction method per opportunity."
    )

    add_table_section(doc,
        name="opportunity_intel_source",
        purpose="Provenance/citation records for each piece of extracted intelligence. Links an intel finding back to the exact text location in the source attachment.",
        ownership="Python DDL (tables/36_attachment.sql)",
        data_source="Populated by attachment_intel_extractor.py alongside opportunity_attachment_intel rows.",
        columns=[
            ("source_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("intel_id", "INT", "FK to the intel record this source supports"),
            ("field_name", "VARCHAR(50)", "Which intel field this source supports (e.g., 'clearance_required', 'eval_method')"),
            ("attachment_id", "INT", "FK to the attachment file containing the source text"),
            ("source_filename", "VARCHAR(500)", "Filename of the source document"),
            ("page_number", "INT", "Page number where the text was found (1-based)"),
            ("char_offset_start", "INT", "Character offset of the match start within the page/document"),
            ("char_offset_end", "INT", "Character offset of the match end"),
            ("matched_text", "VARCHAR(500)", "The exact text that was matched/extracted"),
            ("surrounding_context", "TEXT", "Broader text context around the match for verification"),
            ("pattern_name", "VARCHAR(100)", "Name of the regex/keyword pattern that matched"),
            ("extraction_method", "ENUM", "Method used: keyword, heuristic, ai_haiku, ai_sonnet"),
            ("confidence", "ENUM", "Confidence level: high, medium, low"),
            ("created_at", "DATETIME", "Row creation timestamp"),
        ],
        relationships="References: opportunity_attachment_intel (intel_id), opportunity_attachment (attachment_id).",
        change_detection="None (regenerated when parent intel record is refreshed).",
        growth="200,000-1,000,000 rows. Multiple source citations per intel finding."
    )

    add_table_section(doc,
        name="opportunity_poc",
        purpose="Junction table linking opportunities to their contracting officer points of contact. An opportunity can have primary and secondary POCs.",
        ownership="Python DDL (tables/90_web_api.sql)",
        data_source="Populated by opportunity_loader.py from SAM.gov opportunity POC data.",
        columns=[
            ("poc_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("notice_id", "VARCHAR(100)", "FK to the opportunity"),
            ("officer_id", "INT", "FK to the contracting_officer record"),
            ("poc_type", "VARCHAR(20)", "Contact role: PRIMARY or SECONDARY"),
            ("created_at", "DATETIME", "Row creation timestamp"),
        ],
        relationships="References: opportunity (notice_id), contracting_officer (officer_id).",
        change_detection="None. Upserted on each load.",
        growth="50,000-200,000 rows. 1-2 POCs per opportunity."
    )

    add_table_section(doc,
        name="contracting_officer",
        purpose="Deduplicated directory of government contracting officers encountered across opportunities. Used for relationship tracking and contact intelligence.",
        ownership="Python DDL (tables/90_web_api.sql)",
        data_source="Populated by opportunity_loader.py. Deduplicated on (full_name, email).",
        columns=[
            ("officer_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("full_name", "VARCHAR(500)", "Officer's full name"),
            ("email", "VARCHAR(200)", "Email address"),
            ("phone", "VARCHAR(100)", "Phone number"),
            ("fax", "VARCHAR(100)", "Fax number"),
            ("title", "VARCHAR(200)", "Job title"),
            ("department_name", "VARCHAR(200)", "Department/agency the officer belongs to"),
            ("office_name", "VARCHAR(200)", "Specific contracting office name"),
            ("officer_type", "VARCHAR(50)", "Type of officer role"),
            ("created_at", "DATETIME", "Row creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="Referenced by: opportunity_poc (officer_id).",
        change_detection="None. Upserted by name+email composite key.",
        growth="10,000-50,000 rows. Grows with new unique officers discovered."
    )

    # =========================================================================
    # DOMAIN 2: AWARD / CONTRACT TABLES
    # =========================================================================
    doc.add_heading("2. Award / Contract Tables", level=0)
    doc.add_paragraph(
        "Federal contract award data from FPDS (via SAM.gov Contract Awards API), "
        "USASpending.gov, and SAM.gov subaward data."
    )

    add_table_section(doc,
        name="fpds_contract",
        purpose="FPDS (Federal Procurement Data System) contract records from the SAM.gov Awards API. Each row is a contract action (base award or modification). Provides incumbent analysis, competitive landscape, and historical spend data.",
        ownership="Python DDL (tables/40_federal.sql)",
        data_source="SAM.gov Contract Awards API via awards_loader.py. CLI: python main.py load awards",
        columns=[
            ("contract_id", "VARCHAR(100) PK", "PIID (Procurement Instrument Identifier) -- the contract number"),
            ("modification_number", "VARCHAR(25) PK", "Modification sequence number ('0' for base award)"),
            ("idv_piid", "VARCHAR(50)", "Parent IDV (Indefinite Delivery Vehicle) contract number, if this is a task/delivery order"),
            ("transaction_number", "VARCHAR(10)", "Transaction sequence within the modification"),
            ("agency_id", "VARCHAR(10)", "Awarding agency FPDS code"),
            ("agency_name", "VARCHAR(200)", "Awarding agency name"),
            ("contracting_office_id", "VARCHAR(20)", "FPDS contracting office code"),
            ("contracting_office_name", "VARCHAR(200)", "Contracting office name"),
            ("funding_agency_id", "VARCHAR(10)", "Agency providing the funding (may differ from awarding)"),
            ("funding_agency_name", "VARCHAR(200)", "Funding agency name"),
            ("vendor_uei", "VARCHAR(12)", "Contractor UEI (links to entity table)"),
            ("vendor_name", "VARCHAR(200)", "Contractor legal business name"),
            ("date_signed", "DATE", "Date the contract action was signed"),
            ("effective_date", "DATE", "Contract effective date"),
            ("completion_date", "DATE", "Expected completion date of the current period"),
            ("ultimate_completion_date", "DATE", "Final completion date including all options"),
            ("dollars_obligated", "DECIMAL(15,2)", "Dollar amount obligated in this action"),
            ("base_and_all_options", "DECIMAL(15,2)", "Total potential value including all option years"),
            ("naics_code", "VARCHAR(6)", "NAICS code for the contracted work"),
            ("psc_code", "VARCHAR(10)", "Product Service Code classifying the type of work"),
            ("set_aside_type", "VARCHAR(20)", "Set-aside designation (e.g., 'SBA', 'WOSB', '8AN')"),
            ("type_of_contract", "VARCHAR(10)", "Contract type code (FFP, T&M, CPFF, etc.)"),
            ("description", "TEXT", "Contract description text"),
            ("extent_competed", "VARCHAR(10)", "Competition level (full, limited, sole source, etc.)"),
            ("number_of_offers", "INT", "Number of proposals/offers received"),
            ("solicitation_number", "VARCHAR(200)", "Links back to opportunity solicitation_number"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_loaded_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="References: entity (vendor_uei). Cross-references opportunity via solicitation_number. Referenced by views: expiring_contracts, incumbent_profile, competitor_analysis.",
        change_detection="SHA-256 record_hash. Changes logged but no separate history table.",
        growth="500,000-2,000,000+ rows. Includes multi-year history; ~50,000-100,000 new actions per bulk load."
    )

    add_table_section(doc,
        name="usaspending_award",
        purpose="Award-level spending data from USASpending.gov bulk CSV downloads. Provides obligation totals, recipient info, and contract periods for market intelligence and expiring contract analysis.",
        ownership="Python DDL (tables/70_usaspending.sql)",
        data_source="USASpending.gov bulk CSV files via usaspending_bulk_loader.py. CLI: python main.py load usaspending-bulk",
        columns=[
            ("generated_unique_award_id", "VARCHAR(100) PK", "USASpending-generated unique award identifier"),
            ("piid", "VARCHAR(50)", "PIID for contracts; links to fpds_contract"),
            ("fain", "VARCHAR(30)", "Federal Award Identification Number (grants/assistance)"),
            ("uri", "VARCHAR(70)", "Unique Record Identifier"),
            ("award_type", "VARCHAR(50)", "Award type (contract, grant, loan, etc.)"),
            ("award_description", "TEXT", "Description of the award"),
            ("recipient_name", "VARCHAR(200)", "Prime recipient/contractor name"),
            ("recipient_uei", "VARCHAR(12)", "Recipient UEI (links to entity table)"),
            ("total_obligation", "DECIMAL(15,2)", "Total federal dollars obligated"),
            ("base_and_all_options_value", "DECIMAL(15,2)", "Total potential value"),
            ("start_date", "DATE", "Contract/award start date"),
            ("end_date", "DATE", "Contract/award end date"),
            ("awarding_agency_name", "VARCHAR(200)", "Agency that made the award"),
            ("naics_code", "VARCHAR(6)", "NAICS code"),
            ("psc_code", "VARCHAR(10)", "Product Service Code"),
            ("type_of_set_aside", "VARCHAR(100)", "Set-aside type code"),
            ("pop_state", "VARCHAR(6)", "Place of performance state"),
            ("solicitation_identifier", "VARCHAR(100)", "Solicitation number for cross-referencing"),
            ("fiscal_year", "SMALLINT", "Federal fiscal year of the data"),
            ("fpds_enriched_at", "DATETIME", "When FPDS enrichment was last applied"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("deleted_at", "DATETIME", "Soft-delete timestamp (preserves transaction FK references)"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_loaded_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Referenced by: usaspending_transaction (award_id). Cross-references entity (recipient_uei), fpds_contract (piid), opportunity (solicitation_identifier).",
        change_detection="SHA-256 record_hash. Soft-delete via deleted_at column.",
        growth="2,000,000-5,000,000+ rows. Bulk loaded per fiscal year; one-time historical load plus incremental yearly adds."
    )

    add_table_section(doc,
        name="usaspending_transaction",
        purpose="Transaction-level spending detail for individual contract actions within a USASpending award. Enables burn rate analysis and obligation trend tracking.",
        ownership="Python DDL (tables/70_usaspending.sql)",
        data_source="USASpending.gov API via usaspending_loader.py. CLI: python main.py load usaspending",
        columns=[
            ("id", "BIGINT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("award_id", "VARCHAR(100)", "FK to usaspending_award"),
            ("action_date", "DATE", "Date of the contract action"),
            ("modification_number", "VARCHAR(25)", "Contract modification number"),
            ("action_type", "VARCHAR(5)", "Action type code (A=new, B=continuation, etc.)"),
            ("action_type_description", "VARCHAR(100)", "Human-readable action type"),
            ("federal_action_obligation", "DECIMAL(15,2)", "Dollar amount of this specific action"),
            ("description", "TEXT", "Description of the action"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="References: usaspending_award (award_id).",
        change_detection="Deduplicated on (award_id, modification_number, action_date).",
        growth="5,000,000-20,000,000+ rows. Multiple transactions per award."
    )

    add_table_section(doc,
        name="sam_subaward",
        purpose="Subaward data from SAM.gov showing subcontracting relationships. Reveals which small businesses receive work as subcontractors under prime contracts.",
        ownership="Python DDL (tables/40_federal.sql)",
        data_source="SAM.gov Subaward API via subaward_loader.py. CLI: python main.py load subawards",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("prime_piid", "VARCHAR(50)", "Prime contract number (PIID)"),
            ("prime_agency_id", "VARCHAR(10)", "Prime contract awarding agency code"),
            ("prime_agency_name", "VARCHAR(200)", "Prime contract awarding agency name"),
            ("prime_uei", "VARCHAR(12)", "Prime contractor UEI"),
            ("prime_name", "VARCHAR(500)", "Prime contractor name"),
            ("sub_uei", "VARCHAR(12)", "Subcontractor UEI"),
            ("sub_name", "VARCHAR(500)", "Subcontractor name"),
            ("sub_amount", "DECIMAL(15,2)", "Dollar amount of the subaward"),
            ("sub_date", "DATE", "Date of the subaward"),
            ("sub_description", "TEXT", "Description of subcontracted work"),
            ("naics_code", "VARCHAR(6)", "NAICS code for the subcontracted work"),
            ("sub_business_type", "VARCHAR(50)", "Small business designation codes for the subcontractor"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_updated_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Cross-references: entity (prime_uei, sub_uei), fpds_contract (prime_piid).",
        change_detection="SHA-256 record_hash. Deduplicated on (prime_piid, sub_uei, sub_date).",
        growth="500,000-2,000,000 rows. Grows with subaward reporting requirements."
    )

    add_table_section(doc,
        name="gsa_labor_rate",
        purpose="GSA Schedule labor rates from the CALC+ (Contract-Awarded Labor Category) tool. Used for pricing intelligence and competitive rate analysis.",
        ownership="Python DDL (tables/40_federal.sql)",
        data_source="GSA CALC+ API via calc_loader.py. CLI: python main.py load calc",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("labor_category", "VARCHAR(200)", "Labor category title (e.g., 'Senior Systems Engineer')"),
            ("education_level", "VARCHAR(50)", "Required education level"),
            ("min_years_experience", "INT", "Minimum years of experience"),
            ("current_price", "DECIMAL(10,2)", "Current hourly rate"),
            ("next_year_price", "DECIMAL(10,2)", "Projected rate for next option year"),
            ("second_year_price", "DECIMAL(10,2)", "Projected rate for second option year"),
            ("schedule", "VARCHAR(200)", "GSA Schedule name/number"),
            ("contractor_name", "VARCHAR(200)", "Contractor holding this rate"),
            ("sin", "VARCHAR(500)", "Special Item Number(s)"),
            ("business_size", "VARCHAR(10)", "Small or large business"),
            ("security_clearance", "VARCHAR(50)", "Clearance level priced into the rate"),
            ("idv_piid", "VARCHAR(50)", "GSA contract number"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_loaded_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Cross-references entity via contractor_name (no FK). References GSA schedule contracts.",
        change_detection="Full reload (truncate and reload). No hash-based detection.",
        growth="500,000-1,000,000 rows. Full dataset reloaded periodically."
    )

    # =========================================================================
    # DOMAIN 3: ENTITY TABLES
    # =========================================================================
    doc.add_heading("3. Entity Tables", level=0)
    doc.add_paragraph(
        "SAM.gov Entity (contractor/vendor) registration data. Entities are identified "
        "by UEI (Unique Entity Identifier). Child tables store multi-valued attributes."
    )

    add_table_section(doc,
        name="entity",
        purpose="Core contractor/vendor registration data from SAM.gov Entity API. Each row represents one registered entity identified by its UEI. Contains registration status, business structure, and primary NAICS.",
        ownership="Python DDL (tables/20_entity.sql)",
        data_source="SAM.gov Entity Management API via entity_loader.py. CLI: python main.py load entities",
        columns=[
            ("uei_sam", "VARCHAR(12) PK", "SAM.gov Unique Entity Identifier (12-char alphanumeric)"),
            ("uei_duns", "VARCHAR(9)", "Legacy DUNS number (deprecated, retained for historical linking)"),
            ("cage_code", "VARCHAR(5)", "Commercial and Government Entity code (DoD identifier)"),
            ("registration_status", "VARCHAR(1)", "SAM registration status: A=Active, E=Expired, etc."),
            ("initial_registration_date", "DATE", "When the entity first registered in SAM.gov"),
            ("registration_expiration_date", "DATE", "When the current registration expires"),
            ("legal_business_name", "VARCHAR(120)", "Official legal name of the business"),
            ("dba_name", "VARCHAR(120)", "Doing-business-as name"),
            ("entity_url", "VARCHAR(200)", "Company website URL"),
            ("entity_structure_code", "VARCHAR(2)", "Business structure code (LLC, Corp, etc.) referencing ref_entity_structure"),
            ("primary_naics", "VARCHAR(6)", "Primary NAICS code for the business"),
            ("exclusion_status_flag", "VARCHAR(1)", "Whether entity has active exclusions"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_loaded_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Referenced by: entity_address, entity_naics, entity_psc, entity_business_type, entity_sba_certification, entity_poc, entity_disaster_response, entity_history (all via uei_sam FK). Also referenced by organization_entity. Cross-referenced by fpds_contract (vendor_uei), opportunity (awardee_uei).",
        change_detection="SHA-256 record_hash. Field-level changes logged to entity_history.",
        growth="100,000-500,000 rows. Grows with new SAM registrations discovered during loads."
    )

    for tbl_name, tbl_purpose, tbl_cols, tbl_rels, tbl_growth in [
        ("entity_address",
         "Physical and mailing addresses for SAM-registered entities. Each entity has up to 2 addresses (physical, mailing).",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("address_type", "VARCHAR(10)", "Address type: 'PHYSICAL' or 'MAILING'"),
          ("address_line_1", "VARCHAR(150)", "Street address line 1"),
          ("address_line_2", "VARCHAR(150)", "Street address line 2"),
          ("city", "VARCHAR(40)", "City name"),
          ("state_or_province", "VARCHAR(55)", "State/province name or code"),
          ("zip_code", "VARCHAR(50)", "ZIP/postal code"),
          ("country_code", "VARCHAR(3)", "ISO 3166 alpha-3 country code"),
          ("congressional_district", "VARCHAR(10)", "US congressional district number")],
         "References: entity (uei_sam) with CASCADE delete.",
         "200,000-1,000,000 rows. ~2 per entity."),

        ("entity_naics",
         "NAICS codes an entity is registered to perform work under. Entities typically register for 5-20 NAICS codes.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("naics_code", "VARCHAR(11)", "NAICS code (links to ref_naics_code)"),
          ("is_primary", "CHAR(1)", "Whether this is the entity's primary NAICS (Y/N)"),
          ("sba_small_business", "VARCHAR(1)", "Whether entity qualifies as small business under this NAICS"),
          ("naics_exception", "VARCHAR(20)", "Any size standard exception applied")],
         "References: entity (uei_sam) with CASCADE delete. Cross-references ref_naics_code.",
         "1,000,000-5,000,000 rows. ~10 NAICS per entity."),

        ("entity_psc",
         "Product Service Codes an entity is registered to perform. Maps entities to the types of products/services they provide.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("psc_code", "VARCHAR(10)", "PSC code (links to ref_psc_code)")],
         "References: entity (uei_sam) with CASCADE delete. Cross-references ref_psc_code.",
         "500,000-2,000,000 rows. Variable per entity."),

        ("entity_business_type",
         "Business type classifications for an entity (e.g., Woman-Owned, Veteran-Owned, 8(a), HUBZone). Critical for set-aside eligibility analysis.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("business_type_code", "VARCHAR(4)", "Business type code (links to ref_business_type)")],
         "References: entity (uei_sam) with CASCADE delete. Cross-references ref_business_type.",
         "500,000-2,000,000 rows. ~5 types per entity."),

        ("entity_sba_certification",
         "SBA certifications held by an entity (8(a), HUBZone, WOSB, EDWOSB, VOSB, SDVOSB). Tracks certification dates including graduation/expiration.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("sba_type_code", "VARCHAR(10)", "SBA certification type code (links to ref_sba_type)"),
          ("sba_type_desc", "VARCHAR(200)", "Description of the certification"),
          ("certification_entry_date", "DATE", "When the certification was granted"),
          ("certification_exit_date", "DATE", "When the certification expires or was exited (NULL = active)")],
         "References: entity (uei_sam) with CASCADE delete. Cross-references ref_sba_type.",
         "50,000-200,000 rows. Only entities with SBA certifications."),

        ("entity_poc",
         "Points of contact for SAM-registered entities. Includes government business, electronic business, and alternate contacts.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("poc_type", "VARCHAR(40)", "Contact role type (e.g., 'GOVT_BUSINESS_POC', 'ELECTRONIC_BUSINESS_POC')"),
          ("first_name", "VARCHAR(65)", "Contact first name"),
          ("last_name", "VARCHAR(65)", "Contact last name"),
          ("title", "VARCHAR(50)", "Job title"),
          ("address_line_1", "VARCHAR(150)", "Contact mailing address"),
          ("city", "VARCHAR(40)", "Contact city"),
          ("state_or_province", "VARCHAR(55)", "Contact state"),
          ("country_code", "VARCHAR(3)", "Contact country code")],
         "References: entity (uei_sam) with CASCADE delete.",
         "200,000-1,000,000 rows. ~2-4 POCs per entity."),

        ("entity_disaster_response",
         "Geographic areas where an entity has indicated willingness to provide disaster response services.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "FK to entity"),
          ("state_code", "VARCHAR(10)", "State FIPS code"),
          ("state_name", "VARCHAR(60)", "State name"),
          ("county_code", "VARCHAR(5)", "County FIPS code"),
          ("county_name", "VARCHAR(100)", "County name"),
          ("msa_code", "VARCHAR(10)", "Metropolitan Statistical Area code"),
          ("msa_name", "VARCHAR(100)", "MSA name")],
         "References: entity (uei_sam) with CASCADE delete.",
         "100,000-500,000 rows. Only entities registered for disaster response."),

        ("entity_history",
         "Audit trail of field-level changes to entity records, similar to opportunity_history. Records what changed, when, and in which load.",
         [("id", "BIGINT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("uei_sam", "VARCHAR(12)", "The entity that changed"),
          ("field_name", "VARCHAR(100)", "Name of the column that changed"),
          ("old_value", "TEXT", "Previous value"),
          ("new_value", "TEXT", "New value"),
          ("changed_at", "DATETIME", "When the change was detected"),
          ("load_id", "INT", "FK to etl_load_log")],
         "References: entity (uei_sam), etl_load_log (load_id).",
         "500,000-2,000,000+ rows. Append-only audit log."),
    ]:
        add_table_section(doc, name=tbl_name, purpose=tbl_purpose,
            ownership="Python DDL (tables/20_entity.sql)",
            data_source="SAM.gov Entity Management API via entity_loader.py",
            columns=tbl_cols, relationships=tbl_rels,
            change_detection="Replaced in full on each entity update (child rows deleted and re-inserted when parent entity hash changes)." if tbl_name != "entity_history" else "None (append-only).",
            growth=tbl_growth)

    add_table_section(doc,
        name="sam_exclusion",
        purpose="SAM.gov exclusion (debarment/suspension) records. Entities with active exclusions are barred from receiving federal contracts. Critical for risk assessment.",
        ownership="Python DDL (tables/40_federal.sql)",
        data_source="SAM.gov Exclusions API via exclusions_loader.py. CLI: python main.py load exclusions",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("uei", "VARCHAR(12)", "Entity UEI (may be NULL for individuals)"),
            ("cage_code", "VARCHAR(10)", "CAGE code of excluded entity"),
            ("entity_name", "VARCHAR(500)", "Name of excluded entity/individual"),
            ("first_name", "VARCHAR(100)", "Individual first name (for individual exclusions)"),
            ("last_name", "VARCHAR(100)", "Individual last name"),
            ("exclusion_type", "VARCHAR(50)", "Type: Ineligible (Proceedings Pending), Ineligible (Proceedings Completed), Prohibition/Restriction, Voluntary"),
            ("exclusion_program", "VARCHAR(50)", "Program: Reciprocal, Nonreciprocal, or Voluntary"),
            ("excluding_agency_code", "VARCHAR(10)", "Agency that imposed the exclusion"),
            ("excluding_agency_name", "VARCHAR(200)", "Name of the excluding agency"),
            ("activation_date", "DATE", "Date exclusion became active"),
            ("termination_date", "DATE", "Date exclusion expires (NULL = indefinite)"),
            ("classification_type", "VARCHAR(50)", "Individual or Firm"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_updated_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Cross-references entity (uei). No formal FK to allow exclusions for unregistered entities.",
        change_detection="SHA-256 record_hash. Deduplicated on (uei, exclusion_type, activation_date).",
        growth="50,000-200,000 rows. Relatively stable; new exclusions added, old ones terminated."
    )

    # =========================================================================
    # DOMAIN 4: REFERENCE TABLES
    # =========================================================================
    doc.add_heading("4. Reference Tables", level=0)
    doc.add_paragraph(
        "Lookup/code tables loaded from CSV files or government reference data. "
        "Updated infrequently (annually or less). Used for joins and display labels."
    )

    ref_tables = [
        ("ref_naics_code",
         "North American Industry Classification System codes. Hierarchical 2-6 digit codes classifying industries. Used to determine small business size standards and filter opportunities by industry.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file (workdir/converted/local database/) via reference_loader.py. CLI: python main.py load reference --type=naics",
         [("naics_code", "VARCHAR(11) PK", "NAICS code (2-6 digits; 11 chars allows exception codes)"),
          ("description", "VARCHAR(500)", "Industry description"),
          ("code_level", "TINYINT", "Hierarchy level (2=sector, 3=subsector, 4=group, 5=industry, 6=national)"),
          ("level_name", "VARCHAR(30)", "Level name label"),
          ("parent_code", "VARCHAR(11)", "Parent NAICS code in the hierarchy"),
          ("year_version", "VARCHAR(4)", "NAICS revision year (e.g., '2022')"),
          ("is_active", "CHAR(1)", "Whether the code is currently active (Y/N)"),
          ("footnote_id", "VARCHAR(5)", "Reference to ref_naics_footnote for size standard exceptions")],
         "Referenced by: entity_naics, opportunity, fpds_contract, usaspending_award, ref_sba_size_standard (naics_code FK). Self-referencing via parent_code.",
         "~2,200 rows. Updated every 5 years with NAICS revisions."),

        ("ref_sba_size_standard",
         "SBA small business size standards by NAICS code. Defines revenue or employee thresholds that determine whether a firm qualifies as 'small' under each NAICS.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py. CLI: python main.py load reference --type=sba",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("naics_code", "VARCHAR(11)", "NAICS code this standard applies to (FK to ref_naics_code)"),
          ("industry_description", "VARCHAR(500)", "Industry description"),
          ("size_standard", "DECIMAL(13,2)", "Threshold value (revenue in millions or employee count)"),
          ("size_type", "CHAR(1)", "Type of standard: 'R' = revenue (millions $), 'E' = employees"),
          ("footnote_id", "VARCHAR(5)", "Reference to special rules/exceptions"),
          ("effective_date", "DATE", "When this standard became effective")],
         "References: ref_naics_code (naics_code FK).",
         "~1,800 rows. Updated when SBA revises size standards (~every 5 years)."),

        ("ref_naics_footnote",
         "Footnotes and exceptions for NAICS size standards. Explains special rules like alternative size standards for certain industries.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("footnote_id", "VARCHAR(5) PK", "Footnote identifier"),
          ("section", "VARCHAR(5) PK", "Section within the footnote"),
          ("description", "TEXT", "Full text of the footnote/exception rule")],
         "Referenced by: ref_naics_code (footnote_id), ref_sba_size_standard (footnote_id).",
         "~50 rows. Rarely changes."),

        ("ref_psc_code",
         "Product Service Codes classifying what the government buys. Hierarchical coding from broad categories down to specific products/services. Temporal: codes have start/end dates.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py. CLI: python main.py load reference --type=psc",
         [("psc_code", "VARCHAR(10) PK", "PSC code (1-4 characters, e.g., 'D306' for IT services)"),
          ("psc_name", "VARCHAR(200)", "Short name of the code"),
          ("start_date", "DATE PK", "Effective start date (part of composite PK for temporal versioning)"),
          ("end_date", "DATE", "Effective end date (NULL = currently active)"),
          ("full_description", "TEXT", "Detailed description of what this code covers"),
          ("psc_includes", "TEXT", "What is included under this code"),
          ("psc_excludes", "TEXT", "What is excluded from this code"),
          ("category_type", "CHAR(1)", "P=Product, S=Service, R=R&D"),
          ("level1_category_code", "VARCHAR(10)", "Top-level category code"),
          ("level1_category", "VARCHAR(100)", "Top-level category name"),
          ("level2_category_code", "VARCHAR(10)", "Second-level category code"),
          ("level2_category", "VARCHAR(100)", "Second-level category name")],
         "Referenced by: opportunity (classification_code), fpds_contract (psc_code), entity_psc (psc_code).",
         "~3,000 rows. Updated periodically by GSA."),

        ("ref_set_aside_type",
         "Federal procurement set-aside type codes. Defines which socioeconomic categories qualify for restricted competition opportunities.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py. CLI: python main.py load reference --type=set_aside",
         [("set_aside_code", "VARCHAR(10) PK", "Set-aside code (e.g., 'SBA', 'WOSB', '8AN', 'HZC')"),
          ("description", "VARCHAR(200)", "Full description of the set-aside type"),
          ("is_small_business", "CHAR(1)", "Whether this set-aside is for small businesses (Y/N)"),
          ("category", "VARCHAR(50)", "Grouping category")],
         "Referenced by: opportunity (set_aside_code), fpds_contract (set_aside_type).",
         "~30 rows. Very stable."),

        ("ref_sba_type",
         "SBA certification program types (8(a), HUBZone, WOSB, EDWOSB, VOSB, SDVOSB). Maps codes to program names.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("sba_type_code", "VARCHAR(10) PK", "SBA program type code"),
          ("description", "VARCHAR(200)", "Program description"),
          ("program_name", "VARCHAR(100)", "Official SBA program name")],
         "Referenced by: entity_sba_certification (sba_type_code).",
         "~10 rows. Very stable."),

        ("ref_business_type",
         "Business type classification codes from SAM.gov. Categorizes entities by ownership, size, and socioeconomic status (e.g., 2X=Woman-Owned, A2=Veteran-Owned).",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("business_type_code", "VARCHAR(4) PK", "Business type code from SAM.gov"),
          ("description", "VARCHAR(200)", "Human-readable description"),
          ("classification", "VARCHAR(50)", "Classification category"),
          ("category", "VARCHAR(50)", "Broader grouping"),
          ("is_socioeconomic", "CHAR(1)", "Whether this is a socioeconomic designation (Y/N)"),
          ("is_small_business_related", "CHAR(1)", "Whether this relates to small business status (Y/N)")],
         "Referenced by: entity_business_type (business_type_code).",
         "~50 rows. Updated when SAM.gov adds new types."),

        ("ref_entity_structure",
         "Legal entity structure codes (LLC, Corporation, Partnership, Sole Proprietorship, etc.).",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("structure_code", "VARCHAR(2) PK", "Two-character structure code"),
          ("description", "VARCHAR(200)", "Full description of the entity structure type")],
         "Referenced by: entity (entity_structure_code).",
         "~20 rows. Very stable."),

        ("ref_country_code",
         "ISO country codes (alpha-2, alpha-3, numeric) with SAM.gov recognition flags.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("three_code", "VARCHAR(3) PK", "ISO 3166-1 alpha-3 code (e.g., 'USA', 'GBR')"),
          ("two_code", "VARCHAR(2)", "ISO 3166-1 alpha-2 code (e.g., 'US', 'GB')"),
          ("country_name", "VARCHAR(100)", "Country name"),
          ("numeric_code", "VARCHAR(4)", "ISO 3166-1 numeric code"),
          ("independent", "VARCHAR(3)", "Whether the country is independent"),
          ("is_iso_standard", "CHAR(1)", "Whether this is an official ISO code (Y/N)"),
          ("sam_gov_recognized", "CHAR(1)", "Whether SAM.gov recognizes this country code (Y/N)")],
         "Referenced by: entity_address (country_code), opportunity (pop_country).",
         "~250 rows. Very stable."),

        ("ref_state_code",
         "US state and territory codes. Maps 2-letter state abbreviations to full names.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py",
         [("state_code", "VARCHAR(2) PK", "Two-letter state/territory code (e.g., 'VA', 'DC')"),
          ("state_name", "VARCHAR(60)", "Full state/territory name"),
          ("country_code", "VARCHAR(3) PK", "Country code (default 'USA')")],
         "Referenced by: opportunity (pop_state), entity_address (state_or_province).",
         "~60 rows. Very stable."),

        ("ref_fips_county",
         "FIPS (Federal Information Processing Standards) county codes. Maps 5-digit FIPS codes to county and state names for geographic analysis.",
         "Python DDL (tables/10_reference.sql)",
         "CSV file via reference_loader.py. CLI: python main.py load reference --type=fips",
         [("fips_code", "VARCHAR(5) PK", "5-digit FIPS county code (state 2 + county 3)"),
          ("county_name", "VARCHAR(100)", "County name"),
          ("state_name", "VARCHAR(60)", "State name")],
         "Referenced by: entity_disaster_response (county_code).",
         "~3,200 rows. Updated when counties merge/split (rare)."),
    ]

    for name, purpose, ownership, data_source, columns, rels, growth in ref_tables:
        add_table_section(doc, name=name, purpose=purpose, ownership=ownership,
            data_source=data_source, columns=columns, relationships=rels,
            change_detection="None. Full reload from CSV files; existing rows replaced.",
            growth=growth)

    # =========================================================================
    # DOMAIN 5: ETL / OPERATIONS TABLES
    # =========================================================================
    doc.add_heading("5. ETL / Operations Tables", level=0)
    doc.add_paragraph(
        "Operational tables supporting the Extract-Transform-Load pipeline. "
        "Track load history, errors, data quality rules, and API rate limits."
    )

    add_table_section(doc,
        name="etl_load_log",
        purpose="Master log of every ETL load operation. Each run of a loader creates one row. Tracks record counts, timing, and success/failure status. Central audit trail for all data loading.",
        ownership="Python DDL (tables/50_etl.sql)",
        data_source="Created by load_manager.py at the start of every loader execution.",
        columns=[
            ("load_id", "INT PK AUTO_INCREMENT", "Unique load identifier; referenced by all loaded records via last_load_id"),
            ("source_system", "VARCHAR(50)", "Data source: 'SAM_OPPORTUNITIES', 'SAM_ENTITIES', 'SAM_AWARDS', 'USASPENDING', 'SAM_EXCLUSIONS', 'SAM_SUBAWARDS', 'FED_HIERARCHY', 'GSA_CALC', 'REFERENCE'"),
            ("load_type", "VARCHAR(20)", "Load type: 'FULL', 'INCREMENTAL', 'DELTA', 'MANUAL'"),
            ("status", "VARCHAR(20)", "Current status: 'RUNNING', 'SUCCESS', 'FAILED', 'PARTIAL'"),
            ("started_at", "DATETIME", "Load start timestamp"),
            ("completed_at", "DATETIME", "Load completion timestamp"),
            ("records_read", "INT", "Number of records read from the source"),
            ("records_inserted", "INT", "New records inserted"),
            ("records_updated", "INT", "Existing records updated (hash changed)"),
            ("records_unchanged", "INT", "Records skipped (hash unchanged)"),
            ("records_errored", "INT", "Records that failed to process"),
            ("error_message", "TEXT", "Error message if the load failed"),
            ("parameters", "JSON", "CLI parameters and configuration used for this load"),
            ("source_file", "VARCHAR(500)", "Source file path (for file-based loads)"),
        ],
        relationships="Referenced by: etl_load_error (load_id FK), all data tables (last_load_id), usaspending_load_checkpoint (load_id FK).",
        change_detection="None (append-only log).",
        growth="5,000-20,000 rows. One row per load execution; ~10-50 loads per week."
    )

    add_table_section(doc,
        name="etl_load_error",
        purpose="Detailed error log for individual records that failed during ETL processing. Each row captures one record-level failure with the raw data for debugging.",
        ownership="Python DDL (tables/50_etl.sql)",
        data_source="Created by loader classes when individual record processing fails.",
        columns=[
            ("id", "BIGINT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("load_id", "INT", "FK to etl_load_log identifying the parent load"),
            ("record_identifier", "VARCHAR(100)", "Business key of the failed record (notice_id, uei, etc.)"),
            ("error_type", "VARCHAR(50)", "Error category: 'VALIDATION', 'PARSE', 'DB', 'API', 'TRANSFORM'"),
            ("error_message", "TEXT", "Detailed error message"),
            ("raw_data", "TEXT", "Raw source data that caused the error (for debugging)"),
            ("created_at", "DATETIME", "When the error occurred"),
        ],
        relationships="References: etl_load_log (load_id FK).",
        change_detection="None (append-only).",
        growth="1,000-50,000 rows. Varies with data quality; ideally low."
    )

    add_table_section(doc,
        name="etl_data_quality_rule",
        purpose="Configurable data quality rules applied during ETL processing. Rules are stored in the database rather than hardcoded, allowing runtime configuration.",
        ownership="Python DDL (tables/50_etl.sql)",
        data_source="Manually configured or seeded by reference_loader.py.",
        columns=[
            ("rule_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("rule_name", "VARCHAR(100)", "Unique rule name/identifier"),
            ("description", "TEXT", "Human-readable description of what the rule checks"),
            ("target_table", "VARCHAR(100)", "Which table the rule applies to"),
            ("target_column", "VARCHAR(100)", "Which column the rule checks (NULL for multi-column rules)"),
            ("rule_type", "VARCHAR(20)", "Rule type: 'NOT_NULL', 'RANGE', 'REGEX', 'LOOKUP', 'CUSTOM'"),
            ("rule_definition", "JSON", "Rule parameters (thresholds, patterns, valid values, etc.)"),
            ("is_active", "CHAR(1)", "Whether the rule is currently active (Y/N)"),
            ("priority", "INT", "Execution priority (lower = higher priority)"),
        ],
        relationships="No direct FK references. Applied programmatically by ETL pipeline.",
        change_detection="None. Static configuration.",
        growth="50-200 rows. Grows as new quality checks are added."
    )

    add_table_section(doc,
        name="etl_rate_limit",
        purpose="Tracks daily API request counts per source system to stay within rate limits. Prevents exceeding SAM.gov and other API quotas.",
        ownership="Python DDL (tables/50_etl.sql)",
        data_source="Updated by BaseAPIClient on each API call.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("source_system", "VARCHAR(50)", "API source system name"),
            ("request_date", "DATE", "Calendar date of the requests"),
            ("requests_made", "INT", "Number of requests made on this date"),
            ("max_requests", "INT", "Maximum allowed requests per day"),
            ("last_request_at", "DATETIME", "Timestamp of the most recent request"),
        ],
        relationships="No FK references.",
        change_detection="None. Upserted daily.",
        growth="365-1,000 rows per year per source. Old rows can be purged."
    )

    add_table_section(doc,
        name="etl_health_snapshot",
        purpose="Periodic snapshots of ETL pipeline health status. Records overall health, alert counts, and stale data source detection for monitoring dashboards.",
        ownership="Python DDL (tables/50_etl.sql)",
        data_source="Created by health_check.py. CLI: python main.py health-check",
        columns=[
            ("snapshot_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("checked_at", "DATETIME", "When the health check was performed"),
            ("overall_status", "VARCHAR(20)", "Aggregate status: 'HEALTHY', 'WARNING', 'CRITICAL'"),
            ("results_json", "JSON", "Detailed per-source health check results"),
            ("alert_count", "INT", "Number of warnings generated"),
            ("error_count", "INT", "Number of errors detected"),
            ("stale_source_count", "INT", "Number of data sources that haven't loaded recently"),
        ],
        relationships="No FK references.",
        change_detection="None (append-only).",
        growth="1,000-5,000 rows per year. One row per health check run."
    )

    add_table_section(doc,
        name="data_load_request",
        purpose="On-demand data loading request queue. Allows the web application to request specific data loads (e.g., load FPDS data for a particular contract). Processed asynchronously by the ETL pipeline.",
        ownership="Python DDL (tables/55_data_load_request.sql)",
        data_source="Created by the C# API when users request data enrichment. Processed by demand_loader.py.",
        columns=[
            ("request_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("request_type", "VARCHAR(30)", "Type of data to load: 'FPDS_CONTRACT', 'ENTITY', 'SUBAWARD'"),
            ("lookup_key", "VARCHAR(200)", "The value to look up (contract number, UEI, etc.)"),
            ("lookup_key_type", "VARCHAR(20)", "Type of lookup key: 'PIID', 'UEI', 'SOLICITATION_NUMBER'"),
            ("status", "VARCHAR(20)", "Request status: 'PENDING', 'IN_PROGRESS', 'COMPLETED', 'FAILED'"),
            ("requested_by", "INT", "FK to app_user who made the request"),
            ("requested_at", "DATETIME", "When the request was submitted"),
            ("started_at", "DATETIME", "When processing began"),
            ("completed_at", "DATETIME", "When processing finished"),
            ("load_id", "INT", "FK to etl_load_log for the resulting load"),
            ("error_message", "TEXT", "Error message if the request failed"),
            ("result_summary", "JSON", "Summary of what was loaded"),
        ],
        relationships="References: app_user (requested_by), etl_load_log (load_id).",
        change_detection="None. State machine progression.",
        growth="1,000-10,000 rows. One per user-initiated data request."
    )

    add_table_section(doc,
        name="usaspending_load_checkpoint",
        purpose="Checkpoint/resume tracking for USASpending bulk CSV loads. Allows large multi-file loads to resume after interruption without reprocessing completed files.",
        ownership="Python DDL (tables/70_usaspending.sql)",
        data_source="Created by usaspending_bulk_loader.py during bulk CSV processing.",
        columns=[
            ("checkpoint_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("load_id", "INT", "FK to etl_load_log"),
            ("fiscal_year", "INT", "Federal fiscal year being loaded"),
            ("csv_file_name", "VARCHAR(255)", "Name of the CSV file being processed"),
            ("status", "ENUM", "Processing status: IN_PROGRESS, COMPLETE, FAILED"),
            ("completed_batches", "INT", "Number of batches successfully loaded"),
            ("total_rows_loaded", "INT", "Total rows loaded from this file"),
            ("archive_hash", "VARCHAR(130)", "SHA-256 hash (first 1MB) + file size for FY deduplication"),
            ("started_at", "TIMESTAMP", "When processing of this file started"),
            ("completed_at", "TIMESTAMP", "When processing of this file finished"),
        ],
        relationships="References: etl_load_log (load_id FK).",
        change_detection="None. Operational state tracking.",
        growth="100-500 rows. One per CSV file per bulk load."
    )

    add_table_section(doc,
        name="ai_usage_log",
        purpose="Tracks AI API usage (token counts, costs) for attachment intel extraction and other AI-powered features. Enables cost monitoring and budget controls.",
        ownership="Python DDL (tables/75_ai_usage.sql)",
        data_source="Created by attachment_intel_loader.py when calling Anthropic API for document analysis.",
        columns=[
            ("usage_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("notice_id", "VARCHAR(100)", "Opportunity notice ID the AI call relates to"),
            ("attachment_id", "INT", "FK to sam_attachment (if analyzing a specific document)"),
            ("model", "VARCHAR(50)", "AI model identifier (e.g. claude-3-haiku-20240307)"),
            ("input_tokens", "INT", "Number of input tokens consumed"),
            ("output_tokens", "INT", "Number of output tokens generated"),
            ("cache_read_tokens", "INT", "Tokens read from Anthropic prompt cache"),
            ("cache_write_tokens", "INT", "Tokens written to Anthropic prompt cache"),
            ("cost_usd", "DECIMAL(10,6)", "Estimated cost in USD"),
            ("requested_by", "INT", "FK to app_user who triggered the analysis"),
            ("created_at", "DATETIME", "Timestamp of the API call"),
        ],
        relationships="References: opportunity (notice_id), sam_attachment (attachment_id), app_user (requested_by).",
        change_detection="None (append-only log).",
        growth="10,000-100,000 rows per year. One row per AI API call."
    )

    # =========================================================================
    # DOMAIN 6: STAGING TABLES
    # =========================================================================
    doc.add_heading("6. Staging Tables", level=0)
    doc.add_paragraph(
        "Raw staging tables preserve complete API/file responses before normalization "
        "into production tables. Enable replay, re-processing, and schema evolution safety."
    )

    stg_tables = [
        ("stg_entity_raw",
         "Raw JSON responses from SAM.gov Entity API before normalization into entity and child tables.",
         [("load_id", "INT", "FK to etl_load_log"),
          ("uei_sam", "VARCHAR(12)", "Entity UEI for lookup"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash of the raw JSON for dedup"),
          ("processed", "CHAR(1)", "Processing status: N=not processed, Y=processed, E=error"),
          ("processed_at", "DATETIME", "When the record was processed into production tables"),
          ("error_message", "TEXT", "Error message if processing failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/20_entity.sql", "entity_loader.py"),

        ("stg_opportunity_raw",
         "Raw JSON responses from SAM.gov Opportunities API before normalization.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("notice_id", "VARCHAR(100)", "Opportunity notice ID for lookup"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash of the raw JSON"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if processing failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "opportunity_loader.py"),

        ("stg_fpds_award_raw",
         "Raw JSON responses from SAM.gov Awards API before normalization into fpds_contract.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("contract_id", "VARCHAR(100)", "Contract PIID"),
          ("modification_number", "VARCHAR(25)", "Modification number"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "awards_loader.py"),

        ("stg_usaspending_raw",
         "Raw JSON responses from USASpending.gov API before normalization.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("award_id", "VARCHAR(100)", "USASpending award ID"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "usaspending_loader.py"),

        ("stg_exclusion_raw",
         "Raw JSON responses from SAM.gov Exclusions API before normalization into sam_exclusion.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("record_id", "VARCHAR(100)", "Exclusion record identifier"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "exclusions_loader.py"),

        ("stg_fedhier_raw",
         "Raw JSON responses from SAM.gov Federal Hierarchy API before normalization into federal_organization.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("fh_org_id", "INT", "Federal hierarchy organization ID"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "fedhier_loader.py"),

        ("stg_subaward_raw",
         "Raw JSON responses from SAM.gov Subaward API before normalization into sam_subaward.",
         [("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
          ("load_id", "INT", "FK to etl_load_log"),
          ("prime_piid", "VARCHAR(50)", "Prime contract PIID"),
          ("sub_uei", "VARCHAR(12)", "Subcontractor UEI"),
          ("raw_json", "JSON", "Complete API response JSON"),
          ("raw_record_hash", "CHAR(64)", "SHA-256 hash"),
          ("processed", "CHAR(1)", "Processing status: N/Y/E"),
          ("error_message", "TEXT", "Error message if failed"),
          ("created_at", "DATETIME", "Row creation timestamp")],
         "tables/80_raw_staging.sql", "subaward_loader.py"),
    ]

    for name, purpose, columns, ddl_file, loader in stg_tables:
        add_table_section(doc, name=name, purpose=purpose,
            ownership=f"Python DDL ({ddl_file})",
            data_source=f"Populated by {loader} during API data loading.",
            columns=columns,
            relationships="References: etl_load_log (load_id). Feeds production tables during normalization.",
            change_detection="raw_record_hash used to skip already-seen records.",
            growth="Varies widely. Can grow to millions of rows; should be periodically purged after successful processing.")

    # =========================================================================
    # DOMAIN 7: APPLICATION / MULTI-TENANT TABLES
    # =========================================================================
    doc.add_heading("7. Application / Multi-Tenant Tables", level=0)
    doc.add_paragraph(
        "Application-level tables supporting the web UI. Multi-tenant via organization_id. "
        "Schema owned by EF Core (C# ASP.NET Core API) except where noted."
    )

    add_table_section(doc,
        name="organization",
        purpose="Top-level tenant entity. Each organization has its own users, prospects, saved searches, and configuration. Includes company profile fields for matching against opportunities.",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core migrations",
        data_source="Created via the web application during registration. Profile fields populated by user input or SAM.gov entity linking.",
        columns=[
            ("organization_id", "INT PK AUTO_INCREMENT", "Surrogate primary key and tenant isolation key"),
            ("name", "VARCHAR(200)", "Organization display name"),
            ("slug", "VARCHAR(100) UNIQUE", "URL-safe identifier (e.g., 'acme-consulting')"),
            ("is_active", "CHAR(1)", "Whether the organization account is active (Y/N)"),
            ("max_users", "INT", "Maximum number of user accounts allowed"),
            ("subscription_tier", "VARCHAR(50)", "Subscription level: trial, basic, professional, enterprise"),
            ("legal_name", "VARCHAR(300)", "Legal business name (for SAM.gov matching)"),
            ("uei_sam", "VARCHAR(13)", "Organization's UEI (links to entity table for self-identification)"),
            ("cage_code", "VARCHAR(5)", "CAGE code"),
            ("ein", "VARCHAR(10)", "Employer Identification Number"),
            ("annual_revenue", "DECIMAL(18,2)", "Annual revenue for size standard determination"),
            ("employee_count", "INT", "Employee count for size standard determination"),
            ("profile_completed", "VARCHAR(1)", "Whether company profile setup is complete (Y/N)"),
            ("created_at", "DATETIME", "Organization creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="Referenced by: app_user, prospect, saved_search, activity_log, organization_invite, organization_entity, organization_certification, organization_naics, organization_past_performance (all via organization_id FK).",
        change_detection="None. Updated via application CRUD operations.",
        growth="10-1,000 rows. One per tenant organization."
    )

    add_table_section(doc,
        name="app_user",
        purpose="Application user accounts. Each user belongs to one organization. Supports role-based access, MFA, and account locking.",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core",
        data_source="Created via web application registration or invitation acceptance.",
        columns=[
            ("user_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization (tenant isolation)"),
            ("username", "VARCHAR(50) UNIQUE", "Login username"),
            ("display_name", "VARCHAR(100)", "Display name shown in the UI"),
            ("email", "VARCHAR(200)", "Email address"),
            ("password_hash", "VARCHAR(255)", "Bcrypt password hash"),
            ("role", "VARCHAR(20)", "System role: USER, ADMIN"),
            ("is_active", "CHAR(1)", "Account active status (Y/N)"),
            ("is_org_admin", "CHAR(1)", "Whether user is an organization administrator (Y/N)"),
            ("mfa_enabled", "CHAR(1)", "Whether MFA is enabled (Y/N)"),
            ("org_role", "VARCHAR(50)", "Organization-level role: member, admin, owner"),
            ("is_system_admin", "TINYINT(1)", "System-wide admin flag (0/1)"),
            ("force_password_change", "CHAR(1)", "Whether user must change password on next login (Y/N)"),
            ("failed_login_attempts", "INT", "Consecutive failed login count (resets on success)"),
            ("locked_until", "DATETIME", "Account locked until this time (NULL = not locked)"),
            ("last_login_at", "DATETIME", "Most recent successful login"),
            ("invited_by", "INT", "FK to app_user who invited this user"),
            ("invite_accepted_at", "DATETIME", "When the user accepted their invitation"),
            ("created_at", "DATETIME", "Account creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="References: organization (organization_id FK), app_user (invited_by self-FK). Referenced by: prospect (assigned_to, capture_manager_id), prospect_note (user_id), saved_search (user_id), app_session (user_id), notification (user_id), activity_log (user_id), organization_invite (invited_by).",
        change_detection="None. Updated via application CRUD.",
        growth="50-5,000 rows. ~5-10 users per organization."
    )

    add_table_section(doc,
        name="app_session",
        purpose="Active user sessions with JWT token tracking. Supports token refresh and revocation for security. httpOnly cookie-based authentication.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by the authentication service on login. Updated on token refresh.",
        columns=[
            ("session_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("user_id", "INT", "FK to app_user"),
            ("token_hash", "CHAR(64) UNIQUE", "SHA-256 hash of the JWT access token"),
            ("refresh_token_hash", "CHAR(64)", "SHA-256 hash of the refresh token"),
            ("issued_at", "DATETIME", "When the session was created"),
            ("expires_at", "DATETIME", "When the access token expires"),
            ("revoked_at", "DATETIME", "When the session was explicitly revoked (NULL = active)"),
            ("revoked_reason", "VARCHAR(100)", "Why the session was revoked (logout, password change, admin action)"),
            ("ip_address", "VARCHAR(45)", "Client IP address (supports IPv6)"),
            ("user_agent", "VARCHAR(500)", "Browser/client user agent string"),
        ],
        relationships="References: app_user (user_id FK).",
        change_detection="None. Created/revoked by auth service.",
        growth="1,000-50,000 rows. Should be periodically purged of expired sessions."
    )

    add_table_section(doc,
        name="prospect",
        purpose="Core sales pipeline entity. Links an opportunity to an organization's capture/pursuit effort. Tracks status (NEW through WON/LOST), priority, assigned team members, and financial estimates.",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core",
        data_source="Created by users via the web UI or automatically by saved search auto-prospecting.",
        columns=[
            ("prospect_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization (tenant isolation)"),
            ("source", "VARCHAR(20)", "How the prospect was created: MANUAL, SAVED_SEARCH, AUTO"),
            ("notice_id", "VARCHAR(100)", "FK to opportunity being pursued"),
            ("assigned_to", "INT", "FK to app_user responsible for this pursuit"),
            ("capture_manager_id", "INT", "FK to app_user serving as capture manager"),
            ("status", "VARCHAR(30)", "Pipeline status: NEW, REVIEWING, GO, NO_GO, BID_SUBMITTED, WON, LOST, CANCELLED"),
            ("proposal_status", "VARCHAR(20)", "Proposal preparation status"),
            ("priority", "VARCHAR(10)", "Priority level: LOW, MEDIUM, HIGH, CRITICAL"),
            ("decision_date", "DATE", "Go/no-go decision deadline"),
            ("bid_submitted_date", "DATE", "When the bid/proposal was submitted"),
            ("estimated_value", "DECIMAL(15,2)", "Estimated contract value"),
            ("estimated_effort_hours", "DECIMAL(10,2)", "Estimated level of effort in hours"),
            ("win_probability", "DECIMAL(5,2)", "Probability of win (0-100%)"),
            ("go_no_go_score", "DECIMAL(5,2)", "Calculated go/no-go evaluation score"),
            ("teaming_required", "CHAR(1)", "Whether teaming/JV is needed (Y/N)"),
            ("outcome", "VARCHAR(20)", "Final outcome: WON, LOST, CANCELLED, WITHDRAWN"),
            ("outcome_date", "DATE", "Date of the outcome"),
            ("outcome_notes", "TEXT", "Notes about the outcome"),
            ("created_at", "DATETIME", "Creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="References: organization (organization_id FK), opportunity (notice_id FK), app_user (assigned_to, capture_manager_id FKs). Referenced by: prospect_note (prospect_id), prospect_team_member (prospect_id), proposal (prospect_id).",
        change_detection="None. Updated via application CRUD.",
        growth="100-10,000 rows per organization. ~50-500 active prospects at any time."
    )

    add_table_section(doc,
        name="prospect_note",
        purpose="Timeline of notes, comments, and status changes on a prospect. Provides collaboration history for the capture team.",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core",
        data_source="Created by users via the web UI.",
        columns=[
            ("note_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("prospect_id", "INT", "FK to prospect"),
            ("user_id", "INT", "FK to app_user who created the note"),
            ("note_type", "VARCHAR(30)", "Type: COMMENT, STATUS_CHANGE, GO_NO_GO, INTEL, ACTION_ITEM"),
            ("note_text", "TEXT", "Note content"),
            ("created_at", "DATETIME", "When the note was created"),
        ],
        relationships="References: prospect (prospect_id FK), app_user (user_id FK).",
        change_detection="None (append-only).",
        growth="500-50,000 rows per organization. ~5-20 notes per prospect."
    )

    add_table_section(doc,
        name="prospect_team_member",
        purpose="Team members assigned to a prospect pursuit, including internal users and external teaming partners (identified by UEI).",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core",
        data_source="Created by users via the web UI.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("prospect_id", "INT", "FK to prospect"),
            ("uei_sam", "VARCHAR(12)", "UEI of teaming partner entity (NULL for internal team members)"),
            ("app_user_id", "INT", "FK to app_user (NULL for external partners)"),
            ("role", "VARCHAR(50)", "Team member role (e.g., 'Prime', 'Sub', 'Consultant', 'SME')"),
            ("notes", "TEXT", "Notes about this team member's involvement"),
            ("proposed_hourly_rate", "DECIMAL(10,2)", "Proposed billing rate"),
            ("commitment_pct", "DECIMAL(5,2)", "Percentage of time commitment"),
        ],
        relationships="References: prospect (prospect_id FK), app_user (app_user_id FK). Cross-references entity (uei_sam).",
        change_detection="None. CRUD via application.",
        growth="200-5,000 rows per organization."
    )

    add_table_section(doc,
        name="saved_search",
        purpose="User-defined saved search criteria for finding opportunities. Supports notification on new matches and automatic prospect creation for high-scoring opportunities.",
        ownership="Python DDL (tables/60_prospecting.sql), managed by EF Core",
        data_source="Created by users via the web UI.",
        columns=[
            ("search_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization (tenant isolation)"),
            ("user_id", "INT", "FK to app_user who created the search"),
            ("search_name", "VARCHAR(100)", "User-assigned name for the saved search"),
            ("description", "TEXT", "Description of what this search looks for"),
            ("filter_criteria", "JSON", "Serialized search filters (NAICS, set-aside, keywords, agencies, etc.)"),
            ("notification_enabled", "CHAR(1)", "Whether to notify user of new matching opportunities (Y/N)"),
            ("is_active", "CHAR(1)", "Whether the saved search is active (Y/N)"),
            ("last_run_at", "DATETIME", "When the search was last executed"),
            ("last_new_results", "INT", "Number of new results found in the last run"),
            ("auto_prospect_enabled", "CHAR(1)", "Whether to auto-create prospects from matches (Y/N)"),
            ("min_pwin_score", "DECIMAL(5,2)", "Minimum pWin score for auto-prospect creation"),
            ("auto_assign_to", "INT", "FK to app_user to auto-assign created prospects to"),
            ("created_at", "DATETIME", "Creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="References: organization (organization_id FK), app_user (user_id, auto_assign_to FKs).",
        change_detection="None. CRUD via application.",
        growth="50-500 rows per organization."
    )

    add_table_section(doc,
        name="notification",
        purpose="User notifications for events like new matching opportunities, prospect status changes, and system alerts. Supports read/unread tracking.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by application services (saved search runner, prospect status changes, etc.).",
        columns=[
            ("notification_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("user_id", "INT", "FK to app_user who should see this notification"),
            ("notification_type", "VARCHAR(50)", "Type: NEW_MATCH, STATUS_CHANGE, DEADLINE_APPROACHING, SYSTEM_ALERT"),
            ("title", "VARCHAR(200)", "Notification title/headline"),
            ("message", "TEXT", "Full notification message body"),
            ("entity_type", "VARCHAR(50)", "Type of related entity (e.g., 'opportunity', 'prospect')"),
            ("entity_id", "VARCHAR(100)", "ID of the related entity for deep-linking"),
            ("is_read", "CHAR(1)", "Read status (Y/N)"),
            ("created_at", "DATETIME", "When the notification was created"),
            ("read_at", "DATETIME", "When the user read/dismissed the notification"),
        ],
        relationships="References: app_user (user_id FK).",
        change_detection="None. Append-only with read status update.",
        growth="1,000-100,000 rows per organization. Should be periodically purged."
    )

    add_table_section(doc,
        name="activity_log",
        purpose="Audit trail of user actions in the application. Records who did what, when, and from where. Used for security auditing and usage analytics.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by the C# API middleware and service layer on significant user actions.",
        columns=[
            ("activity_id", "BIGINT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization (tenant isolation)"),
            ("user_id", "INT", "FK to app_user who performed the action (NULL for system actions)"),
            ("action", "VARCHAR(50)", "Action performed: CREATE, UPDATE, DELETE, LOGIN, SEARCH, EXPORT"),
            ("entity_type", "VARCHAR(50)", "Type of entity acted upon: prospect, saved_search, user, etc."),
            ("entity_id", "VARCHAR(100)", "ID of the entity acted upon"),
            ("details", "JSON", "Additional action details (changed fields, search criteria, etc.)"),
            ("ip_address", "VARCHAR(45)", "Client IP address"),
            ("created_at", "DATETIME", "When the action occurred"),
        ],
        relationships="References: organization (organization_id FK), app_user (user_id FK).",
        change_detection="None (append-only audit log).",
        growth="10,000-1,000,000+ rows. High-volume; should be periodically archived."
    )

    add_table_section(doc,
        name="organization_invite",
        purpose="Pending invitations for new users to join an organization. Contains a unique invite code with expiration for secure onboarding.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by org admins via the web UI.",
        columns=[
            ("invite_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization being joined"),
            ("email", "VARCHAR(255)", "Email address the invitation was sent to"),
            ("invite_code", "VARCHAR(64) UNIQUE", "Unique invitation code (used in the invite URL)"),
            ("org_role", "VARCHAR(50)", "Role the invitee will be assigned: member, admin"),
            ("invited_by", "INT", "FK to app_user who sent the invitation"),
            ("expires_at", "DATETIME", "When the invitation expires"),
            ("accepted_at", "DATETIME", "When the invitation was accepted (NULL = pending)"),
            ("created_at", "DATETIME", "When the invitation was created"),
        ],
        relationships="References: organization (organization_id FK), app_user (invited_by FK).",
        change_detection="None. State transitions: created -> accepted or expired.",
        growth="100-1,000 rows. Low volume."
    )

    add_table_section(doc,
        name="organization_entity",
        purpose="Links organizations to entities of interest (competitors, partners, teaming candidates, self). Supports intelligence tracking and relationship management.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users via the web UI or during company profile setup.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization"),
            ("uei_sam", "VARCHAR(12)", "FK to entity being tracked"),
            ("partner_uei", "VARCHAR(13)", "UEI used for JV partnership filings (if applicable)"),
            ("relationship", "VARCHAR(20)", "Relationship type: SELF, COMPETITOR, PARTNER, TEAMING, WATCH"),
            ("is_active", "CHAR(1)", "Whether the tracking is active (Y/N)"),
            ("added_by", "INT", "FK to app_user who added this link"),
            ("notes", "TEXT", "Notes about the relationship"),
            ("created_at", "DATETIME", "Creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="References: organization (organization_id FK), entity (uei_sam FK), app_user (added_by FK).",
        change_detection="None. CRUD via application.",
        growth="10-500 rows per organization."
    )

    add_table_section(doc,
        name="organization_certification",
        purpose="Certifications held by the organization (WOSB, 8(a), HUBZone, ISO, CMMI, etc.). Used for qualification matching against opportunity requirements.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users during company profile setup or synced from SAM.gov entity data.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization"),
            ("certification_type", "VARCHAR(50)", "Certification type (e.g., 'WOSB', '8A', 'HUBZONE', 'ISO_9001')"),
            ("certifying_agency", "VARCHAR(100)", "Agency that issued the certification"),
            ("certification_number", "VARCHAR(100)", "Certification reference number"),
            ("expiration_date", "DATETIME", "When the certification expires"),
            ("is_active", "VARCHAR(1)", "Whether currently active (Y/N)"),
            ("source", "VARCHAR(20)", "How it was added: MANUAL, SAM_SYNC"),
            ("created_at", "DATETIME", "Creation timestamp"),
        ],
        relationships="References: organization (organization_id FK) with CASCADE delete.",
        change_detection="None. CRUD via application.",
        growth="5-20 rows per organization."
    )

    add_table_section(doc,
        name="organization_naics",
        purpose="NAICS codes the organization operates under. Used for pWin scoring and opportunity matching to determine if the org works in the opportunity's industry.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users during company profile setup.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization"),
            ("naics_code", "VARCHAR(11)", "NAICS code (should match ref_naics_code)"),
            ("is_primary", "VARCHAR(1)", "Whether this is the org's primary NAICS (Y/N)"),
            ("size_standard_met", "VARCHAR(1)", "Whether org meets the small business size standard for this NAICS (Y/N)"),
            ("created_at", "DATETIME", "Creation timestamp"),
        ],
        relationships="References: organization (organization_id FK) with CASCADE delete. Cross-references ref_naics_code.",
        change_detection="None. CRUD via application.",
        growth="5-30 rows per organization."
    )

    add_table_section(doc,
        name="organization_past_performance",
        purpose="Past contract performance records for the organization. Used for qualification assessment and pWin scoring -- demonstrates relevant experience.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users via the web UI or imported from FPDS data.",
        columns=[
            ("id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("organization_id", "INT", "FK to organization"),
            ("contract_number", "VARCHAR(50)", "Contract/PIID number"),
            ("agency_name", "VARCHAR(200)", "Contracting agency name"),
            ("description", "TEXT", "Description of work performed"),
            ("naics_code", "VARCHAR(11)", "NAICS code of the contract"),
            ("contract_value", "DECIMAL(18,2)", "Total contract value"),
            ("period_start", "DATETIME", "Contract start date"),
            ("period_end", "DATETIME", "Contract end date"),
            ("created_at", "DATETIME", "Creation timestamp"),
        ],
        relationships="References: organization (organization_id FK) with CASCADE delete.",
        change_detection="None. CRUD via application.",
        growth="10-100 rows per organization."
    )

    add_table_section(doc,
        name="proposal",
        purpose="Formal proposal records linked to prospects. Tracks proposal lifecycle from draft through submission, with deadline tracking and outcome recording.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users via the web UI when a prospect moves to active pursuit.",
        columns=[
            ("proposal_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("prospect_id", "INT", "FK to prospect (1:1 relationship)"),
            ("proposal_number", "VARCHAR(50)", "Internal proposal tracking number"),
            ("submission_deadline", "DATETIME", "Deadline for proposal submission"),
            ("submitted_at", "DATETIME", "When the proposal was actually submitted"),
            ("proposal_status", "VARCHAR(20)", "Status: DRAFT, IN_REVIEW, APPROVED, SUBMITTED, WITHDRAWN"),
            ("estimated_value", "DECIMAL(15,2)", "Estimated value of the proposed contract"),
            ("win_probability_pct", "DECIMAL(5,2)", "Win probability at time of submission (0-100%)"),
            ("lessons_learned", "TEXT", "Post-outcome lessons learned"),
            ("created_at", "DATETIME", "Creation timestamp"),
            ("updated_at", "DATETIME", "Last update timestamp"),
        ],
        relationships="References: prospect (prospect_id FK, unique constraint = 1:1). Referenced by: proposal_document (proposal_id), proposal_milestone (proposal_id).",
        change_detection="None. CRUD via application.",
        growth="50-2,000 rows per organization."
    )

    add_table_section(doc,
        name="proposal_document",
        purpose="Documents uploaded as part of a proposal (technical volume, cost volume, past performance, etc.). Tracks file metadata and upload history.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Uploaded by users via the web UI.",
        columns=[
            ("document_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("proposal_id", "INT", "FK to proposal"),
            ("document_type", "VARCHAR(50)", "Type: TECHNICAL, COST, PAST_PERFORMANCE, MANAGEMENT, OTHER"),
            ("file_name", "VARCHAR(255)", "Original uploaded filename"),
            ("file_path", "VARCHAR(500)", "Storage path on the server"),
            ("file_size_bytes", "BIGINT", "File size in bytes"),
            ("uploaded_by", "INT", "FK to app_user who uploaded"),
            ("uploaded_at", "DATETIME", "Upload timestamp"),
            ("notes", "TEXT", "Notes about this document version"),
        ],
        relationships="References: proposal (proposal_id FK with CASCADE delete), app_user (uploaded_by FK).",
        change_detection="None. Append-only (new versions uploaded as new rows).",
        growth="100-5,000 rows per organization."
    )

    add_table_section(doc,
        name="proposal_milestone",
        purpose="Milestones and tasks within a proposal effort. Tracks deadlines, assignments, and completion status for proposal preparation workflow.",
        ownership="Python DDL (tables/90_web_api.sql), managed by EF Core",
        data_source="Created by users via the web UI.",
        columns=[
            ("milestone_id", "INT PK AUTO_INCREMENT", "Surrogate primary key"),
            ("proposal_id", "INT", "FK to proposal"),
            ("milestone_name", "VARCHAR(100)", "Name of the milestone (e.g., 'Draft Technical Volume', 'Pink Team Review')"),
            ("due_date", "DATE", "When the milestone is due"),
            ("completed_date", "DATE", "When the milestone was actually completed"),
            ("assigned_to", "INT", "FK to app_user responsible for this milestone"),
            ("status", "VARCHAR(20)", "Status: PENDING, IN_PROGRESS, COMPLETED, OVERDUE"),
            ("notes", "TEXT", "Notes about the milestone"),
            ("created_at", "DATETIME", "Creation timestamp"),
        ],
        relationships="References: proposal (proposal_id FK with CASCADE delete), app_user (assigned_to FK).",
        change_detection="None. CRUD via application.",
        growth="200-5,000 rows per organization."
    )

    # =========================================================================
    # FEDERAL HIERARCHY (standalone in domain 2 but fits better with awards)
    # =========================================================================
    # Already covered federal_organization? Let me add it.

    # Insert federal_organization into the awards section -- actually, let me add it as a standalone
    # It wasn't covered yet. Let me add it to the end.

    doc.add_heading("Appendix: Federal Hierarchy", level=0)
    doc.add_paragraph(
        "Federal organizational hierarchy from SAM.gov. Maps departments, agencies, "
        "sub-agencies, and offices."
    )

    add_table_section(doc,
        name="federal_organization",
        purpose="Federal government organizational hierarchy from SAM.gov Federal Hierarchy API. Maps the tree structure from departments down to individual contracting offices. Used for agency filtering and organizational context.",
        ownership="Python DDL (tables/40_federal.sql)",
        data_source="SAM.gov Federal Hierarchy API via fedhier_loader.py. CLI: python main.py load fedhier",
        columns=[
            ("fh_org_id", "INT PK", "SAM.gov Federal Hierarchy organization ID (not auto-increment)"),
            ("fh_org_name", "VARCHAR(500)", "Organization name"),
            ("fh_org_type", "VARCHAR(50)", "Type: DEPARTMENT, AGENCY, SUB_AGENCY, OFFICE, SUB_OFFICE"),
            ("description", "TEXT", "Organization description"),
            ("status", "VARCHAR(20)", "Active or inactive status"),
            ("agency_code", "VARCHAR(20)", "FPDS agency code"),
            ("oldfpds_office_code", "VARCHAR(20)", "Legacy FPDS office code"),
            ("cgac", "VARCHAR(10)", "Common Government-wide Accounting Classification code"),
            ("parent_org_id", "INT", "FK to parent federal_organization (self-referencing hierarchy)"),
            ("level", "INT", "Depth level in the hierarchy (1=department, 2=agency, etc.)"),
            ("created_date", "DATE", "When the org was created in the federal hierarchy"),
            ("last_modified_date", "DATE", "When last modified"),
            ("record_hash", "CHAR(64)", "SHA-256 hash for change detection"),
            ("first_loaded_at", "DATETIME", "Initial load timestamp"),
            ("last_loaded_at", "DATETIME", "Most recent update timestamp"),
            ("last_load_id", "INT", "FK to etl_load_log"),
        ],
        relationships="Self-referencing: parent_org_id -> fh_org_id. Cross-referenced by opportunity (department_name, office) and fpds_contract (agency_id, contracting_office_id).",
        change_detection="SHA-256 record_hash.",
        growth="5,000-15,000 rows. Relatively stable; changes when government reorganizes."
    )

    # -- Save --
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Data dictionary saved to: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE):,} bytes")


if __name__ == "__main__":
    build_document()
