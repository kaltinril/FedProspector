#!/usr/bin/env python3
"""Generate FedProspect System Administration Guide as a Word document."""

from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT = Path(__file__).parent / "FedProspect-Admin-Guide.docx"


def add_code_block(doc, text):
    """Add a formatted code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Code"]
    p.add_run(text)


def setup_styles(doc):
    """Configure document styles."""
    # Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Code style
    code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = "Consolas"
    code_font.size = Pt(9)
    code_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    code_fmt = code_style.paragraph_format
    code_fmt.space_before = Pt(4)
    code_fmt.space_after = Pt(4)
    code_fmt.left_indent = Inches(0.3)

    # Heading styles
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_title_page(doc):
    """Add title page."""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FedProspect\nSystem Administration Guide")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Version 1.0 - {date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Federal Contract Prospecting System\n"
        "WOSB and 8(a) Contract Discovery Platform"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_toc(doc):
    """Add a table of contents placeholder."""
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. System Overview",
        "2. Service Management",
        "3. Daily Data Loading",
        "4. CLI Command Reference",
        "5. Organization Management",
        "6. User Management",
        "7. Data Management",
        "8. Monitoring & Troubleshooting",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()


def ch1_system_overview(doc):
    """Chapter 1: System Overview."""
    doc.add_heading("1. System Overview", level=1)

    doc.add_heading("1.1 Architecture", level=2)
    doc.add_paragraph(
        "FedProspect is a four-tier system for discovering WOSB and 8(a) "
        "federal contract opportunities. Data flows through the following pipeline:"
    )

    arch_text = (
        "Government APIs          Python ETL           MySQL 8.4          C# API            React UI\n"
        "(SAM.gov, USASpending,   (fed_prospector/)    (fed_contracts     (ASP.NET Core)    (Vite + React 19\n"
        " GSA CALC+, FedHier)     Load, Transform,      database)        REST endpoints     + MUI v7)\n"
        "                          Validate                               httpOnly cookies\n"
        "                                                                 CSRF protection\n"
        "\n"
        "  [External APIs] ---> [Python CLI] ---> [MySQL] ---> [C# API] ---> [React UI]\n"
        "                        port N/A          3306          5056          5173"
    )
    add_code_block(doc, arch_text)

    doc.add_heading("1.2 Component Locations", level=2)
    table = doc.add_table(rows=7, cols=3, style="Light Grid Accent 1")
    headers = ["Component", "Location", "Default Port"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ("Python ETL / CLI", "fed_prospector/", "N/A"),
        ("MySQL Database", "E:\\mysql (standalone)", "3306"),
        ("C# Web API", "api/src/FedProspector.Api/", "5056"),
        ("React UI", "ui/", "5173"),
        ("Service Manager", "fed_prospector.py", "N/A"),
        ("Configuration", "fed_prospector/.env", "N/A"),
    ]
    for row_idx, (comp, loc, port) in enumerate(data, 1):
        table.rows[row_idx].cells[0].text = comp
        table.rows[row_idx].cells[1].text = loc
        table.rows[row_idx].cells[2].text = port

    doc.add_paragraph()

    doc.add_heading("1.3 Technology Stack", level=2)
    doc.add_paragraph(
        "Python 3.14 for all ETL and data gathering. MySQL 8.4 LTS with InnoDB engine "
        "and utf8mb4 charset. ASP.NET Core Web API with httpOnly cookie authentication "
        "and CSRF protection. Vite 8 + React 19 + TypeScript + MUI v7 frontend."
    )
    doc.add_page_break()


def ch2_service_management(doc):
    """Chapter 2: Service Management."""
    doc.add_heading("2. Service Management", level=1)

    doc.add_heading("2.1 Service Manager Overview", level=2)
    doc.add_paragraph(
        "The service manager (fed_prospector.py) controls all three services: "
        "MySQL database, C# API, and React UI. It is invoked via the batch wrapper "
        "or directly with Python."
    )

    doc.add_heading("2.2 Command Syntax", level=2)
    add_code_block(doc, "fed_prospector.bat <command> [service]")
    doc.add_paragraph("Or equivalently:")
    add_code_block(doc, "python fed_prospector.py <command> [service]")

    doc.add_heading("2.3 Commands", level=2)
    table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    table.rows[0].cells[0].text = "Command"
    table.rows[0].cells[1].text = "Description"
    cmds = [
        ("start [all|db|api|ui]", "Start services. Default: all (db, then api, then ui)."),
        ("stop [all|db|api|ui]", "Stop services. Default: all (ui, then api, then db). Use --force/-f for force-kill."),
        ("restart [all|db|api|ui]", "Stop then start. Supports --force/-f flag."),
        ("status [all]", "Show running/stopped status of all services with ports."),
        ("build [all|api|ui]", "Build API (.NET) and/or UI (npm). Stops, builds, restarts."),
    ]
    for i, (cmd, desc) in enumerate(cmds, 1):
        table.rows[i].cells[0].text = cmd
        table.rows[i].cells[1].text = desc

    doc.add_heading("2.4 Starting the Full Stack", level=2)
    add_code_block(doc, "fed_prospector.bat start all")
    doc.add_paragraph(
        "This starts MySQL first (waits up to 30s for ping), then the .NET API "
        "(waits up to 60s for /health), then the Vite dev server (waits for port)."
    )

    doc.add_heading("2.5 Starting Individual Services", level=2)
    add_code_block(doc, "fed_prospector.bat start db      # Start MySQL only")
    add_code_block(doc, "fed_prospector.bat start api     # Start C# API only")
    add_code_block(doc, "fed_prospector.bat start ui      # Start React UI only")

    doc.add_heading("2.6 Checking Status", level=2)
    add_code_block(doc, "fed_prospector.bat status")
    doc.add_paragraph("Example output:")
    add_code_block(doc,
        "  FedProspector Service Status\n"
        "  ============================\n"
        "  [DB]  Running  (port 3306)\n"
        "  [API] Running  (http://localhost:5056/swagger)\n"
        "  [UI]  Running  (http://localhost:5173)"
    )

    doc.add_heading("2.7 Stopping Services", level=2)
    add_code_block(doc, "fed_prospector.bat stop all          # Graceful shutdown")
    add_code_block(doc, "fed_prospector.bat stop db --force   # Force-kill MySQL")
    doc.add_paragraph(
        "Graceful stop sends mysqladmin shutdown to MySQL. The --force flag uses "
        "taskkill /F instead. The API and UI are always force-killed (no graceful "
        "shutdown protocol)."
    )

    doc.add_heading("2.8 Building", level=2)
    add_code_block(doc, "fed_prospector.bat build all    # Build API + UI, then restart")
    add_code_block(doc, "fed_prospector.bat build api    # Build .NET API only")
    add_code_block(doc, "fed_prospector.bat build ui     # Build React UI only (npm run build)")
    doc.add_paragraph(
        "The build command stops the target service, runs the build, then restarts it."
    )

    doc.add_heading("2.9 Port Configuration", level=2)
    doc.add_paragraph(
        "Default ports are configured via environment variables in fed_prospector/.env:"
    )
    add_code_block(doc,
        "DB_PORT=3306        # MySQL\n"
        "API_PORT=5056       # ASP.NET Core API\n"
        "UI_PORT=5173        # Vite dev server"
    )
    doc.add_page_break()


def ch3_daily_loading(doc):
    """Chapter 3: Daily Data Loading."""
    doc.add_heading("3. Daily Data Loading", level=1)

    doc.add_heading("3.1 Overview", level=2)
    doc.add_paragraph(
        "The daily load script (daily_load.bat) runs a 13-step pipeline that fetches "
        "data from government APIs, backfills descriptions, downloads attachments, "
        "extracts text, and performs intelligence extraction. It is designed to run "
        "once per day."
    )

    doc.add_heading("3.2 Automated Batch Loads", level=2)
    doc.add_paragraph(
        "The CLI also provides structured batch commands that run job sequences with "
        "freshness checks (skipping recently-loaded data):"
    )
    add_code_block(doc,
        "python main.py load daily              # Opportunities + awards + saved searches\n"
        "python main.py load daily --full       # Full: entities, opps, awards, subawards, hierarchy\n"
        "python main.py load weekly             # Entities, hierarchy, awards, exclusions, subawards\n"
        "python main.py load monthly            # All sources including calc rates and USASpending\n"
        "\n"
        "Options:\n"
        "  --key 1|2                  SAM.gov API key (default: 2)\n"
        "  --skip JOB                 Skip specific jobs (repeatable)\n"
        "  --dry-run                  Preview without executing\n"
        "  --force                    Override freshness check\n"
        "  --continue-on-failure      Don't stop on first failure"
    )

    doc.add_heading("3.3 daily_load.bat Step-by-Step", level=2)

    steps = [
        (
            "Step 1: Load Opportunities",
            "python main.py load opportunities --max-calls 300 --key 2 --days-back 31 --force",
            "Fetches contract opportunities posted in the last 31 days from SAM.gov "
            "Opportunities API using key 2 (1000/day limit). Uses up to 300 API calls. "
            "The --force flag starts a fresh load ignoring any previous partial progress.",
            "2-10 minutes depending on volume."
        ),
        (
            "Step 2: Fetch Missing Descriptions",
            "python main.py update fetch-descriptions --key 2 --limit 100 --naics <NAICS> --set-aside WOSB,8A,SBA",
            "Backfills opportunity description text from SAM.gov. Uses a two-pass "
            "prioritized approach: first fetches descriptions for high-value opportunities "
            "(matching target NAICS codes and WOSB/8(a)/SBA set-asides), then uses remaining "
            "budget for the general backlog. Limited to 100 descriptions per daily run to "
            "stay within Key 2's 1,000/day quota (shared with other steps).",
            "1-5 minutes."
        ),
        (
            "Step 3: Load USASpending Bulk",
            "python main.py load usaspending-bulk --days-back 5",
            "Downloads bulk award data from USASpending.gov for the last 5 days. "
            "No API rate limits (USASpending has no daily quotas). Loads contract "
            "awards with obligation amounts and vendor information.",
            "5-15 minutes depending on data volume."
        ),
        (
            "Step 4: Load Awards - 8(a) Set-Aside",
            "python main.py load awards --naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside 8a",
            "Fetches 8(a) set-aside contract awards from SAM.gov for the configured "
            "NAICS codes (24 codes covering IT, consulting, logistics, etc.) signed "
            "in the last 10 days.",
            "5-15 minutes."
        ),
        (
            "Step 5: Load Awards - WOSB Set-Aside",
            "python main.py load awards --naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside WOSB",
            "Same as Step 4 but for WOSB (Women-Owned Small Business) set-aside awards.",
            "5-15 minutes."
        ),
        (
            "Step 6: Load Awards - SBA Set-Aside",
            "python main.py load awards --naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside SBA",
            "Same as Step 4 but for SBA (Small Business Administration) set-aside awards.",
            "5-15 minutes."
        ),
        (
            "Step 7: Enrich Resource Link Metadata",
            "python main.py update link-metadata",
            "HEAD-requests SAM.gov resource link URLs to extract filenames and content "
            "types. Updates the resource_links JSON column with enriched data. This enables "
            "proper file naming during attachment downloads.",
            "1-5 minutes."
        ),
        (
            "Step 8: Download Attachments",
            "python main.py download attachments --missing-only --active-only --batch-size 5000",
            "Downloads attachment files (PDFs, DOCXs, etc.) for active opportunities. "
            "The --missing-only flag skips already-downloaded files. The --active-only flag "
            "limits to opportunities with future response deadlines. Files are stored in the "
            "attachment directory (E:\\fedprospector\\attachments\\).",
            "10-30 minutes depending on volume."
        ),
        (
            "Step 9: Extract Attachment Text",
            "python main.py extract attachment-text --batch-size 5000 --workers 10",
            "Parses downloaded documents (PDF, DOCX, DOC, XLSX, etc.) to extract raw "
            "text content. Uses parallel processing with multiple worker threads. "
            "Extracted text is stored in the database for downstream intelligence extraction.",
            "5-20 minutes."
        ),
        (
            "Step 10: Extract Attachment Intelligence",
            "python main.py extract attachment-intel --batch-size 5000",
            "Analyzes extracted text using keyword/heuristic methods to identify "
            "requirements, evaluation criteria, set-aside details, incumbent information, "
            "and other bid-relevant intelligence.",
            "2-10 minutes."
        ),
        (
            "Step 11: AI Analysis (Disabled)",
            "python main.py extract attachment-ai --model haiku --batch-size 50",
            "Runs Claude AI deep analysis on attachment text. Currently disabled in "
            "daily_load.bat to save on costs. Can be run on demand for specific "
            "opportunities when needed.",
            "5-30 minutes (when enabled)."
        ),
        (
            "Step 12: Backfill Opportunity Intel",
            "python main.py backfill opportunity-intel",
            "Backfills opportunity-level intelligence columns from attachment analysis "
            "results. Aggregates extracted data across all attachments for each opportunity.",
            "1-5 minutes."
        ),
        (
            "Step 13: Clean Up Attachment Files",
            "python main.py maintain attachment-files",
            "Removes original attachment files from disk after they have completed the "
            "full analysis pipeline (downloaded, text extracted, keyword intel extracted, "
            "and AI analyzed). All extracted data remains in the database. Only files that "
            "passed ALL pipeline stages are eligible for cleanup.",
            "1-2 minutes."
        ),
    ]

    for title, cmd, desc, duration in steps:
        doc.add_heading(title, level=3)
        add_code_block(doc, cmd)
        doc.add_paragraph(desc)
        p = doc.add_paragraph()
        run = p.add_run(f"Expected duration: {duration}")
        run.italic = True

    doc.add_heading("3.4 NAICS Codes", level=2)
    doc.add_paragraph(
        "The daily_load.bat script uses 24 NAICS codes covering the organization's "
        "target industries. These are set in the NAICS variable at the top of the script:"
    )
    add_code_block(doc,
        "336611, 488190, 519210, 541219, 541330, 541511, 541512, 541513,\n"
        "541519, 541611, 541612, 541613, 541690, 541990, 561110, 561210,\n"
        "561510, 561990, 611430, 611710, 621111, 621399, 624190, 812910"
    )

    doc.add_heading("3.5 Common Issues", level=2)
    doc.add_paragraph(
        "Rate limiting: If you see '429' errors or 'Rate limit reached', the daily "
        "SAM.gov API quota has been exhausted. Use --key 2 for the 1000/day tier. "
        "Progress is saved automatically; re-run the same command to resume."
    )
    doc.add_paragraph(
        "Partial loads: All load commands save progress after every page. If interrupted "
        "(Ctrl+C, network failure, etc.), re-run the same command to continue from "
        "where it left off. Use --force to start over."
    )
    doc.add_paragraph(
        "Disk space: Attachment downloads consume disk space in E:\\fedprospector\\attachments\\. "
        "Step 10 (cleanup) reclaims space by deleting fully-analyzed files."
    )
    doc.add_page_break()


def ch4_cli_reference(doc):
    """Chapter 4: CLI Command Reference."""
    doc.add_heading("4. CLI Command Reference", level=1)

    doc.add_paragraph(
        "All CLI commands are run from the fed_prospector/ directory:"
    )
    add_code_block(doc, "cd fed_prospector\npython main.py <group> <command> [options]")
    doc.add_paragraph("To see help for any group or command:")
    add_code_block(doc,
        "python main.py --help                    # List all groups\n"
        "python main.py load --help               # List commands in the load group\n"
        "python main.py load opportunities --help # Show options for a specific command"
    )

    # -- setup --
    doc.add_heading("4.1 setup - Database Setup & Verification", level=2)
    cmds = [
        ("setup build", "Build/rebuild the database schema (creates tables, views, indexes)"),
        ("setup seed-lookups", "Load reference data (NAICS codes, PSC codes, SBA types, FIPS)"),
        ("setup seed-rules", "Seed data quality rules into etl_data_quality_rule table"),
        ("setup verify", "Verify all prerequisites (Python, MySQL, .env, API keys, etc.)"),
        ("setup schedule-jobs", "Set up Windows Task Scheduler jobs for automated loading"),
        ("setup test-api", "Test connectivity to SAM.gov and other APIs"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- load --
    doc.add_heading("4.2 load - Data Loading from Government APIs", level=2)

    doc.add_heading("load opportunities", level=3)
    doc.add_paragraph(
        "Load contract opportunities from SAM.gov Opportunities API. Supports "
        "incremental loading with automatic resume."
    )
    add_code_block(doc,
        "Options:\n"
        "  --days-back N        Days of history (default: 7)\n"
        "  --set-aside CODE     Set-aside filter (WOSB, 8A, all, none)\n"
        "  --naics CODES        NAICS codes, comma-separated\n"
        "  --posted-from DATE   Start date (MM/dd/yyyy)\n"
        "  --posted-to DATE     End date (MM/dd/yyyy)\n"
        "  --historical         Load 2 years of data\n"
        "  --max-calls N        API call budget (default: 5)\n"
        "  --key 1|2            SAM API key (default: 1)\n"
        "  --force              Ignore resume state, start fresh\n"
        "\n"
        "Examples:\n"
        "  python main.py load opportunities\n"
        "  python main.py load opportunities --days-back 30 --key 2\n"
        "  python main.py load opportunities --set-aside WOSB --naics 541512\n"
        "  python main.py load opportunities --historical --max-calls 20"
    )

    doc.add_heading("load awards", level=3)
    doc.add_paragraph(
        "Load contract awards from SAM.gov Contract Awards API. Supports watermark-based "
        "incremental loading and crash-safe resume."
    )
    add_code_block(doc,
        "Options:\n"
        "  --naics CODES        NAICS codes, comma-separated\n"
        "  --set-aside CODE     Set-aside type (WOSB, 8A, SBA, etc.)\n"
        "  --agency CODE        Contracting department CGAC code\n"
        "  --awardee-uei UEI    Awardee UEI\n"
        "  --piid PIID          Contract PIID\n"
        "  --for-org NAME|ID    Load awards for all UEIs linked to an org\n"
        "  --years-back N       Years of history\n"
        "  --days-back N        Days of history\n"
        "  --fiscal-year YYYY   Specific fiscal year\n"
        "  --date-from DATE     Start date (YYYY-MM-DD)\n"
        "  --date-to DATE       End date (YYYY-MM-DD)\n"
        "  --max-calls N        API call budget (default: 10)\n"
        "  --key 1|2            SAM API key (default: 2)\n"
        "  --force              Skip resume, start fresh\n"
        "  --dry-run            Preview without API calls\n"
        "\n"
        "Examples:\n"
        "  python main.py load awards\n"
        "  python main.py load awards --naics 541512 --years-back 5\n"
        "  python main.py load awards --for-org \"Acme Corp\" --dry-run"
    )

    doc.add_heading("load entities", level=3)
    doc.add_paragraph("Load SAM.gov entity (vendor) data.")
    add_code_block(doc,
        "Options:\n"
        "  --type monthly|api   Load method (default: api)\n"
        "  --date YYYY-MM-DD    Date for API query\n"
        "  --year/--month       Year/month for monthly extract\n"
        "  --file PATH          Load from local file\n"
        "  --key 1|2            SAM API key\n"
        "  --uei UEI            Filter by UEI\n"
        "  --name NAME          Filter by entity name\n"
        "  --naics CODES        NAICS codes filter\n"
        "  --set-aside CODE     Business type (8W=WOSB, 8E=EDWOSB, A4=8(a))\n"
        "  --max-calls N        API call cap (default: 100)\n"
        "  --force              Force reload"
    )

    other_load_cmds = [
        ("load hierarchy", "Load federal organization hierarchy from SAM.gov Federal Hierarchy API. "
         "Options: --status, --max-calls, --full-refresh, --key, --force"),
        ("load offices", "Load office-level data from the Federal Hierarchy API"),
        ("load exclusions", "Load exclusion records from SAM.gov. Options: --exclusion-type, --agency, --max-calls, --key"),
        ("load labor-rates", "Load labor rate data from GSA CALC+ API"),
        ("load subawards", "Load subaward data from SAM.gov. Options: --naics, --agency, --piid, --years-back, --max-calls, --min-amount, --key"),
        ("load usaspending", "Load transaction history for a USASpending award. Options: --award-id, --solicitation, --piid"),
        ("load usaspending-bulk", "Bulk download award data from USASpending.gov. Options: --days-back"),
        ("load replay-awards", "Replay staged award records through the loader. Options: --load-id, --status"),
        ("load daily", "Run the daily load job sequence. Options: --key, --full, --skip, --dry-run, --force"),
        ("load weekly", "Run the weekly load job sequence. Options: --key, --skip, --dry-run, --force"),
        ("load monthly", "Run the monthly load job sequence. Options: --key, --skip, --dry-run, --force"),
    ]
    for cmd, desc in other_load_cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- search --
    doc.add_heading("4.3 search - Local Database Queries", level=2)
    cmds = [
        ("search opportunities", "Search loaded opportunities. Options: --set-aside, --naics, --open-only, --days, --keyword, --department, --limit"),
        ("search entities", "Search loaded entity/vendor data"),
        ("search awards", "Search loaded contract awards. Options: --naics, --set-aside, --agency, --vendor, --piid, --from-date, --to-date, --limit"),
        ("search agencies", "Search federal agencies in the hierarchy"),
        ("search subawards", "Search loaded subaward data"),
        ("search exclusions", "Check exclusion status for an entity"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- prospect --
    doc.add_heading("4.4 prospect - Bid Pipeline Management", level=2)
    cmds = [
        ("prospect dashboard", "Show bid pipeline dashboard with counts by stage"),
        ("prospect create", "Create a new prospect from an opportunity"),
        ("prospect list", "List prospects with filters"),
        ("prospect show", "Show detailed prospect information"),
        ("prospect update", "Update prospect status, stage, or notes"),
        ("prospect assign", "Reassign a prospect to a different user"),
        ("prospect add-note", "Add a note to a prospect"),
        ("prospect add-partner", "Add a teaming partner to a prospect"),
        ("prospect add-user", "Add a user to the prospecting system"),
        ("prospect list-users", "List prospecting system users"),
        ("prospect save-search", "Save a search with filters for reuse"),
        ("prospect run-search", "Execute a saved search"),
        ("prospect list-searches", "List all saved searches"),
        ("prospect auto-generate", "Auto-generate prospects from matching opportunities"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- analyze --
    doc.add_heading("4.5 analyze - Analysis Tools", level=2)
    cmds = [
        ("analyze burn-rate", "Calculate contract burn rate from USASpending transaction data"),
        ("analyze teaming", "Discover potential teaming partners from subaward data"),
        ("analyze scan-exclusions", "Scan prospects for excluded entities"),
        ("extract attachment-ai", "Analyze attachment content using Claude AI (Phase 110C)"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- admin --
    doc.add_heading("4.6 admin - System Administration", level=2)
    doc.add_paragraph("See Chapters 5 and 6 for detailed usage.")
    cmds = [
        ("admin create-sysadmin", "Bootstrap a system admin user. Options: --username, --email, --display-name, --org-name"),
        ("admin create-org", "Create a new organization. Options: --name, --slug, --max-users, --tier"),
        ("admin list-orgs", "List all organizations"),
        ("admin invite-user", "Create an invitation for a user. Options: --email, --org-id, --role"),
        ("admin list-org-members", "List members of an organization. Options: --org-id"),
        ("admin disable-user", "Disable a user account and revoke sessions. Options: --user-id"),
        ("admin enable-user", "Re-enable a disabled user account. Options: --user-id"),
        ("admin reset-password", "Reset a user's password. Options: --user-id"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- update --
    doc.add_heading("4.7 update - Data Enrichment", level=2)
    cmds = [
        ("update link-metadata", "Enrich opportunity resource links with filenames and content types. Options: --missing-only, --batch-size"),
        ("update build-relationships", "Detect opportunity lifecycle relationships (RFI->RFP, PRESOL->SOL, SOL->AWARD)"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    doc.add_heading("update fetch-descriptions", level=3)
    doc.add_paragraph(
        "Fetch and cache opportunity description text from SAM.gov. Each description "
        "requires one API call. Supports a two-pass prioritized fetch: first fetches "
        "descriptions for high-value opportunities matching target NAICS codes and "
        "set-aside types (WOSB/8(a)/SBA), then uses remaining budget for the general "
        "backlog. Integrated into daily_load.bat (100 descriptions/day)."
    )
    add_code_block(doc,
        "Options:\n"
        "  --missing-only/--all     Only fetch for opportunities without description_text (default: missing-only)\n"
        "  --batch-size N           Commit batch size (default: 100)\n"
        "  --days-back N            Only fetch for opportunities posted in the last N days\n"
        "  --notice-id ID           Fetch description for a single notice ID\n"
        "  --key 1|2                SAM API key (default: 2). Key 2 has 1,000/day quota.\n"
        "  --naics CODES            Priority NAICS codes (comma-separated). Fetches these first.\n"
        "  --set-aside CODES        Priority set-aside codes (e.g. WOSB,8A,SBA). Fetches these first.\n"
        "  --limit N                Max total descriptions to fetch (default: unlimited)\n"
        "\n"
        "Examples:\n"
        "  python main.py update fetch-descriptions --days-back 7\n"
        "  python main.py update fetch-descriptions --notice-id abc123\n"
        "  python main.py update fetch-descriptions --key 2 --limit 100 --naics 541511,541512 --set-aside WOSB,8A\n"
        "  python main.py update fetch-descriptions --all --batch-size 900"
    )
    doc.add_paragraph(
        "On-demand fetch via the C# API is also available. If the description is "
        "already cached, the cached text is returned without making an API call:"
    )
    add_code_block(doc,
        "POST /api/v1/opportunities/{noticeId}/fetch-description\n"
        "Authorization: (authenticated user cookie)"
    )

    # -- download --
    doc.add_heading("4.8 download - File Downloads", level=2)
    p = doc.add_paragraph()
    run = p.add_run("download attachments")
    run.bold = True
    p.add_run(
        " - Download opportunity attachments from SAM.gov. "
        "Options: --notice-id, --batch-size, --max-file-size, --missing-only, "
        "--check-changed, --delay, --active-only, --workers"
    )

    # -- extract --
    doc.add_heading("4.9 extract - Text & Intelligence Extraction", level=2)
    cmds = [
        ("extract attachment-text", "Extract text from downloaded documents. Options: --notice-id, --batch-size, --force, --workers"),
        ("extract attachment-intel", "Extract structured intelligence from text. Options: --notice-id, --batch-size, --method (keyword|regex|hybrid), --force"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    # -- cleanup --
    doc.add_heading("4.10 cleanup - File Cleanup", level=2)
    p = doc.add_paragraph()
    run = p.add_run("cleanup attachment-files")
    run.bold = True
    p.add_run(
        " - Remove fully-analyzed attachment files from disk. Only deletes files "
        "that completed ALL 4 pipeline stages. Options: --notice-id, --batch-size, --dry-run"
    )

    # -- job (on-demand) --
    doc.add_heading("4.11 job - On-Demand Loading", level=2)
    p = doc.add_paragraph()
    run = p.add_run("job process-requests")
    run.bold = True
    p.add_run(" - Process queued on-demand data loading requests")

    # -- health --
    doc.add_heading("4.12 health - Monitoring & Maintenance", level=2)
    doc.add_paragraph("See Chapter 8 for detailed usage.")
    cmds = [
        ("health status", "Show database connection status and table counts"),
        ("health check", "Comprehensive health check with data freshness, API usage, alerts. Options: --json"),
        ("health load-history", "Show ETL load history. Options: --source, --days, --status, --limit"),
        ("health catchup", "Auto-refresh stale data sources. Options: --dry-run, --include-all"),
        ("health run-job JOB", "Manually trigger a scheduled job. Options: --list"),
        ("health maintain-app-data", "Clean up old history, staging, and error records. Options: --dry-run"),
        ("health maintain-db", "MySQL maintenance (analyze, optimize, purge logs). Options: --optimize, --skip-analyze, --purge-binlog-days, --tables, --sizes, --dry-run"),
        ("health run-all-searches", "Execute all active saved searches"),
        ("health check-schema", "Verify database schema matches expected DDL"),
        ("health pipeline-status", "Show attachment pipeline stage counts and disk usage"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.bold = True
        p.add_run(f" - {desc}")

    doc.add_page_break()


def ch5_org_management(doc):
    """Chapter 5: Organization Management."""
    doc.add_heading("5. Organization Management", level=1)

    doc.add_heading("5.1 Creating an Organization (CLI)", level=2)
    doc.add_paragraph(
        "Organizations are the top-level tenant container. Each organization has its "
        "own users, prospects, saved searches, and entity links. Create an organization "
        "using the CLI:"
    )
    add_code_block(doc,
        "python main.py admin create-org --name \"Acme Corp\" --slug acme-corp\n"
        "python main.py admin create-org --name \"Acme Corp\" --max-users 25 --tier professional"
    )

    doc.add_heading("5.2 Creating an Organization (API)", level=2)
    doc.add_paragraph(
        "System admins can also create organizations via the REST API:"
    )
    add_code_block(doc,
        "POST /api/v1/admin/organizations\n"
        "Authorization: (system admin cookie)\n"
        "Content-Type: application/json\n"
        "\n"
        '{"name": "Acme Corp", "slug": "acme-corp"}'
    )

    doc.add_heading("5.3 Organization Properties", level=2)
    table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    table.rows[0].cells[0].text = "Property"
    table.rows[0].cells[1].text = "Description"
    props = [
        ("name", "Display name of the organization"),
        ("slug", "URL-safe identifier (auto-generated from name if omitted)"),
        ("max_users", "Maximum number of users allowed (default: 10)"),
        ("subscription_tier", "Subscription level: trial, professional, enterprise"),
        ("is_active", "Y/N - whether the organization is active"),
    ]
    for i, (prop, desc) in enumerate(props, 1):
        table.rows[i].cells[0].text = prop
        table.rows[i].cells[1].text = desc

    doc.add_heading("5.4 Listing Organizations", level=2)
    add_code_block(doc, "python main.py admin list-orgs")
    doc.add_paragraph("Or via API:")
    add_code_block(doc, "GET /api/v1/admin/organizations   (System Admin only)")

    doc.add_heading("5.5 NAICS Codes and Set-Aside Types", level=2)
    doc.add_paragraph(
        "Each organization can be configured with target NAICS codes and set-aside "
        "types that drive automated data loading and prospect generation. These are "
        "stored in the organization_entity table linking the org to SAM.gov entities."
    )

    doc.add_heading("5.6 Organization Entity Links", level=2)
    doc.add_paragraph(
        "Organizations link to SAM.gov entities via UEI (Unique Entity Identifier). "
        "Entity links enable features like award loading by org (load awards --for-org), "
        "exclusion scanning, and incumbent detection."
    )
    doc.add_page_break()


def ch6_user_management(doc):
    """Chapter 6: User Management."""
    doc.add_heading("6. User Management", level=1)

    doc.add_heading("6.1 Bootstrapping the First Admin", level=2)
    doc.add_paragraph(
        "The first user must be created via CLI. This bootstraps a system admin who "
        "can then create organizations and invite users through the UI or API."
    )
    add_code_block(doc,
        'python main.py admin create-sysadmin \\\n'
        '    --username admin \\\n'
        '    --email admin@mycompany.com \\\n'
        '    --display-name "Admin User"'
    )
    doc.add_paragraph(
        "You will be prompted for a password (with confirmation). The user is created "
        "with is_system_admin=1, is_org_admin='Y', and org_role='owner'."
    )

    doc.add_heading("6.2 Roles", level=2)
    table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Capabilities"
    roles = [
        ("sysadmin (is_system_admin=1)", "Full system access. Can create/manage organizations, view ETL status, manage all users across orgs. Access to /api/v1/admin endpoints with SystemAdmin policy."),
        ("org_admin (is_org_admin='Y')", "Organization administrator. Can manage users within their org, update roles, reset passwords. Access to /api/v1/admin endpoints within their org."),
        ("member (default)", "Standard user. Can view opportunities, manage prospects, run saved searches within their organization."),
    ]
    for i, (role, caps) in enumerate(roles, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = caps

    doc.add_heading("6.3 Inviting Users", level=2)
    add_code_block(doc,
        "python main.py admin invite-user --email user@acme.com --org-id 2\n"
        "python main.py admin invite-user --email user@acme.com --org-id 2 --role admin"
    )
    doc.add_paragraph(
        "This generates a SHA-256 invite code valid for 7 days. Share the code with "
        "the user to complete registration. Available roles: member, admin, owner."
    )

    doc.add_heading("6.4 Listing Organization Members", level=2)
    add_code_block(doc, "python main.py admin list-org-members --org-id 2")
    doc.add_paragraph("Or via API:")
    add_code_block(doc, "GET /api/v1/admin/users?page=1&pageSize=25")

    doc.add_heading("6.5 Disabling a User", level=2)
    add_code_block(doc, "python main.py admin disable-user --user-id 5")
    doc.add_paragraph(
        "Disabling a user sets is_active='N' and revokes all active sessions. "
        "The user will be immediately logged out and unable to log back in."
    )

    doc.add_heading("6.6 Enabling a User", level=2)
    add_code_block(doc, "python main.py admin enable-user --user-id 5")
    doc.add_paragraph(
        "Re-enables a disabled user account. Also resets failed_login_attempts "
        "and clears any locked_until timestamp."
    )

    doc.add_heading("6.7 Password Management", level=2)
    doc.add_paragraph("Reset a user's password via CLI:")
    add_code_block(doc, "python main.py admin reset-password --user-id 5")
    doc.add_paragraph(
        "Generates a random temporary password, revokes all active sessions, and "
        "sets force_password_change='Y'. The user must change their password on "
        "next login. The temporary password is displayed on screen."
    )
    doc.add_paragraph("Or via API (org admin or sys admin):")
    add_code_block(doc, "POST /api/v1/admin/users/{id}/reset-password")

    doc.add_heading("6.8 Updating User Roles (API)", level=2)
    add_code_block(doc,
        "PATCH /api/v1/admin/users/{id}\n"
        "Content-Type: application/json\n"
        "\n"
        '{"orgRole": "admin", "isOrgAdmin": true}'
    )
    doc.add_page_break()


def ch7_data_management(doc):
    """Chapter 7: Data Management."""
    doc.add_heading("7. Data Management", level=1)

    doc.add_heading("7.1 Data Sources Overview", level=2)
    table = doc.add_table(rows=6, cols=4, style="Light Grid Accent 1")
    headers = ["Source", "Data Type", "API Rate Limit", "Python Command"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    sources = [
        ("SAM.gov Opportunities", "Contract solicitations", "10/day (key 1), 1000/day (key 2)", "load opportunities"),
        ("SAM.gov Awards (FPDS)", "Contract awards", "Same as above", "load awards"),
        ("SAM.gov Entities", "Vendor registrations", "Same as above", "load entities"),
        ("SAM.gov Exclusions", "Debarred/excluded entities", "Same as above", "load exclusions"),
        ("USASpending.gov", "Award obligations & transactions", "No daily limit", "load usaspending-bulk"),
    ]
    for i, (src, dtype, limit, cmd) in enumerate(sources, 1):
        table.rows[i].cells[0].text = src
        table.rows[i].cells[1].text = dtype
        table.rows[i].cells[2].text = limit
        table.rows[i].cells[3].text = cmd

    doc.add_paragraph()

    table2 = doc.add_table(rows=4, cols=4, style="Light Grid Accent 1")
    for i, h in enumerate(headers):
        table2.rows[0].cells[i].text = h
    sources2 = [
        ("SAM.gov Federal Hierarchy", "Agency org structure", "Same as SAM.gov", "load hierarchy"),
        ("SAM.gov Subawards", "Subcontract data", "Same as SAM.gov", "load subawards"),
        ("GSA CALC+", "Labor rates", "No daily limit", "load labor-rates"),
    ]
    for i, (src, dtype, limit, cmd) in enumerate(sources2, 1):
        table2.rows[i].cells[0].text = src
        table2.rows[i].cells[1].text = dtype
        table2.rows[i].cells[2].text = limit
        table2.rows[i].cells[3].text = cmd

    doc.add_heading("7.2 Federal Hierarchy Management", level=2)
    doc.add_paragraph(
        "The federal hierarchy (departments, sub-tiers, and offices) is loaded from "
        "the SAM.gov Federal Hierarchy API. This data powers the hierarchy browser "
        "in the UI and provides agency context for opportunities and awards."
    )

    doc.add_heading("CLI Hierarchy Commands", level=3)
    add_code_block(doc,
        "python main.py load hierarchy                    # Load departments + sub-tiers (levels 1-2)\n"
        "python main.py load hierarchy --full-refresh     # Truncate and reload all hierarchy data\n"
        "python main.py load offices                      # Load offices (level 3) under existing sub-tiers\n"
        "python main.py load offices --max-calls 300      # Limit API calls for office loading\n"
        "python main.py load offices --days-back 30       # Incremental: only recently updated offices"
    )

    doc.add_heading("Admin UI Refresh Panel (Phase 113)", level=3)
    doc.add_paragraph(
        "System admins can refresh hierarchy data from the Federal Hierarchy Browser "
        "page in the UI. The Hierarchy Data Refresh panel provides:"
    )
    items = [
        "API Key Selection: Choose Key 1 (10/day) or Key 2 (1,000/day) for the refresh",
        "Refresh Hierarchy: Reloads departments and sub-tiers (levels 1-2)",
        "Refresh Offices: Reloads all offices (level 3). May use ~738 API calls.",
        "Full Refresh: Truncates and reloads ALL hierarchy levels (destructive operation)",
        "Status Polling: Shows progress while a refresh job is running",
        "Last Refresh: Displays timestamp, total record count, and per-level breakdown",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The refresh is executed by the C# API, which invokes the Python CLI in "
        "the background. The UI polls the refresh status endpoint until the job completes."
    )

    doc.add_heading("Hierarchy Refresh API Endpoints", level=3)
    table = doc.add_table(rows=4, cols=4, style="Light Grid Accent 1")
    headers = ["Method", "Endpoint", "Access", "Description"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    endpoints = [
        ("POST", "/api/v1/federal-hierarchy/refresh", "System Admin", "Trigger a full hierarchy refresh (body: {level, apiKey})"),
        ("GET", "/api/v1/federal-hierarchy/refresh/status", "Authenticated", "Check status of last refresh job"),
        ("POST", "/api/v1/federal-hierarchy/{fhOrgId}/refresh", "Authenticated", "Queue a single-org refresh via the data load poller"),
    ]
    for i, (method, ep, access, desc) in enumerate(endpoints, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[1].text = ep
        table.rows[i].cells[2].text = access
        table.rows[i].cells[3].text = desc

    doc.add_paragraph()

    doc.add_heading("7.3 API Key Management", level=2)
    doc.add_paragraph("SAM.gov supports two API keys, configured in fed_prospector/.env:")
    add_code_block(doc,
        "SAM_API_KEY=<key-1>           # Free tier: 10 calls/day\n"
        "SAM_API_KEY_2=<key-2>         # Registered tier: 1,000 calls/day\n"
        "SAM_API_KEY_CREATED=YYYY-MM-DD  # For expiry tracking"
    )
    doc.add_paragraph(
        "Key 1 is used by default for most commands. Key 2 is recommended for daily "
        "loads and any operation requiring more than 10 API calls. Select the key with "
        "the --key option on any load command."
    )
    doc.add_paragraph(
        "SAM.gov API keys expire annually. Set SAM_API_KEY_CREATED in .env to enable "
        "expiry warnings in health checks."
    )

    doc.add_heading("7.4 Attachment Pipeline", level=2)
    doc.add_paragraph(
        "The attachment intelligence pipeline processes opportunity documents through "
        "a 7-stage state machine:"
    )
    stages = [
        "1. Load - Opportunity data loaded with resource links",
        "2. Enrich URLs - HEAD requests to get filenames and content types",
        "3. Download - Files downloaded to E:\\fedprospector\\attachments\\",
        "4. Extract Text - PDF/DOCX/DOC/XLSX parsed to raw text",
        "5. Keyword Intel - Keyword/heuristic extraction of requirements",
        "6. AI Analysis - Claude AI deep analysis (Phase 110C)",
        "7. Cleanup - Original files deleted (data stays in DB)",
    ]
    for stage in stages:
        doc.add_paragraph(stage, style="List Number")

    doc.add_heading("7.5 Pipeline Status", level=2)
    add_code_block(doc, "python main.py health pipeline-status")
    doc.add_paragraph(
        "Shows counts for each stage: downloads (pending/downloaded/failed/skipped), "
        "text extraction status, keyword intel completion, AI analysis status, "
        "cleanup eligibility, and total disk usage."
    )

    doc.add_heading("7.6 Change Detection", level=2)
    doc.add_paragraph(
        "All ETL loaders use SHA-256 record hashing to detect changes between loads. "
        "When a record is loaded, a hash of its key fields is computed and stored. On "
        "subsequent loads, records with matching hashes are skipped (unchanged), while "
        "records with different hashes are updated. This minimizes database writes and "
        "enables accurate insert/update/unchanged counts in load reports."
    )

    doc.add_heading("7.7 Data Quality Rules", level=2)
    doc.add_paragraph(
        "Data quality rules are stored in the etl_data_quality_rule table. They are "
        "not hardcoded -- add new rules via SQL or the seed-rules setup command. "
        "Rules are evaluated during ETL loading and violations are logged."
    )
    doc.add_page_break()


def ch8_monitoring(doc):
    """Chapter 8: Monitoring & Troubleshooting."""
    doc.add_heading("8. Monitoring & Troubleshooting", level=1)

    doc.add_heading("8.1 Health Check", level=2)
    add_code_block(doc, "python main.py health check")
    doc.add_paragraph("The comprehensive health check reports:")
    items = [
        "Data Freshness: Last load time and staleness status for each data source",
        "Table Statistics: Row counts for all tables",
        "API Usage Today: Calls used vs. daily limit per API",
        "API Key Status: Configuration and expiry tracking",
        "Alerts: Actionable warnings (stale data, expiring keys, errors)",
        "Recent Errors: Failed ETL loads in the last 7 days",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Use --json for machine-readable output:")
    add_code_block(doc, "python main.py health check --json")

    doc.add_heading("8.2 ETL Load History", level=2)
    add_code_block(doc,
        "python main.py health load-history                      # Last 20 loads\n"
        "python main.py health load-history --source SAM_OPPORTUNITY  # Filter by source\n"
        "python main.py health load-history --days 7              # Last 7 days\n"
        "python main.py health load-history --status FAILED       # Failed loads only\n"
        "python main.py health load-history --limit 50            # Show more results"
    )
    doc.add_paragraph(
        "All ETL loads are recorded in the etl_load_log table with timing, record "
        "counts (read/inserted/updated/errored), and error messages. This is the "
        "primary audit trail for data loading."
    )

    doc.add_heading("8.3 Database Status", level=2)
    add_code_block(doc, "python main.py health status")
    doc.add_paragraph(
        "Shows database connection status and row counts for all tables."
    )

    doc.add_heading("8.4 Catching Up Stale Data", level=2)
    add_code_block(doc,
        "python main.py health catchup --dry-run    # Preview what would run\n"
        "python main.py health catchup              # Auto-refresh stale sources\n"
        "python main.py health catchup --include-all  # Include unsafe jobs too"
    )
    doc.add_paragraph(
        "The catchup command checks data freshness and automatically runs safe "
        "refresh jobs for stale data sources. Unsafe jobs (entity_daily, usaspending) "
        "are skipped unless --include-all is specified."
    )

    doc.add_heading("8.5 Running Scheduled Jobs Manually", level=2)
    add_code_block(doc,
        "python main.py health run-job --list            # List all available jobs\n"
        "python main.py health run-job opportunities     # Run a specific job\n"
        "python main.py health run-job exclusions"
    )

    doc.add_heading("8.6 Database Maintenance", level=2)
    add_code_block(doc,
        "# Application data cleanup (old history, staging, errors)\n"
        "python main.py health maintain-app-data --dry-run\n"
        "python main.py health maintain-app-data\n"
        "\n"
        "# MySQL engine maintenance\n"
        "python main.py health maintain-db                    # ANALYZE all tables\n"
        "python main.py health maintain-db --optimize         # ANALYZE + OPTIMIZE\n"
        "python main.py health maintain-db --sizes            # Show table sizes\n"
        "python main.py health maintain-db --purge-binlog-days 7  # Purge old binlogs"
    )

    doc.add_heading("8.7 Schema Verification", level=2)
    add_code_block(doc, "python main.py health check-schema")
    doc.add_paragraph(
        "Compares the live database schema against the expected DDL files. Reports "
        "missing tables, missing columns, and type mismatches."
    )

    doc.add_heading("8.8 Admin API Endpoints", level=2)
    doc.add_paragraph(
        "The C# API provides admin endpoints at /api/v1/admin (requires AdminAccess "
        "or SystemAdmin authorization):"
    )
    table = doc.add_table(rows=13, cols=3, style="Light Grid Accent 1")
    table.rows[0].cells[0].text = "Method"
    table.rows[0].cells[1].text = "Endpoint"
    table.rows[0].cells[2].text = "Access"
    endpoints = [
        ("GET", "/api/v1/admin/etl-status", "System Admin"),
        ("GET", "/api/v1/admin/load-history", "System Admin"),
        ("GET", "/api/v1/admin/health-snapshots", "System Admin"),
        ("GET", "/api/v1/admin/api-keys", "System Admin"),
        ("GET", "/api/v1/admin/jobs", "System Admin"),
        ("GET", "/api/v1/admin/users", "Org Admin"),
        ("PATCH", "/api/v1/admin/users/{id}", "Org Admin"),
        ("POST", "/api/v1/admin/users/{id}/reset-password", "Org Admin"),
        ("POST", "/api/v1/federal-hierarchy/refresh", "System Admin"),
        ("GET", "/api/v1/federal-hierarchy/refresh/status", "Authenticated"),
        ("POST", "/api/v1/opportunities/{noticeId}/fetch-description", "Authenticated"),
        ("POST", "/api/v1/federal-hierarchy/{fhOrgId}/refresh", "Authenticated"),
    ]
    for i, (method, ep, access) in enumerate(endpoints, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[1].text = ep
        table.rows[i].cells[2].text = access

    doc.add_heading("8.9 Common Errors and Fixes", level=2)

    errors = [
        (
            "429 Too Many Requests / Rate Limit Exceeded",
            "The daily SAM.gov API quota is exhausted.",
            "Wait until tomorrow, or switch to key 2 (--key 2) which has 1000/day. "
            "Progress is saved; re-run the same command to resume."
        ),
        (
            "MySQL connection refused",
            "MySQL is not running.",
            "Run: fed_prospector.bat start db"
        ),
        (
            "API /health returns 503",
            "API is running but database connection failed.",
            "Check MySQL is running and DB_PASSWORD in .env is correct."
        ),
        (
            "JWT_SECRET_KEY not set",
            "API authentication will fail without a JWT secret.",
            "Add JWT_SECRET_KEY=<at-least-32-chars> to fed_prospector/.env"
        ),
        (
            "LOAD DATA INFILE permission denied",
            "MySQL user lacks FILE privilege.",
            "Run: GRANT FILE ON *.* TO 'fed_app'@'localhost'; FLUSH PRIVILEGES;"
        ),
        (
            "Port already in use",
            "Another process is using the port.",
            "Check with: netstat -ano | findstr :<port>  -- then kill the process or change the port in .env"
        ),
        (
            "Attachment download failures",
            "SAM.gov may throttle or block rapid downloads.",
            "Use --delay to add wait time between requests. Use --workers to reduce parallelism."
        ),
        (
            "Stale data warnings in health check",
            "A data source has not been loaded recently.",
            "Run: python main.py health catchup  -- or manually load the specific source."
        ),
    ]

    for title, cause, fix in errors:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        run = p.add_run("Cause: ")
        run.bold = True
        p.add_run(cause)
        p = doc.add_paragraph()
        run = p.add_run("Fix: ")
        run.bold = True
        p.add_run(fix)

    doc.add_heading("8.10 Log Files", level=2)
    doc.add_paragraph(
        "Python ETL commands log to stdout using the Python logging module with "
        "structured output. Redirect to a file for persistent logs:"
    )
    add_code_block(doc, "python main.py load opportunities >> load.log 2>&1")
    doc.add_paragraph(
        "The C# API logs to the console by default. In production, configure "
        "Serilog or another logging provider in appsettings.json."
    )

    doc.add_heading("8.11 Backup Recommendations", level=2)
    doc.add_paragraph(
        "The MySQL database (fed_contracts) contains all loaded data and application "
        "state. Regular backups are recommended:"
    )
    add_code_block(doc,
        "mysqldump -u root -p fed_contracts > backup_$(date +%%Y%%m%%d).sql\n"
        "\n"
        "# Or for faster backup with --single-transaction:\n"
        "mysqldump -u root -p --single-transaction fed_contracts > backup.sql"
    )
    doc.add_paragraph(
        "Attachment files in E:\\fedprospector\\attachments\\ are transient (deleted "
        "after full pipeline processing) and do not need regular backup."
    )


def main():
    doc = Document()
    setup_styles(doc)
    add_title_page(doc)
    add_toc(doc)
    ch1_system_overview(doc)
    ch2_service_management(doc)
    ch3_daily_loading(doc)
    ch4_cli_reference(doc)
    ch5_org_management(doc)
    ch6_user_management(doc)
    ch7_data_management(doc)
    ch8_monitoring(doc)

    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
