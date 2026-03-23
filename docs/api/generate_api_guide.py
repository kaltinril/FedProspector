"""Generate FedProspect API Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Global style tweaks --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def add_code_block(text: str):
    """Add a monospaced, shaded paragraph to simulate a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Shading via XML
    shading = p.paragraph_format.element.makeelement(
        qn("w:shd"),
        {qn("w:fill"): "F2F2F2", qn("w:val"): "clear"},
    )
    p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_table(headers: list[str], rows: list[list[str]]):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer


def endpoint(method: str, path: str, desc: str, *,
             params: list[list[str]] | None = None,
             request_body: str | None = None,
             response_body: str | None = None,
             notes: str | None = None,
             auth: str = "JWT (cookie or Bearer header)"):
    """Document one endpoint."""
    doc.add_heading(f"{method} {path}", level=3)
    doc.add_paragraph(desc)
    add_table(["Attribute", "Value"], [
        ["Auth", auth],
    ])
    if params:
        add_table(["Parameter", "Type", "Required", "Description"], params)
    if request_body:
        doc.add_paragraph("Request body:", style="List Bullet")
        add_code_block(request_body)
    if response_body:
        doc.add_paragraph("Response:", style="List Bullet")
        add_code_block(response_body)
    if notes:
        doc.add_paragraph(f"Note: {notes}")


# =========================================================================
# TITLE PAGE
# =========================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("FedProspect API Engineer's Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph("Federal Contract Prospecting System")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14)
p = doc.add_paragraph("REST API Reference")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(12)
doc.add_paragraph()
p = doc.add_paragraph("Base URL: http://localhost:5062")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# =========================================================================
# CHAPTER 1: Architecture Overview
# =========================================================================
doc.add_heading("Chapter 1: Architecture Overview", level=1)

doc.add_paragraph(
    "FedProspect is a federal contract prospecting system built to help WOSB and 8(a) "
    "businesses find, qualify, and pursue government contracts. The API layer is the backend "
    "REST interface consumed by the React single-page application."
)

doc.add_heading("Technology Stack", level=2)
add_table(["Component", "Technology"], [
    ["Runtime", "ASP.NET Core 9 (C#)"],
    ["ORM", "Entity Framework Core 9 (MySQL provider)"],
    ["Database", "MySQL 8.4 LTS (InnoDB, utf8mb4)"],
    ["Authentication", "JWT Bearer tokens with httpOnly cookie transport"],
    ["Validation", "FluentValidation"],
    ["Logging", "Serilog (structured)"],
    ["API Docs", "Swagger / OpenAPI (development only)"],
    ["Rate Limiting", "ASP.NET Core RateLimiter middleware"],
])

doc.add_heading("Middleware Pipeline", level=2)
doc.add_paragraph(
    "The middleware pipeline is executed in this order for every request:"
)
add_table(["Order", "Middleware", "Purpose"], [
    ["1", "ExceptionHandlerMiddleware", "Global exception catch, returns ApiErrorResponse JSON"],
    ["2", "SecurityHeadersMiddleware", "Adds X-Content-Type-Options, X-Frame-Options, etc."],
    ["3", "Swagger (dev only)", "Serves OpenAPI spec at /swagger"],
    ["4", "HTTPS Redirect (prod only)", "Redirects HTTP to HTTPS"],
    ["5", "CORS", "Cross-origin policy enforcement"],
    ["6", "Authentication", "JWT validation (cookie or header)"],
    ["7", "Authorization", "Role/policy checks"],
    ["8", "CsrfMiddleware", "Double-submit CSRF token validation for mutating requests"],
    ["9", "ForcePasswordChangeMiddleware", "Blocks non-password-change requests if user must change password"],
    ["10", "RateLimiter", "Per-user/IP rate limiting"],
    ["11", "Controllers", "Route dispatch to controller actions"],
])

doc.add_heading("Data Architecture", level=2)
doc.add_paragraph(
    "The system has two distinct data domains:"
)
doc.add_paragraph(
    "Public ETL Data: Opportunities, awards, entities, NAICS codes, exclusions, subawards, "
    "and federal hierarchy data loaded from SAM.gov and USASpending.gov by the Python ETL pipeline. "
    "This data is shared across all organizations and is read-only from the API's perspective.",
    style="List Bullet",
)
doc.add_paragraph(
    "Organization Capture Data: Prospects, proposals, saved searches, notifications, team assignments, "
    "Go/No-Go scores, and company profiles. This data is strictly isolated per organization.",
    style="List Bullet",
)

# =========================================================================
# CHAPTER 2: Authentication
# =========================================================================
doc.add_heading("Chapter 2: Authentication", level=1)

doc.add_heading("Overview", level=2)
doc.add_paragraph(
    "FedProspect uses JWT Bearer tokens with a dual transport mechanism. Tokens can be sent via "
    "the standard Authorization header (for programmatic clients) or via httpOnly cookies (for "
    "browser-based SPA clients). The cookie approach prevents XSS token theft."
)

doc.add_heading("Cookie-Based Flow (Browser)", level=2)
doc.add_paragraph("1. Client POSTs credentials to /api/v1/auth/login.")
doc.add_paragraph("2. Server validates credentials and returns three cookies:")
add_table(["Cookie", "HttpOnly", "Path", "Purpose"], [
    ["access_token", "Yes", "/api", "JWT access token (short-lived)"],
    ["refresh_token", "Yes", "/api/v1/auth", "Refresh token (7-day expiry)"],
    ["XSRF-TOKEN", "No", "/", "CSRF double-submit token (readable by JS)"],
])
doc.add_paragraph(
    "3. For all subsequent requests, the browser automatically sends the access_token cookie. "
    "For mutating requests (POST, PUT, PATCH, DELETE), the client must also send the XSRF-TOKEN "
    "value in the X-XSRF-TOKEN request header."
)
doc.add_paragraph(
    "4. When the access token expires, the client calls POST /api/v1/auth/refresh. The refresh_token "
    "cookie is automatically sent (scoped to /api/v1/auth). New cookies are issued."
)

doc.add_heading("Header-Based Flow (Programmatic)", level=2)
doc.add_paragraph(
    "Clients can also use the Authorization header: Authorization: Bearer <token>. "
    "The token is the same JWT returned in the login response. CSRF validation is skipped "
    "when the Authorization header is present."
)

doc.add_heading("Session Validation", level=2)
doc.add_paragraph(
    "Every authenticated request is validated against the session store. When a user logs out "
    "or an admin revokes sessions, the token hash is invalidated. Even a non-expired JWT will "
    "be rejected if the session is revoked. Sessions are validated via SHA-256 hash of the raw token."
)

doc.add_heading("Auth Endpoints", level=2)

endpoint("POST", "/api/v1/auth/login", "Authenticate and receive JWT tokens.",
         auth="None (public)",
         request_body='{\n  "email": "user@example.com",\n  "password": "s3cret"\n}',
         response_body='{\n  "success": true,\n  "userId": 1,\n  "userName": "Jane Smith",\n  "expiresAt": "2026-03-22T15:00:00Z",\n  "forcePasswordChange": false\n}',
         notes="Rate limited: 10/min per IP, 100/min global. Sets httpOnly cookies on success.")

endpoint("POST", "/api/v1/auth/logout", "Revoke the current session and clear cookies.",
         auth="Optional (clears cookies regardless)",
         response_body='{ "message": "Logged out successfully." }')

endpoint("POST", "/api/v1/auth/register", "Register a new user with an invite code.",
         auth="None (public)",
         request_body='{\n  "email": "newuser@example.com",\n  "password": "Str0ngP@ss!",\n  "displayName": "New User",\n  "inviteCode": "abc123"\n}',
         response_body='{\n  "success": true,\n  "userId": 5,\n  "userName": "New User",\n  "expiresAt": "2026-03-22T15:00:00Z",\n  "forcePasswordChange": false\n}',
         notes="Rate limited: 3 attempts/min per IP.")

endpoint("POST", "/api/v1/auth/refresh", "Refresh access token using refresh_token cookie.",
         auth="None (uses refresh_token cookie)",
         response_body='{\n  "success": true,\n  "userId": 1,\n  "userName": "Jane Smith",\n  "expiresAt": "2026-03-22T16:00:00Z",\n  "forcePasswordChange": false\n}')

endpoint("POST", "/api/v1/auth/change-password", "Change the current user's password. Revokes all sessions.",
         request_body='{\n  "currentPassword": "oldPass",\n  "newPassword": "N3wStr0ng!"\n}',
         response_body='{ "message": "Password changed successfully. Please log in again." }')

endpoint("GET", "/api/v1/auth/me", "Get the current user's profile.",
         response_body='{\n  "userId": 1,\n  "email": "user@example.com",\n  "displayName": "Jane Smith",\n  "organizationId": 1,\n  "organizationName": "Acme Corp",\n  "orgRole": "admin",\n  "isSystemAdmin": false\n}')

endpoint("PATCH", "/api/v1/auth/me", "Update the current user's display name or email.",
         request_body='{\n  "displayName": "Jane Q. Smith",\n  "email": "jane@example.com"\n}')

# =========================================================================
# CHAPTER 3: Multi-Tenancy
# =========================================================================
doc.add_heading("Chapter 3: Multi-Tenancy", level=1)

doc.add_heading("Data Isolation Model", level=2)
doc.add_paragraph(
    "FedProspect uses a shared-database, row-level isolation model for multi-tenancy. "
    "All organizations share the same MySQL database, but capture data (prospects, proposals, "
    "saved searches, team members, notifications) is isolated by organization_id."
)

doc.add_heading("Public vs. Org-Scoped Data", level=2)
add_table(["Data Type", "Scope", "Examples"], [
    ["ETL/Public", "Shared across all orgs", "Opportunities, awards, entities, NAICS, exclusions, subawards"],
    ["Capture/Private", "Isolated per organization", "Prospects, proposals, saved searches, notes, team members, Go/No-Go scores"],
    ["User", "Belongs to one org", "User accounts, sessions, activity logs"],
    ["Intelligence", "Computed per org context", "pWin scores, qualification checks, recommended opportunities"],
])

doc.add_heading("Organization Context Resolution", level=2)
doc.add_paragraph(
    "The API resolves the current user's organization in two ways:"
)
doc.add_paragraph(
    "1. JWT claim (org_id): Set during login and embedded in the token. Used by GetCurrentOrganizationId().",
    style="List Bullet",
)
doc.add_paragraph(
    "2. Database fallback: If the org_id claim is missing (legacy tokens), ResolveOrganizationIdAsync() "
    "looks up the user's organization from the database.",
    style="List Bullet",
)

doc.add_heading("Authorization Policies", level=2)
add_table(["Policy", "Requirement", "Used By"], [
    ["(default)", "Any authenticated user", "Most read endpoints"],
    ["OrgAdmin", 'org_role claim is "owner" or "admin"', "Org settings, invites, profile updates"],
    ["AdminAccess", "org_role=admin OR is_system_admin=true", "Admin controller (users, ETL status)"],
    ["SystemAdmin", "is_system_admin=true claim", "System-wide admin (orgs, ETL, API keys, jobs)"],
])

# =========================================================================
# CHAPTER 4: Opportunities
# =========================================================================
doc.add_heading("Chapter 4: Opportunities", level=1)
doc.add_paragraph(
    "Opportunity endpoints provide access to federal contracting opportunities sourced from SAM.gov. "
    "All endpoints require authentication and are rate-limited at 60 requests/minute."
)
doc.add_paragraph("Base path: /api/v1/opportunities")

endpoint("GET", "/api/v1/opportunities", "Search opportunities with filters and pagination.",
         params=[
             ["keyword", "string", "No", "Full-text search in title and description"],
             ["naics", "string", "No", "NAICS code filter (e.g., 541512)"],
             ["setAside", "string", "No", "Set-aside type (e.g., SBA, WOSB, 8A)"],
             ["solicitation", "string", "No", "Solicitation number filter"],
             ["daysOut", "int", "No", "Only show opportunities closing within N days"],
             ["openOnly", "bool", "No", "Only active/open opportunities (default: true)"],
             ["department", "string", "No", "Awarding department filter"],
             ["state", "string", "No", "Place of performance state"],
             ["page", "int", "No", "Page number (default: 1)"],
             ["pageSize", "int", "No", "Results per page (default: 25, max: 100)"],
             ["sortBy", "string", "No", "Sort field"],
             ["sortDescending", "bool", "No", "Sort direction (default: false)"],
         ],
         response_body='{\n  "items": [\n    {\n      "noticeId": "abc123",\n      "title": "IT Support Services",\n      "solicitationNumber": "W91QV1-24-R-0001",\n      "department": "Department of Defense",\n      "setAsideDescription": "Women-Owned Small Business",\n      "naicsCode": "541512",\n      "responseDeadLine": "2026-04-15",\n      "postedDate": "2026-03-01",\n      "estimatedValue": 2500000.00,\n      "prospectStatus": "QUALIFIED"\n    }\n  ],\n  "page": 1,\n  "pageSize": 25,\n  "totalCount": 142,\n  "totalPages": 6,\n  "hasPreviousPage": false,\n  "hasNextPage": true\n}')

endpoint("GET", "/api/v1/opportunities/targets", "Search target opportunities (pre-filtered for org relevance).",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
         ])

endpoint("GET", "/api/v1/opportunities/{noticeId}", "Get detailed information for a single opportunity.",
         response_body='{\n  "noticeId": "abc123",\n  "title": "IT Support Services",\n  "solicitationNumber": "W91QV1-24-R-0001",\n  "fullDescription": "...",\n  "department": "Department of Defense",\n  "subtier": "U.S. Army",\n  "office": "ACC-APG",\n  "setAsideDescription": "Women-Owned Small Business",\n  "naicsCode": "541512",\n  "classificationCode": "D302",\n  "responseDeadLine": "2026-04-15T14:00:00",\n  "postedDate": "2026-03-01",\n  "type": "Solicitation",\n  "estimatedValue": 2500000.00,\n  "placeOfPerformance": "Aberdeen, MD",\n  "pointOfContact": "John Doe, john.doe@army.mil",\n  "resourceLinks": [...],\n  "prospectId": 42,\n  "prospectStatus": "QUALIFIED"\n}')

endpoint("GET", "/api/v1/opportunities/{noticeId}/pwin",
         "Calculate probability of win for an opportunity against the org's profile.",
         response_body='{\n  "noticeId": "abc123",\n  "overallScore": 72.5,\n  "confidence": "MEDIUM",\n  "factors": [\n    { "name": "NAICS Match", "score": 90, "weight": 0.25 },\n    { "name": "Set-Aside Eligibility", "score": 100, "weight": 0.20 },\n    { "name": "Past Performance", "score": 60, "weight": 0.20 },\n    { "name": "Geographic Proximity", "score": 50, "weight": 0.15 },\n    { "name": "Size Standard", "score": 80, "weight": 0.20 }\n  ]\n}')

endpoint("POST", "/api/v1/opportunities/pwin/batch",
         "Calculate pWin for multiple opportunities in one request.",
         request_body='{\n  "noticeIds": ["abc123", "def456", "ghi789"]\n}',
         response_body='{\n  "results": [\n    { "noticeId": "abc123", "overallScore": 72.5, "confidence": "MEDIUM" },\n    { "noticeId": "def456", "overallScore": 45.0, "confidence": "LOW" }\n  ]\n}',
         notes="Maximum 200 notice IDs per batch request.")

endpoint("GET", "/api/v1/opportunities/export", "Export opportunity search results as CSV file.",
         params=[
             ["(same as search)", "", "", "Same filters as GET /api/v1/opportunities"],
         ],
         notes="Returns Content-Type: text/csv with filename opportunities_export.csv.")

endpoint("GET", "/api/v1/opportunities/recommended",
         "Get opportunities scored and ranked for the org's profile.",
         params=[
             ["limit", "int", "No", "Number of results (default: 10)"],
         ],
         response_body='[\n  {\n    "noticeId": "abc123",\n    "title": "IT Support Services",\n    "matchScore": 85.2,\n    "matchReasons": ["NAICS match", "Set-aside eligible", "Past performance in agency"],\n    "responseDeadLine": "2026-04-15"\n  }\n]')

endpoint("GET", "/api/v1/opportunities/{noticeId}/incumbent",
         "Get incumbent analysis - current contract holder and vulnerability signals.",
         response_body='{\n  "noticeId": "abc123",\n  "incumbentName": "Incumbent Corp",\n  "incumbentUei": "XYZ123",\n  "contractValue": 5000000.00,\n  "periodOfPerformance": "2022-01-01 to 2026-12-31",\n  "vulnerabilitySignals": [\n    "Contract ending within 12 months",\n    "Declining obligation trend"\n  ],\n  "vulnerabilityScore": 7.5\n}')

endpoint("GET", "/api/v1/opportunities/{noticeId}/competitive-landscape",
         "Get competitive landscape scoped to the opportunity's agency and NAICS.",
         response_body='{\n  "noticeId": "abc123",\n  "agencyName": "U.S. Army",\n  "naicsCode": "541512",\n  "topVendors": [\n    { "name": "Vendor A", "uei": "AAA111", "contractCount": 15, "totalValue": 25000000 }\n  ],\n  "competitorCount": 42\n}')

endpoint("GET", "/api/v1/opportunities/{noticeId}/set-aside-shift",
         "Analyze set-aside changes compared to predecessor contract.",
         response_body='{\n  "noticeId": "abc123",\n  "currentSetAside": "WOSB",\n  "predecessorSetAside": "Full and Open",\n  "shiftType": "NARROWED",\n  "analysis": "Set-aside narrowed from full competition to WOSB"\n}')

endpoint("GET", "/api/v1/opportunities/{noticeId}/qualification",
         "Run qualification checks against the org's profile.",
         response_body='{\n  "noticeId": "abc123",\n  "overallStatus": "QUALIFIED",\n  "checks": [\n    { "name": "NAICS Code", "status": "PASS", "detail": "541512 is in org NAICS list" },\n    { "name": "Set-Aside", "status": "PASS", "detail": "WOSB certification active" },\n    { "name": "Size Standard", "status": "WARN", "detail": "Within 80% of size threshold" },\n    { "name": "Clearance", "status": "UNKNOWN", "detail": "Cannot determine from solicitation" }\n  ]\n}')

endpoint("GET", "/api/v1/opportunities/{noticeId}/document-intelligence",
         "Get intelligence extracted from opportunity attachments (SOW, RFP, etc.).",
         response_body='{\n  "noticeId": "abc123",\n  "hasAttachments": true,\n  "analysisStatus": "COMPLETE",\n  "keywords": ["cybersecurity", "cloud migration", "FedRAMP"],\n  "requirements": [...],\n  "evaluationCriteria": [...]\n}')

endpoint("POST", "/api/v1/opportunities/{noticeId}/analyze",
         "Request AI analysis of opportunity attachments.",
         params=[
             ["tier", "string", "No", "AI model tier: 'haiku' (default) or 'sonnet'"],
         ],
         response_body='{\n  "requestId": 42,\n  "status": "PENDING",\n  "message": "Analysis request queued"\n}',
         notes="Inserts a data_load_request for the Python pipeline to process. Rate limited at 30/min.")

# =========================================================================
# CHAPTER 5: Awards
# =========================================================================
doc.add_heading("Chapter 5: Awards", level=1)
doc.add_paragraph(
    "Award endpoints provide access to federal contract award data from USASpending.gov. "
    "Includes burn rate analysis, market share, and expiring contract tracking."
)
doc.add_paragraph("Base path: /api/v1/awards")

endpoint("GET", "/api/v1/awards", "Search contract awards with filters and pagination.",
         params=[
             ["page", "int", "No", "Page number (default: 1)"],
             ["pageSize", "int", "No", "Results per page (default: 25, max: 100)"],
             ["sortBy", "string", "No", "Sort field"],
             ["sortDescending", "bool", "No", "Sort direction"],
         ])

endpoint("GET", "/api/v1/awards/{contractId}", "Get detailed information for a single contract award.")

endpoint("GET", "/api/v1/awards/{contractId}/burn-rate",
         "Get obligation burn rate analysis for a contract.",
         response_body='{\n  "contractId": "W91QV1-20-C-0001",\n  "totalObligated": 5000000.00,\n  "totalSpent": 3200000.00,\n  "burnRate": 0.64,\n  "monthlySpend": [\n    { "month": "2026-01", "amount": 150000.00 },\n    { "month": "2026-02", "amount": 175000.00 }\n  ],\n  "projectedExhaustion": "2026-09-15"\n}')

endpoint("GET", "/api/v1/awards/market-share",
         "Get market share data for a NAICS code.",
         params=[
             ["naicsCode", "string", "Yes", "NAICS code to analyze"],
             ["years", "int", "No", "Lookback period in years (default: 3, max: 10)"],
             ["limit", "int", "No", "Top N vendors (default: 10, max: 100)"],
         ],
         response_body='{\n  "naicsCode": "541512",\n  "totalContractValue": 150000000.00,\n  "vendors": [\n    { "name": "Vendor A", "uei": "AAA111", "contractCount": 25, "totalValue": 35000000 },\n    { "name": "Vendor B", "uei": "BBB222", "contractCount": 18, "totalValue": 22000000 }\n  ]\n}')

endpoint("GET", "/api/v1/awards/expiring",
         "Get contracts expiring within a time window, filtered for org relevance.",
         params=[
             ["monthsAhead", "int", "No", "Months to look ahead (default: 12, max: 24)"],
             ["naicsCode", "string", "No", "Filter by NAICS code"],
             ["setAsideType", "string", "No", "Filter by set-aside type"],
             ["limit", "int", "No", "Max results (default: 20, max: 100)"],
             ["offset", "int", "No", "Offset for pagination (default: 0)"],
         ],
         response_body='[\n  {\n    "contractId": "W91QV1-20-C-0001",\n    "vendorName": "Incumbent Corp",\n    "naicsCode": "541512",\n    "expirationDate": "2026-12-31",\n    "totalValue": 5000000.00,\n    "agency": "Department of Defense",\n    "setAside": "WOSB"\n  }\n]')

endpoint("GET", "/api/v1/awards/set-aside-trends/{naicsCode}",
         "Get set-aside trends over time for a NAICS code.",
         response_body='[\n  {\n    "year": 2024,\n    "setAsideType": "WOSB",\n    "contractCount": 45,\n    "totalValue": 12000000.00\n  }\n]')

endpoint("POST", "/api/v1/awards/{contractId}/load",
         "Request detailed data load for a contract.",
         request_body='{ "tier": "haiku" }')

endpoint("GET", "/api/v1/awards/{contractId}/load-status",
         "Check the status of a data load request for a contract.",
         response_body='{\n  "requestId": 10,\n  "status": "COMPLETE",\n  "message": "Data loaded successfully"\n}')

# =========================================================================
# CHAPTER 6: Entities
# =========================================================================
doc.add_heading("Chapter 6: Entities", level=1)
doc.add_paragraph(
    "Entity endpoints provide access to SAM.gov entity registration data. Entities represent "
    "government contractors and their registration details, capabilities, and exclusion status."
)
doc.add_paragraph("Base path: /api/v1/entities")

endpoint("GET", "/api/v1/entities", "Search registered entities.",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
         ])

endpoint("GET", "/api/v1/entities/{uei}", "Get detailed entity information by UEI (Unique Entity Identifier).")

endpoint("GET", "/api/v1/entities/{uei}/competitor-profile",
         "Get a competitor profile built from award history, NAICS codes, and agency relationships.",
         response_body='{\n  "uei": "XYZ123",\n  "legalBusinessName": "Competitor Corp",\n  "naicsCodes": ["541512", "541519"],\n  "topAgencies": ["DoD", "DHS"],\n  "recentAwards": [...],\n  "totalAwardValue": 45000000.00,\n  "contractCount": 32\n}')

endpoint("GET", "/api/v1/entities/{uei}/exclusion-check",
         "Check if an entity has active exclusions (debarment, suspension, etc.).",
         response_body='{\n  "uei": "XYZ123",\n  "hasActiveExclusions": false,\n  "exclusions": []\n}')

# =========================================================================
# CHAPTER 7: Prospects
# =========================================================================
doc.add_heading("Chapter 7: Prospects", level=1)
doc.add_paragraph(
    "Prospect endpoints manage the bid/no-bid pipeline. Prospects are org-scoped records that "
    "link an opportunity to a capture workflow with status tracking, team assignments, notes, "
    "and Go/No-Go scoring."
)
doc.add_paragraph("Base path: /api/v1/prospects")

endpoint("POST", "/api/v1/prospects", "Create a new prospect from an opportunity.",
         request_body='{\n  "noticeId": "abc123",\n  "initialStatus": "IDENTIFIED",\n  "notes": "Looks like a good fit for our cybersecurity team"\n}',
         response_body='{\n  "prospect": {\n    "prospectId": 42,\n    "noticeId": "abc123",\n    "status": "IDENTIFIED",\n    "createdAt": "2026-03-22T10:00:00Z"\n  }\n}',
         notes="Returns 201 Created with Location header.")

endpoint("GET", "/api/v1/prospects", "Search prospects with filters and pagination.",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
             ["sortBy", "string", "No", "Sort field"],
             ["sortDescending", "bool", "No", "Sort direction"],
         ])

endpoint("GET", "/api/v1/prospects/{id}", "Get full prospect detail including opportunity, notes, team, proposal, and score.",
         response_body='{\n  "prospectId": 42,\n  "noticeId": "abc123",\n  "status": "QUALIFIED",\n  "assignedUser": { "userId": 1, "displayName": "Jane Smith" },\n  "opportunity": { ... },\n  "notes": [ ... ],\n  "teamMembers": [ ... ],\n  "proposal": { ... },\n  "score": { "totalScore": 78, "criteria": [...] }\n}')

endpoint("PATCH", "/api/v1/prospects/{id}/status", "Update a prospect's pipeline status.",
         request_body='{\n  "status": "QUALIFIED",\n  "reason": "Passes all qualification checks"\n}',
         notes="Status must follow defined flow: IDENTIFIED -> QUALIFIED -> PURSUING -> PROPOSED -> WON/LOST/NO_BID.")

endpoint("PATCH", "/api/v1/prospects/{id}/reassign", "Reassign a prospect to a different user.",
         request_body='{\n  "assignToUserId": 3\n}')

endpoint("POST", "/api/v1/prospects/{id}/notes", "Add a note to a prospect.",
         request_body='{\n  "noteType": "GENERAL",\n  "content": "Spoke with the contracting officer about requirements."\n}',
         notes="Returns 201 Created.")

endpoint("POST", "/api/v1/prospects/{id}/team-members", "Add a team member to a prospect.",
         request_body='{\n  "userId": 3,\n  "role": "Technical Lead"\n}',
         notes="Returns 201 Created.")

endpoint("DELETE", "/api/v1/prospects/{id}/team-members/{memberId}",
         "Remove a team member from a prospect.",
         notes="Returns 204 No Content on success.")

endpoint("POST", "/api/v1/prospects/{id}/recalculate-score",
         "Recalculate the Go/No-Go score for a prospect.",
         response_body='{\n  "totalScore": 78,\n  "maxScore": 100,\n  "criteria": [\n    { "name": "Technical Capability", "score": 8, "maxScore": 10, "weight": 2.0 },\n    { "name": "Past Performance", "score": 7, "maxScore": 10, "weight": 1.5 }\n  ]\n}')

endpoint("POST", "/api/v1/prospects/auto-generate",
         "Generate prospects automatically from saved searches with auto-prospect enabled.",
         auth="JWT + OrgAdmin policy",
         request_body='{\n  "organizationId": 1\n}',
         response_body='{\n  "evaluated": 50,\n  "created": 8,\n  "skipped": 42,\n  "errors": [],\n  "searchResults": [...]\n}')

# =========================================================================
# CHAPTER 8: Proposals
# =========================================================================
doc.add_heading("Chapter 8: Proposals", level=1)
doc.add_paragraph(
    "Proposal endpoints manage the proposal lifecycle linked to prospects. "
    "Includes document tracking and milestone management."
)
doc.add_paragraph("Base path: /api/v1/proposals")

endpoint("GET", "/api/v1/proposals", "List/search proposals with pagination (org-wide view).",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
         ])

endpoint("POST", "/api/v1/proposals", "Create a new proposal linked to a prospect.",
         request_body='{\n  "prospectId": 42,\n  "title": "IT Support Services Proposal",\n  "estimatedValue": 2500000.00\n}',
         notes="Returns 201 Created.")

endpoint("PATCH", "/api/v1/proposals/{id}", "Update a proposal's status, value, win probability, or lessons learned.",
         request_body='{\n  "status": "SUBMITTED",\n  "winProbability": 0.65,\n  "lessonsLearned": "Strong technical approach"\n}',
         notes="Returns 409 Conflict if optimistic concurrency check fails.")

endpoint("POST", "/api/v1/proposals/{id}/documents", "Add a document record to a proposal.",
         request_body='{\n  "documentType": "TECHNICAL_VOLUME",\n  "title": "Technical Approach v2",\n  "url": "https://sharepoint.example.com/doc/123"\n}',
         notes="Returns 201 Created.")

endpoint("GET", "/api/v1/proposals/{id}/milestones", "List milestones for a proposal.")

endpoint("POST", "/api/v1/proposals/{id}/milestones", "Create a milestone for a proposal.",
         request_body='{\n  "title": "Pink Team Review",\n  "dueDate": "2026-04-01",\n  "description": "Initial compliance review"\n}',
         notes="Returns 201 Created.")

endpoint("PATCH", "/api/v1/proposals/{id}/milestones/{milestoneId}", "Update a milestone.",
         request_body='{\n  "status": "COMPLETED",\n  "completedAt": "2026-03-28T16:00:00Z"\n}')

# =========================================================================
# CHAPTER 9: Dashboard & Notifications
# =========================================================================
doc.add_heading("Chapter 9: Dashboard & Notifications", level=1)

doc.add_heading("Dashboard", level=2)
doc.add_paragraph("Base path: /api/v1/dashboard")

endpoint("GET", "/api/v1/dashboard", "Get the prospect pipeline dashboard with status counts, due items, and workload.",
         response_body='{\n  "statusCounts": {\n    "IDENTIFIED": 12,\n    "QUALIFIED": 8,\n    "PURSUING": 5,\n    "PROPOSED": 3\n  },\n  "dueSoon": [...],\n  "workloadByUser": [...],\n  "recentActivity": [...]\n}')

doc.add_heading("Notifications", level=2)
doc.add_paragraph("Base path: /api/v1/notifications")

endpoint("GET", "/api/v1/notifications", "List notifications for the current user.",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
             ["unreadOnly", "bool", "No", "Filter to unread only"],
         ])

endpoint("GET", "/api/v1/notifications/unread-count",
         "Get unread notification count (lightweight, for UI bell badge).",
         response_body='{ "unreadCount": 5 }')

endpoint("PATCH", "/api/v1/notifications/{id}/read", "Mark a single notification as read.",
         response_body='{ "message": "Notification marked as read." }')

endpoint("POST", "/api/v1/notifications/mark-all-read", "Mark all notifications as read.",
         response_body='{ "message": "All notifications marked as read." }')

# =========================================================================
# CHAPTER 10: Organization & Company Profile
# =========================================================================
doc.add_heading("Chapter 10: Organization & Company Profile", level=1)
doc.add_paragraph(
    "Organization endpoints manage org settings, membership, invites, company profile, "
    "certifications, NAICS codes, past performance, and entity linking."
)
doc.add_paragraph("Base path: /api/v1/org")

doc.add_heading("Organization Management", level=2)

endpoint("GET", "/api/v1/org", "Get the current user's organization details.")
endpoint("PATCH", "/api/v1/org", "Update the organization name.",
         auth="JWT + OrgAdmin",
         request_body='{ "name": "Acme Federal Solutions" }')
endpoint("GET", "/api/v1/org/members", "List members of the current organization.")
endpoint("POST", "/api/v1/org/invites", "Create an invite for a new user.",
         auth="JWT + OrgAdmin",
         request_body='{\n  "email": "newmember@example.com",\n  "orgRole": "member"\n}',
         notes="Returns 201 Created with invite code.")
endpoint("GET", "/api/v1/org/invites", "List pending invites.",
         auth="JWT + OrgAdmin")
endpoint("DELETE", "/api/v1/org/invites/{id}", "Revoke a pending invite.",
         auth="JWT + OrgAdmin",
         notes="Returns 204 No Content.")

doc.add_heading("Company Profile", level=2)

endpoint("GET", "/api/v1/org/profile", "Get the organization's company profile (size, revenue, capabilities).")
endpoint("PUT", "/api/v1/org/profile", "Update the company profile.",
         auth="JWT + OrgAdmin",
         request_body='{\n  "companySize": "Small",\n  "annualRevenue": 5000000,\n  "cageCode": "ABC12",\n  "dunsNumber": "123456789"\n}')

doc.add_heading("NAICS Codes", level=2)
endpoint("GET", "/api/v1/org/naics", "Get the organization's NAICS codes.")
endpoint("PUT", "/api/v1/org/naics", "Set (bulk replace) the organization's NAICS codes.",
         auth="JWT + OrgAdmin",
         request_body='[\n  { "naicsCode": "541512", "isPrimary": true },\n  { "naicsCode": "541519", "isPrimary": false }\n]')

doc.add_heading("Certifications", level=2)
endpoint("GET", "/api/v1/org/certifications", "Get the organization's certifications.")
endpoint("PUT", "/api/v1/org/certifications", "Set (bulk replace) certifications.",
         auth="JWT + OrgAdmin",
         request_body='[\n  { "certType": "WOSB", "certNumber": "WO-12345", "expirationDate": "2027-01-01" }\n]')

doc.add_heading("Past Performance", level=2)
endpoint("GET", "/api/v1/org/past-performance", "Get past performance records.")
endpoint("POST", "/api/v1/org/past-performance", "Add a past performance record.",
         auth="JWT + OrgAdmin",
         request_body='{\n  "contractNumber": "W91QV1-20-C-0001",\n  "agency": "U.S. Army",\n  "description": "IT Support Services",\n  "value": 2500000.00,\n  "startDate": "2020-01-01",\n  "endDate": "2024-12-31"\n}',
         notes="Returns 201 Created.")
endpoint("DELETE", "/api/v1/org/past-performance/{id}", "Delete a past performance record.",
         auth="JWT + OrgAdmin",
         notes="Returns 204 No Content.")

doc.add_heading("Entity Linking", level=2)
doc.add_paragraph(
    "Entity linking connects your organization to SAM.gov entity registrations. "
    "A 'SELF' link auto-populates company profile data from the entity record."
)
endpoint("GET", "/api/v1/org/entities", "Get all linked entities for the current organization.")
endpoint("POST", "/api/v1/org/entities", "Link an entity to the organization.",
         auth="JWT + OrgAdmin",
         request_body='{\n  "uei": "XYZ123",\n  "linkType": "SELF",\n  "isPrimary": true\n}',
         notes="Returns 201 Created. linkType values: SELF, PARTNER, COMPETITOR.")
endpoint("DELETE", "/api/v1/org/entities/{linkId}", "Deactivate an entity link.",
         auth="JWT + OrgAdmin",
         notes="Returns 204 No Content.")
endpoint("POST", "/api/v1/org/entities/refresh-self",
         "Refresh organization profile from the linked SELF entity.",
         auth="JWT + OrgAdmin")
endpoint("POST", "/api/v1/org/entities/resync-certs",
         "Re-sync certifications from all linked entities.",
         auth="JWT + OrgAdmin",
         response_body='{ "message": "Re-synced 5 certifications from linked entities" }')
endpoint("GET", "/api/v1/org/entities/aggregate-naics",
         "Get aggregate NAICS codes across all linked entities and manual entries.")

# =========================================================================
# CHAPTER 11: Admin, Reference, Saved Searches, Subawards
# =========================================================================
doc.add_heading("Chapter 11: Admin & Supporting Endpoints", level=1)

doc.add_heading("Admin Endpoints", level=2)
doc.add_paragraph("Base path: /api/v1/admin. Requires AdminAccess or SystemAdmin policy.")

endpoint("GET", "/api/v1/admin/etl-status", "Get ETL pipeline status, API usage, and recent errors.",
         auth="JWT + SystemAdmin")
endpoint("GET", "/api/v1/admin/load-history", "Get paginated ETL load history.",
         auth="JWT + SystemAdmin",
         params=[
             ["source", "string", "No", "Filter by source system"],
             ["status", "string", "No", "Filter by status (SUCCESS, FAILED)"],
             ["days", "int", "No", "Lookback period (default: 7)"],
             ["page", "int", "No", "Page number (default: 1)"],
             ["pageSize", "int", "No", "Results per page (default: 25)"],
         ])
endpoint("GET", "/api/v1/admin/health-snapshots", "Get health check snapshots.",
         auth="JWT + SystemAdmin",
         params=[["days", "int", "No", "Lookback period (default: 30)"]])
endpoint("GET", "/api/v1/admin/api-keys", "Get API key usage status.",
         auth="JWT + SystemAdmin")
endpoint("GET", "/api/v1/admin/jobs", "Get ETL job definitions with last-run status.",
         auth="JWT + SystemAdmin")

endpoint("GET", "/api/v1/admin/users", "List users in the admin's organization.",
         auth="JWT + AdminAccess",
         params=[
             ["page", "int", "No", "Page number (default: 1)"],
             ["pageSize", "int", "No", "Results per page (default: 25)"],
         ])
endpoint("PATCH", "/api/v1/admin/users/{id}", "Update a user's role, admin status, or active status.",
         auth="JWT + AdminAccess",
         request_body='{\n  "orgRole": "admin",\n  "isActive": true\n}')
endpoint("POST", "/api/v1/admin/users/{id}/reset-password", "Force-reset a user's password and revoke all sessions.",
         auth="JWT + AdminAccess")

endpoint("GET", "/api/v1/admin/organizations", "List all organizations.",
         auth="JWT + SystemAdmin")
endpoint("POST", "/api/v1/admin/organizations", "Create a new organization.",
         auth="JWT + SystemAdmin",
         request_body='{ "name": "Acme Corp", "slug": "acme" }',
         notes="Returns 201 Created. Returns 409 if slug already exists.")
endpoint("POST", "/api/v1/admin/organizations/{id}/owner", "Create the initial owner user for an organization.",
         auth="JWT + SystemAdmin",
         request_body='{\n  "email": "owner@acme.com",\n  "password": "InitialP@ss1",\n  "displayName": "Org Owner"\n}',
         notes="Returns 201 Created.")

doc.add_heading("Reference Data", level=2)
doc.add_paragraph("Base path: /api/v1/reference. Read-only lookup data.")

endpoint("GET", "/api/v1/reference/naics", "Search NAICS codes by code or description.",
         params=[["q", "string", "Yes", "Search query (min 2 chars)"]],
         notes="Returns up to 50 results.")
endpoint("GET", "/api/v1/reference/naics/{code}", "Get NAICS code detail with SBA size standard.")
endpoint("GET", "/api/v1/reference/certifications", "Get the list of available certification types.")

doc.add_heading("Saved Searches", level=2)
doc.add_paragraph("Base path: /api/v1/saved-searches. Per-user saved search management.")

endpoint("GET", "/api/v1/saved-searches", "List all active saved searches for the current user.")
endpoint("GET", "/api/v1/saved-searches/{id}", "Get a saved search by ID.")
endpoint("POST", "/api/v1/saved-searches", "Create a new saved search.",
         request_body='{\n  "name": "WOSB IT Opportunities",\n  "searchType": "opportunity",\n  "criteria": {\n    "setAside": "WOSB",\n    "naics": "541512",\n    "openOnly": true\n  },\n  "autoProspectEnabled": false\n}',
         notes="Returns 201 Created with Location header.")
endpoint("PATCH", "/api/v1/saved-searches/{id}", "Update a saved search.",
         request_body='{\n  "name": "Updated Search Name",\n  "autoProspectEnabled": true\n}')
endpoint("POST", "/api/v1/saved-searches/{id}/run", "Execute a saved search and return matching opportunities.")
endpoint("DELETE", "/api/v1/saved-searches/{id}", "Soft-delete a saved search.",
         notes="Returns 204 No Content.")

doc.add_heading("Subawards", level=2)
doc.add_paragraph("Base path: /api/v1/subawards. Subaward data for teaming analysis.")

endpoint("GET", "/api/v1/subawards/teaming-partners",
         "Search for potential teaming partners based on subaward history.",
         params=[
             ["page", "int", "No", "Page number"],
             ["pageSize", "int", "No", "Results per page"],
         ])
endpoint("GET", "/api/v1/subawards/by-prime/{primePiid}",
         "Get subawards for a specific prime contract.")

doc.add_heading("Health Check", level=2)
doc.add_paragraph("Base path: /health. Public endpoint, no authentication required.")

endpoint("GET", "/health", "Get system health status including database connectivity and ETL data freshness.",
         auth="None (public)",
         response_body='{\n  "status": "Healthy",\n  "database": {\n    "status": "Healthy",\n    "description": "Database connection successful"\n  },\n  "etlFreshness": {\n    "status": "Healthy",\n    "description": "6 data sources - most recent load 4h ago",\n    "data": {\n      "SAM_OPPORTUNITY_lastLoad": "2026-03-22 06:00",\n      "SAM_OPPORTUNITY_age": "4h",\n      "SAM_OPPORTUNITY_totalLoads": "142"\n    }\n  }\n}',
         notes="Status values: Healthy, Degraded, Unhealthy. A source is considered stale after 168 hours (7 days).")

# =========================================================================
# CHAPTER 12: Error Handling
# =========================================================================
doc.add_heading("Chapter 12: Error Handling", level=1)

doc.add_heading("Standard Error Response", level=2)
doc.add_paragraph(
    "All error responses follow a consistent JSON structure defined by ApiErrorResponse:"
)
add_code_block(
    '{\n'
    '  "statusCode": 400,\n'
    '  "message": "Human-readable error message",\n'
    '  "detail": "Optional technical detail for debugging",\n'
    '  "errors": {\n'
    '    "fieldName": ["Validation error 1", "Validation error 2"]\n'
    '  },\n'
    '  "traceId": "0HN1234ABCD:00000001"\n'
    '}'
)

doc.add_heading("HTTP Status Codes", level=2)
add_table(["Code", "Meaning", "When Used"], [
    ["200", "OK", "Successful GET, PATCH, or POST that returns data"],
    ["201", "Created", "Successful POST that creates a resource (prospects, notes, invites)"],
    ["204", "No Content", "Successful DELETE"],
    ["400", "Bad Request", "Validation errors, invalid state transitions, malformed input"],
    ["401", "Unauthorized", "Missing or invalid JWT, revoked session, missing org context"],
    ["403", "Forbidden", "Insufficient role/policy (e.g., non-admin accessing admin endpoints)"],
    ["404", "Not Found", "Resource does not exist or is not accessible to the current org"],
    ["409", "Conflict", "Duplicate resource (org slug), optimistic concurrency failure"],
    ["429", "Too Many Requests", "Rate limit exceeded. Check Retry-After header."],
    ["500", "Internal Server Error", "Unhandled exception (caught by ExceptionHandlerMiddleware)"],
])

doc.add_heading("Validation Errors", level=2)
doc.add_paragraph(
    "FluentValidation errors are caught by the FluentValidationFilter and returned as 400 responses "
    "with the errors dictionary populated by field name:"
)
add_code_block(
    '{\n'
    '  "statusCode": 400,\n'
    '  "message": "Validation failed",\n'
    '  "errors": {\n'
    '    "email": ["Email is required", "Email must be a valid email address"],\n'
    '    "password": ["Password must be at least 8 characters"]\n'
    '  },\n'
    '  "traceId": "0HN1234ABCD:00000002"\n'
    '}'
)

doc.add_heading("Rate Limit Response", level=2)
doc.add_paragraph("When rate limited, the response includes a Retry-After header:")
add_code_block(
    'HTTP/1.1 429 Too Many Requests\n'
    'Retry-After: 60\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '  "statusCode": 429,\n'
    '  "message": "Too many requests. Please try again later.",\n'
    '  "retryAfterSeconds": 60\n'
    '}'
)

# =========================================================================
# CHAPTER 13: Pagination
# =========================================================================
doc.add_heading("Chapter 13: Pagination", level=1)

doc.add_heading("Offset-Based Pagination", level=2)
doc.add_paragraph(
    "FedProspect uses offset-based pagination with page and pageSize query parameters. "
    "All paginated endpoints return a PagedResponse wrapper."
)

doc.add_heading("Request Parameters", level=2)
add_table(["Parameter", "Type", "Default", "Constraints"], [
    ["page", "int", "1", "Minimum 1"],
    ["pageSize", "int", "25", "Minimum 1, maximum 100"],
    ["sortBy", "string", "null", "Field name for sorting (endpoint-specific)"],
    ["sortDescending", "bool", "false", "true for descending order"],
])

doc.add_heading("Response Envelope", level=2)
add_code_block(
    '{\n'
    '  "items": [ ... ],\n'
    '  "page": 2,\n'
    '  "pageSize": 25,\n'
    '  "totalCount": 142,\n'
    '  "totalPages": 6,\n'
    '  "hasPreviousPage": true,\n'
    '  "hasNextPage": true\n'
    '}'
)

doc.add_heading("Rate Limits by Endpoint Category", level=2)
add_table(["Category", "Rate Limit", "Partitioned By", "Endpoints"], [
    ["auth", "10/min", "IP address", "Login, register, refresh"],
    ["login_global", "100/min", "Global (all IPs)", "Login only"],
    ["search", "60/min", "User ID (or IP)", "All GET/read endpoints"],
    ["write", "30/min", "User ID (or IP)", "POST, PATCH, DELETE mutations"],
    ["admin", "30/min", "User ID (or IP)", "Admin-only endpoints"],
])

doc.add_heading("Expiring Contracts (Legacy Pagination)", level=2)
doc.add_paragraph(
    "The GET /api/v1/awards/expiring endpoint uses limit/offset pagination instead of "
    "page/pageSize. This is a legacy pattern; new endpoints should use PagedRequest."
)
add_table(["Parameter", "Type", "Default", "Constraints"], [
    ["limit", "int", "20", "1-100"],
    ["offset", "int", "0", "0-10000"],
])

# =========================================================================
# Save
# =========================================================================
out_path = r"c:\git\fedProspect\docs\api\FedProspect-API-Guide.docx"
doc.save(out_path)
print(f"Saved API guide to {out_path}")
print(f"File size: {os.path.getsize(out_path):,} bytes")
