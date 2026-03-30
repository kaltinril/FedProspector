"""Generate FedProspect-Loader-Reference.docx using python-docx.

Run from fed_prospector/:
    .venv/Scripts/python ../docs/etl/generate_loader_reference.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import date


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading_elm.append(shading)


def add_code_block(doc, text):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def build_document():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("FedProspect ETL Loader Reference", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Generated {date.today().isoformat()}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Comprehensive reference for all 15 ETL loaders in the "
        "Federal Contract Prospecting System."
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS placeholder
    # =========================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  Opportunity Loader",
        "2.  Award Loader",
        "3.  Entity Loader",
        "4.  USASpending Bulk Loader",
        "5.  USASpending API Loader",
        "6.  GSA CALC+ Loader",
        "7.  Federal Hierarchy Loader",
        "8.  Exclusion Loader",
        "9.  Subaward Loader",
        "10. Resource Link Resolver",
        "11. Attachment Downloader",
        "12. Attachment Text Extractor",
        "13. Attachment Intel Extractor",
        "14. Description Backfill Loader",
        "15. On-Demand Loader",
        "A.  Daily Load Sequence",
        "B.  Batch Load Commands",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =========================================================================
    # LOADER SECTIONS
    # =========================================================================

    # ----- 1. Opportunity Loader -----
    doc.add_heading("1. Opportunity Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads contract opportunities from the SAM.gov Opportunities API. "
        "Fetches opportunities matching date range, NAICS, and set-aside filters, "
        "then upserts into the local database with SHA-256 change detection and "
        "field-level history tracking. Supports resumable partial loads."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load opportunities [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Opportunities API (v3)")
    doc.add_paragraph("API Client: SAMOpportunityClient (api_clients/sam_opportunity_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity", "Target", "Main opportunity records"],
            ["opportunity_history", "History", "Field-level change tracking"],
            ["stg_opportunity_raw", "Staging", "Raw JSON staging before normalization"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Incremental upsert with SHA-256 change detection. "
        "Each API page is processed as a batch: raw JSON is staged, then normalized "
        "and upserted. Changed fields are tracked in opportunity_history. "
        "Progress is saved after every page for crash-safe resume. "
        "Historical mode breaks 2-year ranges into 1-year chunks."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 1: 10 calls/day (free tier). "
        "API key 2: 1,000 calls/day. "
        "Default call budget: 5 (reserves calls for other loaders). "
        "Each page = 1 API call."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of 26 business fields (notice_id, title, solicitation_number, "
        "dates, set-aside, NAICS, location, award info). "
        "Excludes timestamps, description, links, and load-tracking columns."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--days-back", "7", "Load opportunities posted in the last N days"],
            ["--set-aside", "None", "Set-aside filter: WOSB, 8A, all, none, or comma-separated"],
            ["--naics", "None", "NAICS code(s), comma-separated"],
            ["--posted-from", "None", "Start date MM/dd/yyyy (overrides --days-back)"],
            ["--posted-to", "None", "End date MM/dd/yyyy (defaults to today)"],
            ["--historical", "False", "Load 2 years of historical data"],
            ["--max-calls", "5", "Max API calls for this invocation"],
            ["--key", "1", "Which SAM API key to use (1 or 2)"],
            ["--force", "False", "Ignore previous progress, start fresh"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load opportunities")
    add_code_block(doc, "python main.py load opportunities --days-back=30 --key 2")
    add_code_block(doc, "python main.py load opportunities --set-aside=WOSB --naics=541511")
    add_code_block(doc, "python main.py load opportunities --historical --max-calls=20")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- API key 1 has only 10 calls/day; use --key=2 for bulk loads.\n"
        "- Multiple NAICS + set-aside combos multiply API calls.\n"
        "- 429 rate limit errors save progress automatically; re-run to resume.\n"
        "- Feb 29 start dates are adjusted to Mar 1 (SAM.gov rejects leap day)."
    )

    doc.add_page_break()

    # ----- 2. Award Loader -----
    doc.add_heading("2. Award Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads FPDS contract awards from the SAM.gov Contract Awards API. "
        "Supports watermark-based incremental loading, resume after interruption, "
        "and default NAICS + set-aside filters for scheduled runs. "
        "Also supports per-organization loading by UEI."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load awards [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Contract Awards API (FPDS)")
    doc.add_paragraph("API Client: SAMAwardsClient (api_clients/sam_awards_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["fpds_contract", "Target", "FPDS contract award records (PK: contract_id + modification_number)"],
            ["stg_fpds_award_raw", "Staging", "Raw JSON staging before normalization"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Offset-based pagination (100 records/page). Each NAICS + set-aside combination "
        "is iterated separately. Watermark-based incremental loading resumes from the "
        "last successful load date. Composite PK (contract_id, modification_number) "
        "with SHA-256 change detection on 16 business fields."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 2 (default): 1,000 calls/day. "
        "Default call budget: 10. Each page = 1 API call."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of 16 fields: contract_id, modification_number, vendor UEI/name, "
        "dollars, NAICS, PSC, set-aside, dates, competition info. "
        "Uses pipe-separated composite key for hash lookups."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--naics", "None", "NAICS code(s), comma-separated"],
            ["--set-aside", "None", "Set-aside type (WOSB, 8A, SBA, etc.)"],
            ["--agency", "None", "Contracting department CGAC code"],
            ["--awardee-uei", "None", "Awardee UEI to search"],
            ["--piid", "None", "Contract PIID to search"],
            ["--for-org", "None", "Load awards for all UEIs linked to an org (name or ID)"],
            ["--years-back", "None", "Years of history to load"],
            ["--days-back", "None", "Days of history (overrides --years-back)"],
            ["--fiscal-year", "None", "Specific fiscal year"],
            ["--date-from", "None", "Start date YYYY-MM-DD"],
            ["--date-to", "None", "End date YYYY-MM-DD"],
            ["--max-calls", "10", "Max API calls for this invocation"],
            ["--key", "2", "Which SAM API key to use (1 or 2)"],
            ["--force", "False", "Skip resume, start fresh"],
            ["--dry-run", "False", "Show what would load without API calls"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load awards")
    add_code_block(doc, "python main.py load awards --naics 541512 --years-back 5")
    add_code_block(doc, "python main.py load awards --set-aside WOSB --days-back 10 --key 2")
    add_code_block(doc, "python main.py load awards --for-org \"Acme Corp\" --dry-run")
    add_code_block(doc, "python main.py load awards --fiscal-year 2025")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- When no filters are specified, uses tracked NAICS and set-asides from settings.\n"
        "- Stale RUNNING loads are auto-cleaned on startup.\n"
        "- Budget exhaustion saves progress; re-run to resume.\n"
        "- Use replay-awards to reprocess staged records after fixing normalization bugs."
    )

    doc.add_page_break()

    # ----- 3. Entity Loader -----
    doc.add_heading("3. Entity Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads SAM.gov entity (vendor) registration data. Supports two methods: "
        "API queries with date-based incremental loading (default), and monthly "
        "bulk extract via ZIP download with LOAD DATA INFILE. Normalizes nested "
        "JSON into 1 parent table and 8 child tables."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load entities [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Entity Management API / Monthly Bulk Extract ZIP")
    doc.add_paragraph(
        "API Clients: SAMEntityClient (API), SAMExtractClient (bulk extract)"
    )

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["entity", "Target", "Main entity records (PK: uei_sam)"],
            ["entity_address", "Child", "Physical and mailing addresses"],
            ["entity_naics", "Child", "NAICS codes per entity"],
            ["entity_poc", "Child", "Points of contact"],
            ["entity_sba_certification", "Child", "SBA certifications (WOSB, 8a, HUBZone)"],
            ["entity_far_response", "Child", "FAR provision responses"],
            ["entity_goods_and_services", "Child", "Product/service codes"],
            ["entity_url", "Child", "Entity website URLs"],
            ["entity_congressional_district", "Child", "Congressional district mappings"],
            ["stg_entity_raw", "Staging", "Raw JSON staging"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "API mode (default): Date-based page-by-page loading (10 entities/page) with resume support. "
        "Monthly mode: Downloads bulk extract ZIP, parses DAT file format, loads via LOAD DATA INFILE. "
        "Both use SHA-256 change detection. Entity data is normalized from deeply nested JSON."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "API mode: SAM.gov API key limits (10 or 1000/day). Default max-calls: 100 (~1000 entities). "
        "Monthly extract: 1 API call to download."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of ~30 business fields including UEI, CAGE, registration status, "
        "legal name, DBA, entity structure, and incorporation info."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--type", "api", "Load method: api or monthly"],
            ["--date", "today", "Date for API query (YYYY-MM-DD)"],
            ["--year", "current", "Year for monthly extract"],
            ["--month", "current", "Month for monthly extract"],
            ["--file", "None", "Load from local file (.dat or .json)"],
            ["--key", "1", "SAM API key (1 or 2)"],
            ["--uei", "None", "Filter by UEI (API only)"],
            ["--name", "None", "Filter by entity name (API only)"],
            ["--naics", "None", "NAICS codes, comma-separated (API only)"],
            ["--set-aside", "None", "Business type code: 8W=WOSB, 8E=EDWOSB, A4=8(a)"],
            ["--status", "A", "Registration status: A=active, E=expired"],
            ["--max-calls", "100", "Max API calls safety cap (API only)"],
            ["--force", "False", "Force reload even if already loaded"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load entities")
    add_code_block(doc, "python main.py load entities --type=monthly")
    add_code_block(doc, "python main.py load entities --naics=541512 --set-aside=8W --key=2")
    add_code_block(doc, "python main.py load entities --file=SAM_PUBLIC_MONTHLY.dat")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Monthly extract files are large (multi-GB); ensure sufficient disk space.\n"
        "- DAT files are auto-cleaned after successful load.\n"
        "- API filter options are ignored when --type=monthly.\n"
        "- Resumable page-by-page loading for API mode."
    )

    doc.add_page_break()

    # ----- 4. USASpending Bulk Loader -----
    doc.add_heading("4. USASpending Bulk Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Bulk loads USASpending.gov contract award data from fiscal-year CSV archives. "
        "Downloads pre-built ZIP archives or on-demand exports, extracts CSV files, "
        "and loads via LOAD DATA INFILE for maximum performance. Supports full fiscal "
        "year loads, delta (monthly change) loads, and recent-days loads."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load usaspending-bulk [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("USASpending.gov Bulk Download API / Pre-built Archive Files")
    doc.add_paragraph("API Client: USASpendingClient (api_clients/usaspending_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["usaspending_award", "Target", "USASpending contract award records"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Downloads fiscal-year ZIP archives containing CSV files. CSVs are transformed "
        "into tab-separated temp files and loaded via LOAD DATA INFILE with upsert "
        "semantics. Fast mode drops secondary indexes during load and rebuilds after. "
        "Delta mode processes the monthly change file (inserts, updates, and deletes)."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "No API key required. No daily call limits. Archive downloads are bandwidth-limited "
        "by USASpending.gov servers."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "No hash-based change detection. Uses LOAD DATA INFILE with ON DUPLICATE KEY UPDATE "
        "for upsert semantics. Delta files include explicit delete markers."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--years-back", "5", "Number of recent fiscal years to load"],
            ["--fiscal-year", "None", "Load a single fiscal year"],
            ["--days-back", "None", "Load recent N days via on-demand API"],
            ["--delta", "False", "Download and process monthly delta file"],
            ["--skip-download", "False", "Use previously downloaded files"],
            ["--source", "archive", "Download source: archive (fast) or api (on-demand)"],
            ["--fast", "False", "Drop secondary indexes during load (faster bulk inserts)"],
            ["--check-available", "False", "List available archive/delta files without loading"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load usaspending-bulk")
    add_code_block(doc, "python main.py load usaspending-bulk --fiscal-year 2025")
    add_code_block(doc, "python main.py load usaspending-bulk --days-back 5")
    add_code_block(doc, "python main.py load usaspending-bulk --delta")
    add_code_block(doc, "python main.py load usaspending-bulk --check-available")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Archive files can be several GB; ensure sufficient disk space.\n"
        "- --fast mode makes queries slow until indexes are rebuilt.\n"
        "- If interrupted during --fast, indexes are auto-rebuilt on next run.\n"
        "- --days-back and --delta are mutually exclusive.\n"
        "- MySQL secure-file-priv must allow LOAD DATA INFILE from temp directory."
    )

    doc.add_page_break()

    # ----- 5. USASpending API Loader -----
    doc.add_heading("5. USASpending API Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads individual award transaction history from the USASpending.gov API. "
        "Fetches per-modification funding timeline for a specific award, enabling "
        "burn rate analysis. Uses the award's generated_unique_award_id."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load usaspending [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("USASpending.gov Award Transaction API")
    doc.add_paragraph("API Client: USASpendingClient (api_clients/usaspending_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["usaspending_transaction", "Target", "Per-modification transaction records"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Fetches all transactions for a single award ID and batch-inserts into "
        "usaspending_transaction. Award must exist in usaspending_award first "
        "(loaded via usaspending-bulk). Can resolve award by --solicitation or --piid."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph("No API rate limits. USASpending has no daily quotas.")

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of business fields. Uses batch upsert with change detection."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--award-id", "None", "USASpending generated_unique_award_id"],
            ["--solicitation", "None", "Solicitation number to find award"],
            ["--piid", "None", "Contract PIID to find award"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load usaspending --award-id CONT_AWD_W911NF25C0001_9700_-NONE-_-NONE-")
    add_code_block(doc, "python main.py load usaspending --solicitation W911NF-25-R-0001")
    add_code_block(doc, "python main.py load usaspending --piid W911NF25C0001")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Award must exist in usaspending_award table first.\n"
        "- One of --award-id, --solicitation, or --piid is required.\n"
        "- Used primarily for burn rate analysis."
    )

    doc.add_page_break()

    # ----- 6. GSA CALC+ Loader -----
    doc.add_heading("6. GSA CALC+ Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads GSA CALC+ labor rate data. Fetches the full ceiling rates dataset via "
        "the CALC+ v3 API, truncates the gsa_labor_rate table, and reloads all records. "
        "Uses multi-sort de-duplication to work around the Elasticsearch 10K window limit, "
        "recovering ~122K unique rates from ~230K total."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load labor-rates")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("GSA CALC+ API v3 (Elasticsearch-backed, refreshed nightly by GSA)")
    doc.add_paragraph("API Client: CalcPlusClient (api_clients/calc_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["gsa_labor_rate", "Target", "GSA schedule ceiling rates (truncate + reload)"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Full refresh: TRUNCATE gsa_labor_rate, then bulk insert all rates. "
        "No incremental mode since the entire dataset is replaced on each load. "
        "Multiple API queries with different sort orderings maximize coverage beyond "
        "the Elasticsearch 10K result window."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph("No authentication or API key required. No rate limits.")

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph("None (full truncate + reload on every run).")

    doc.add_heading("CLI Options", level=2)
    doc.add_paragraph("No options. Simply run the command.")

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load labor-rates")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- These are GSA schedule CEILING rates, not SCA wage determinations.\n"
        "- The Elasticsearch 10K window limit means not all rates can be retrieved in a single query.\n"
        "- Typically takes 30-60 seconds to complete."
    )

    doc.add_page_break()

    # ----- 7. Federal Hierarchy Loader -----
    doc.add_heading("7. Federal Hierarchy Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads the federal agency organizational hierarchy from the SAM.gov Federal "
        "Hierarchy API. Includes departments (Level 1), sub-tier agencies (Level 2), "
        "and offices (Level 3). Two commands: 'load hierarchy' for Levels 1-2, "
        "and 'load offices' for Level 3."
    )

    doc.add_heading("CLI Commands", level=2)
    add_code_block(doc, "python main.py load hierarchy [OPTIONS]")
    add_code_block(doc, "python main.py load offices [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Federal Hierarchy API")
    doc.add_paragraph("API Client: SAMFedHierClient (api_clients/sam_fedhier_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["federal_organization", "Target", "Org hierarchy (PK: fh_org_id), self-referencing parent_org_id"],
            ["stg_fedhier_raw", "Staging", "Raw JSON staging before normalization"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Hierarchy (L1-L2): Offset-based pagination (100 records/page). Supports full "
        "refresh (TRUNCATE + reload) and incremental upsert with SHA-256 change detection. "
        "Page-by-page progress saves for crash-safe resume.\n\n"
        "Offices (L3): Iterates Level 2 sub-tiers from the database, fetching child "
        "organizations for each. ~738 sub-tiers, one API call per sub-tier. Saves "
        "progress per sub-tier for resume."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 2 (default): 1,000 calls/day. "
        "Hierarchy default budget: 50. Office default budget: 200."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of 10 fields: fh_org_id, name, type, description, status, "
        "agency_code, CGAC, parent_org_id, level."
    )

    doc.add_heading("CLI Options (hierarchy)", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--status", "Active", "Organization status filter (Active/Inactive)"],
            ["--max-calls", "50", "Max API calls"],
            ["--full-refresh", "False", "TRUNCATE and reload all data"],
            ["--key", "2", "SAM API key (1 or 2)"],
            ["--force", "False", "Ignore previous progress"],
        ]
    )

    doc.add_heading("CLI Options (offices)", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--max-calls", "200", "Max API calls"],
            ["--key", "2", "SAM API key (1 or 2)"],
            ["--force", "False", "Ignore previous progress"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load hierarchy")
    add_code_block(doc, "python main.py load hierarchy --full-refresh --key 2")
    add_code_block(doc, "python main.py load offices --max-calls 100")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Office load requires Level 2 sub-tiers to be loaded first.\n"
        "- ~738 sub-tiers means the office load needs multiple runs with --max-calls budgeting.\n"
        "- Both commands support resume; re-run to continue from where you left off."
    )

    doc.add_page_break()

    # ----- 8. Exclusion Loader -----
    doc.add_heading("8. Exclusion Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads active exclusion records from the SAM.gov Exclusions API. "
        "These records identify debarred/suspended entities that are ineligible "
        "for federal contracts. Used for teaming partner risk assessment."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load exclusions [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Exclusions API")
    doc.add_paragraph("API Client: SAMExclusionsClient (api_clients/sam_exclusions_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["sam_exclusion", "Target", "Exclusion records (auto-increment PK)"],
            ["stg_exclusion_raw", "Staging", "Raw JSON staging"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Page-based pagination (max 10 records per page). Collects all exclusions first, "
        "then batch upserts with SHA-256 change detection. Logical key: "
        "uei + activation_date + exclusion_type."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 2 (default): 1,000 calls/day. Default budget: 20. "
        "Max 10 records per page = many pages needed for full load."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of 19 fields including UEI, CAGE, entity/person name, "
        "exclusion type/program, agency, dates, and classification."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--exclusion-type", "None", "Filter by exclusion type"],
            ["--agency", "None", "Excluding agency code filter"],
            ["--max-calls", "20", "Max API calls"],
            ["--key", "2", "SAM API key (1 or 2)"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load exclusions")
    add_code_block(doc, "python main.py load exclusions --agency DOD --key 2")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Only 10 records per page makes this loader call-intensive.\n"
        "- Budget exhaustion is reported with remaining record count.\n"
        "- Related commands: 'search exclusions' and 'analyze scan-exclusions'."
    )

    doc.add_page_break()

    # ----- 9. Subaward Loader -----
    doc.add_heading("9. Subaward Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Loads subaward/subcontract data from the SAM.gov Acquisition Subaward "
        "Reporting API. Supports PIID-driven loading (from local fpds_contract table), "
        "direct PIID lookup, and agency-based pagination. Enables teaming partner analysis."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py load subawards [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Acquisition Subaward Reporting API")
    doc.add_paragraph("API Client: SAMSubawardClient (api_clients/sam_subaward_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["sam_subaward", "Target", "Subaward records (auto-increment PK)"],
            ["stg_subaward_raw", "Staging", "Raw JSON staging"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Three modes:\n"
        "  1. PIID-driven (--naics): Queries fpds_contract for PIIDs in given NAICS codes, "
        "then fetches subawards per PIID. Resumable per PIID.\n"
        "  2. Direct PIID (--piid): Fetches subawards for a single prime contract.\n"
        "  3. Agency mode (--agency): Pages through all subawards for an agency.\n\n"
        "Note: FAR 52.204-10 requires subaward reporting only for primes >= $750K. "
        "Default --min-amount filters accordingly."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 2 (default): 1,000 calls/day. Default budget: 20."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of 18 fields including prime/sub UEI/name, amounts, dates, "
        "NAICS, PSC, business type, and location."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--naics", "None", "NAICS code(s), comma-separated (PIID-driven strategy)"],
            ["--agency", "None", "Agency code filter"],
            ["--piid", "None", "Prime contract PIID (direct lookup)"],
            ["--years-back", "2", "Years of contract history to search for PIIDs"],
            ["--max-calls", "20", "Max API calls"],
            ["--min-amount", "750000", "Min prime contract amount to include"],
            ["--force", "False", "Ignore previous progress"],
            ["--key", "2", "SAM API key (1 or 2)"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py load subawards --naics 541512")
    add_code_block(doc, "python main.py load subawards --piid W91QVN-20-C-0001")
    add_code_block(doc, "python main.py load subawards --agency 9700 --key 2")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- At least one filter required (--naics, --agency, or --piid).\n"
        "- PIID-driven mode requires fpds_contract data (load awards first).\n"
        "- Without API filters, the API returns 2.7M+ records across 2,700+ pages.\n"
        "- Related commands: 'search subawards' and 'analyze teaming'."
    )

    doc.add_page_break()

    # ----- 10. Resource Link Resolver -----
    doc.add_heading("10. Resource Link Resolver", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Enriches opportunity resource links with filename and content-type metadata. "
        "Makes HEAD requests to SAM.gov resource link URLs, extracting Content-Disposition "
        "filenames and Content-Type headers from the 303 redirect response. "
        "Does not follow redirects to S3."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py update link-metadata [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph(
        "SAM.gov resource link URLs (embedded in opportunity.resource_links JSON column). "
        "No API key required."
    )

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity", "Target", "Updates resource_links JSON with filename/content_type"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Batch processing: finds opportunities with un-enriched resource links, "
        "sends parallel HEAD requests (up to 5 concurrent), extracts metadata from "
        "303 responses, and updates the resource_links JSON column with enriched objects. "
        "SSRF protection: only allows https://sam.gov/ and https://api.sam.gov/ prefixes. "
        "~30ms per request."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph("No API key needed. No rate limits. 100ms delay between requests.")

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph("Processes only un-enriched links (--missing-only is default behavior).")

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--missing-only", "True", "Only process un-enriched resource links"],
            ["--batch-size", "100", "Number of opportunities per batch"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py update link-metadata")
    add_code_block(doc, "python main.py update link-metadata --batch-size 50")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Prerequisite: opportunities must be loaded first.\n"
        "- This step is a prerequisite for attachment downloading.\n"
        "- SSRF protection blocks non-SAM.gov URLs."
    )

    doc.add_page_break()

    # ----- 11. Attachment Downloader -----
    doc.add_heading("11. Attachment Downloader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Downloads opportunity attachment files (PDFs, DOCXs, etc.) from SAM.gov "
        "and stores them locally for text extraction. Validates redirects (must go to "
        "*.amazonaws.com), computes SHA-256 content hashes, and tracks metadata in "
        "the opportunity_attachment table."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py download attachments [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph(
        "SAM.gov resource link download URLs. Redirects to S3 (amazonaws.com)."
    )

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity_attachment", "Target", "Attachment metadata and download status"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )
    doc.add_paragraph(
        "Files stored in: E:\\fedprospector\\attachments\\ (configurable via ATTACHMENT_DIR env var)"
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Concurrent downloads (default 5 threads). Stream-downloads files in 8KB chunks. "
        "Validates redirects against *.amazonaws.com. Computes SHA-256 content hash. "
        "Supports --missing-only to skip already-downloaded files and --active-only "
        "to limit to opportunities with future response deadlines."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "No API key needed. Configurable delay between requests (default 0.1s). "
        "60-second download timeout."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 content hash stored per attachment. --check-changed flag re-downloads "
        "if remote file has changed."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--notice-id", "None", "Process a single opportunity by notice ID"],
            ["--batch-size", "100", "Attachments per batch"],
            ["--max-file-size", "50", "Maximum file size in MB"],
            ["--missing-only", "False", "Only download attachments not yet stored"],
            ["--check-changed", "False", "Re-download if remote file changed"],
            ["--delay", "0.1", "Delay in seconds between requests"],
            ["--active-only", "False", "Only for opps with future deadlines"],
            ["--workers", "5", "Concurrent download threads"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py download attachments --missing-only --active-only --batch-size 5000")
    add_code_block(doc, "python main.py download attachments --notice-id abc123")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Prerequisite: resource link metadata must be enriched first (update link-metadata).\n"
        "- Files can consume significant disk space; use --active-only for daily loads.\n"
        "- File cleanup is gated on completion of all downstream pipeline stages."
    )

    doc.add_page_break()

    # ----- 12. Attachment Text Extractor -----
    doc.add_heading("12. Attachment Text Extractor", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Extracts text content from downloaded attachment files. Supports PDF (via PyMuPDF), "
        "DOCX, XLSX, PPTX, DOC (via LibreOffice), and plain text formats. Produces "
        "annotated Markdown with heading hierarchy, bold markers, and table formatting. "
        "PDF table detection uses find_tables() for documents under 30 pages."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py extract attachment-text [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("Local attachment files downloaded by the Attachment Downloader.")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity_attachment", "Target", "Updates extracted_text, extraction_status, text_hash"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Batch processing with concurrent extraction threads (default 10). "
        "Auto-detects file format by magic bytes (not just extension). "
        "Extracts text page-by-page for PDFs, with parallel page extraction for large documents. "
        "Stores extracted text in the opportunity_attachment table."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph("Local processing only. No API calls.")

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "SHA-256 hash of extracted text. Skips files already extracted unless --force."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--notice-id", "None", "Process a single opportunity by notice ID"],
            ["--batch-size", "100", "Attachments per batch"],
            ["--force", "False", "Re-extract even if already extracted"],
            ["--workers", "10", "Concurrent extraction threads"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py extract attachment-text")
    add_code_block(doc, "python main.py extract attachment-text --batch-size 5000")
    add_code_block(doc, "python main.py extract attachment-text --force --notice-id abc123")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- PDF table detection (find_tables) is skipped for >30 page PDFs (performance).\n"
        "- DOC files require LibreOffice to be installed for conversion.\n"
        "- Malformed PDFs may produce MuPDF warnings (suppressed by default).\n"
        "- openpyxl warnings about unsupported Excel features are suppressed."
    )

    doc.add_page_break()

    # ----- 13. Attachment Intel Extractor -----
    doc.add_heading("13. Attachment Intel Extractor", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Extracts structured intelligence from attachment text and opportunity descriptions. "
        "Scans annotated Markdown for security clearance requirements, evaluation methods "
        "(LPTA vs Best Value), contract vehicle types, and recompete signals. "
        "Uses regex/keyword extraction with structure-aware confidence boosting."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py extract attachment-intel [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("Extracted text from opportunity_attachment.extracted_text column.")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity_attachment_intel", "Target", "Structured intel records per attachment"],
            ["opportunity_intel_source", "Provenance", "Source tracking for each intel finding"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Batch processing of attachments with extraction_status='extracted'. "
        "Applies compiled regex patterns against text, boosting confidence when "
        "matches appear under relevant headings. Intel categories: clearance_level, "
        "eval_method, vehicle_type, recompete."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph("Local processing only. No API calls.")

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph("Skips attachments already processed unless --force.")

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--notice-id", "None", "Process a single opportunity by notice ID"],
            ["--batch-size", "100", "Attachments per batch"],
            ["--method", "keyword", "Extraction method: keyword, regex, or hybrid"],
            ["--force", "False", "Re-extract even if already extracted"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py extract attachment-intel")
    add_code_block(doc, "python main.py extract attachment-intel --method regex")
    add_code_block(doc, "python main.py extract attachment-intel --notice-id abc123 --force")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Requires text extraction to be completed first.\n"
        "- 'Secret' matches may produce false positives (mitigated by negative lookbehind).\n"
        "- Structure-aware confidence boosting depends on heading detection in annotated Markdown."
    )

    doc.add_page_break()

    # ----- 14. Description Backfill Loader -----
    doc.add_heading("14. Description Backfill Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Fetches and caches full description text for opportunities from SAM.gov. "
        "Opportunities are loaded with a description_url but not the full HTML text; "
        "this loader backfills the description_text column by fetching each URL individually. "
        "Uses a two-pass priority strategy: high-value opportunities (matching target NAICS "
        "and WOSB/8(a)/SBA set-asides) are fetched first, then remaining budget is used "
        "for the general backlog."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py update fetch-descriptions [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph("SAM.gov Opportunity Description URLs (one GET request per opportunity)")
    doc.add_paragraph("API Client: SAMOpportunityClient (api_clients/sam_opportunity_client.py)")

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["opportunity", "Target", "Updates description_text and description_fetched_at columns"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Two-pass priority fetch:\n"
        "  Pass 1 (priority): Fetch descriptions for opportunities matching --naics and "
        "--set-aside filters (e.g., WOSB, 8A, SBA + target NAICS codes).\n"
        "  Pass 2 (general): Use remaining budget for all other opportunities missing descriptions.\n\n"
        "Each description is fetched individually (1 API call per opportunity). "
        "HTML response is parsed and stored as plain text. "
        "Integrated into daily_load.bat as step 2 with --limit 100."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "SAM.gov API key 2 (default): 1,000 calls/day. "
        "Each description = 1 API call. Daily batch limited to 100 via --limit. "
        "Uses --key 2 by default."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "Processes only opportunities where description_text IS NULL (--missing-only default). "
        "Use --all to re-fetch previously cached descriptions."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--missing-only/--all", "missing-only", "Only fetch for rows missing description_text"],
            ["--batch-size", "100", "Opportunities per commit batch"],
            ["--days-back", "None", "Only fetch for opps posted in the last N days"],
            ["--notice-id", "None", "Fetch for a single notice ID"],
            ["--key", "2", "SAM API key (1 or 2)"],
            ["--naics", "None", "Priority NAICS codes (comma-separated, fetched first)"],
            ["--set-aside", "None", "Priority set-aside codes (e.g., WOSB,8A,SBA)"],
            ["--limit", "None", "Max total descriptions to fetch"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py update fetch-descriptions")
    add_code_block(doc, "python main.py update fetch-descriptions --days-back 7")
    add_code_block(doc, "python main.py update fetch-descriptions --notice-id abc123 --key 1")
    add_code_block(doc, "python main.py update fetch-descriptions --naics 541511,541512 --set-aside WOSB,8A --limit 100")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Each description = 1 API call; use --limit to control daily budget consumption.\n"
        "- Priority pass ensures WOSB/8(a)/SBA opportunities are filled first.\n"
        "- Some opportunities have no description URL (description_url IS NULL); these are skipped.\n"
        "- Integrated into daily_load.bat as step 2 (100/day batch)."
    )

    doc.add_page_break()

    # ----- 15. On-Demand Loader -----
    doc.add_heading("15. On-Demand Loader", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Processes pending on-demand data load requests inserted by the C# API. "
        "Polls the data_load_request table for PENDING rows and routes to the "
        "appropriate loader. Supports four request types: "
        "USASPENDING_AWARD, FPDS_AWARD, REFRESH_FEDHIER_ORG (single org refresh "
        "from SAM.gov Federal Hierarchy API), and ATTACHMENT_ANALYSIS (AI analysis "
        "of opportunity attachments)."
    )

    doc.add_heading("CLI Command", level=2)
    add_code_block(doc, "python main.py job process-requests [OPTIONS]")

    doc.add_heading("Data Source", level=2)
    doc.add_paragraph(
        "USASpending.gov API, SAM.gov Awards API, SAM.gov Federal Hierarchy API, "
        "and local attachment processing, triggered by rows in data_load_request."
    )
    doc.add_paragraph(
        "API Clients: USASpendingClient, SAMAwardsClient, SAMFedHierClient"
    )

    doc.add_heading("Tables Populated", level=2)
    add_table(doc,
        ["Table", "Type", "Description"],
        [
            ["data_load_request", "Queue", "Request queue (status: PENDING -> COMPLETE/FAILED)"],
            ["usaspending_award", "Target", "USASpending award records (for USASPENDING_AWARD type)"],
            ["usaspending_transaction", "Target", "Transaction records (for USASPENDING_AWARD type)"],
            ["fpds_contract", "Target", "FPDS contract records (for FPDS_AWARD type)"],
            ["federal_organization", "Target", "Agency/office org records (for REFRESH_FEDHIER_ORG type)"],
            ["opportunity_attachment_intel", "Target", "AI analysis results (for ATTACHMENT_ANALYSIS type)"],
            ["etl_load_log", "Tracking", "Load progress and statistics"],
        ]
    )

    doc.add_heading("Load Strategy", level=2)
    doc.add_paragraph(
        "Polls data_load_request for up to 10 PENDING rows per invocation. "
        "For each request, determines the request type and delegates to the appropriate "
        "loader: USASpendingLoader, AwardsLoader, FedHierLoader, or AttachmentPipeline. "
        "Updates request status to COMPLETE or FAILED. "
        "Use --watch for continuous polling (5-second interval)."
    )

    doc.add_heading("Rate Limits", level=2)
    doc.add_paragraph(
        "USASpending: no rate limits. SAM.gov (FPDS + Fed Hierarchy): uses API key 2 (1,000/day). "
        "ATTACHMENT_ANALYSIS: no API calls (local processing)."
    )

    doc.add_heading("Change Detection", level=2)
    doc.add_paragraph(
        "Delegates to underlying loaders (USASpendingLoader, AwardsLoader, FedHierLoader) "
        "which use SHA-256 change detection."
    )

    doc.add_heading("CLI Options", level=2)
    add_table(doc,
        ["Option", "Default", "Description"],
        [
            ["--watch", "False", "Run continuously, polling every 5 seconds"],
        ]
    )

    doc.add_heading("Example Usage", level=2)
    add_code_block(doc, "python main.py job process-requests")
    add_code_block(doc, "python main.py job process-requests --watch")

    doc.add_heading("Common Issues", level=2)
    doc.add_paragraph(
        "- Requests are inserted by the C# API backend.\n"
        "- Failed requests are marked FAILED with an error message.\n"
        "- Use --watch for background service mode (Ctrl+C to stop).\n"
        "- Processes at most 10 requests per poll cycle.\n"
        "- REFRESH_FEDHIER_ORG: lookup_key = fh_org_id; refreshes single org from SAM.gov.\n"
        "- ATTACHMENT_ANALYSIS: lookup_key = notice_id; runs AI analysis on attachments."
    )

    doc.add_page_break()

    # =========================================================================
    # APPENDIX A: Daily Load Sequence
    # =========================================================================
    doc.add_heading("Appendix A: Daily Load Sequence", level=1)
    doc.add_paragraph(
        "The daily_load.bat script runs the following 13 steps in order:"
    )

    steps = [
        ("1", "load opportunities", "--max-calls 300 --key 2 --days-back 31 --force",
         "Fetch recent opportunities from SAM.gov"),
        ("2", "update fetch-descriptions", "--key 2 --limit 100 --naics <NAICS> --set-aside WOSB,8A,SBA",
         "Backfill description text (priority NAICS+set-aside first)"),
        ("3", "load usaspending-bulk", "--days-back 5",
         "Bulk load recent USASpending data"),
        ("4", "load awards", "--naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside 8a",
         "Load 8(a) set-aside awards"),
        ("5", "load awards", "--naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside WOSB",
         "Load WOSB set-aside awards"),
        ("6", "load awards", "--naics <NAICS> --days-back 10 --max-calls 100 --key 2 --set-aside SBA",
         "Load SBA set-aside awards"),
        ("7", "update link-metadata", "",
         "Enrich resource link filenames and content types"),
        ("8", "download attachments", "--missing-only --active-only --batch-size 5000",
         "Download attachment files for active opportunities"),
        ("9", "extract attachment-text", "--batch-size 5000 --workers 10",
         "Extract text from downloaded attachments"),
        ("10", "extract attachment-intel", "--batch-size 5000",
         "Extract structured intelligence from text"),
        ("11", "backfill opportunity-intel", "",
         "Propagate intel findings to opportunity table"),
        ("12", "maintain attachment-files", "",
         "Remove fully-analyzed attachment files from disk"),
    ]

    add_table(doc,
        ["Step", "Command", "Key Options", "Purpose"],
        [[s[0], s[1], s[2], s[3]] for s in steps]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "NAICS codes used in daily load: 336611, 488190, 519210, 541219, 541330, "
        "541511, 541512, 541513, 541519, 541611, 541612, 541613, 541690, 541990, "
        "561110, 561210, 561510, 561990, 611430, 611710, 621111, 621399, 624190, 812910"
    )

    doc.add_page_break()

    # =========================================================================
    # APPENDIX B: Batch Load Commands
    # =========================================================================
    doc.add_heading("Appendix B: Batch Load Commands", level=1)
    doc.add_paragraph(
        "The system provides three batch orchestration commands that run "
        "multiple loaders in sequence:"
    )

    add_table(doc,
        ["Command", "Frequency", "Description"],
        [
            ["python main.py load daily", "Daily",
             "Runs opportunities, awards (8a/WOSB/SBA), usaspending-bulk, "
             "link-metadata, attachments pipeline"],
            ["python main.py load weekly", "Weekly",
             "Includes daily loads plus entities, hierarchy, exclusions, subawards"],
            ["python main.py load monthly", "Monthly",
             "Includes weekly loads plus full entity refresh, CALC+ labor rates, "
             "full usaspending-bulk, hierarchy full-refresh"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading("Shared Infrastructure", level=2)
    doc.add_paragraph(
        "All loaders share common infrastructure:"
    )
    add_table(doc,
        ["Module", "Purpose"],
        [
            ["etl/load_manager.py", "Load orchestration, etl_load_log tracking, watermark management"],
            ["etl/change_detector.py", "SHA-256 hash computation and comparison"],
            ["etl/staging_mixin.py", "Raw JSON staging to stg_*_raw tables"],
            ["etl/batch_upsert.py", "Efficient batch INSERT ... ON DUPLICATE KEY UPDATE"],
            ["etl/etl_utils.py", "Date/decimal parsing, hash fetching utilities"],
            ["api_clients/base_client.py", "Rate limiting, retries, pagination for all API clients"],
            ["db/connection.py", "MySQL connection pooling"],
        ]
    )

    # =========================================================================
    # Save
    # =========================================================================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "FedProspect-Loader-Reference.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
