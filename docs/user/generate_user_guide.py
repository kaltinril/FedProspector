"""Generate FedProspect User Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# Styles
# ============================================================

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [
    ('Heading 1', 24, RGBColor(0x1A, 0x23, 0x7E)),
    ('Heading 2', 18, RGBColor(0x1A, 0x23, 0x7E)),
    ('Heading 3', 14, RGBColor(0x2E, 0x7D, 0x32)),
]:
    h = doc.styles[level]
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 'Heading 1' else 12)
    h.paragraph_format.space_after = Pt(6)

# Create a "Tip" style (indented, italic, with green left border effect)
tip_style = doc.styles.add_style('Tip', WD_STYLE_TYPE.PARAGRAPH)
tip_style.font.name = 'Calibri'
tip_style.font.size = Pt(10)
tip_style.font.italic = True
tip_style.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
tip_style.paragraph_format.left_indent = Cm(1)
tip_style.paragraph_format.space_before = Pt(6)
tip_style.paragraph_format.space_after = Pt(6)

# Bullet style
bullet_style = doc.styles['List Bullet']
bullet_style.font.name = 'Calibri'
bullet_style.font.size = Pt(11)


def add_tip(text):
    """Add a green tip callout paragraph."""
    p = doc.add_paragraph(style='Tip')
    run = p.add_run('TIP: ')
    run.bold = True
    p.add_run(text)


def add_bullets(items):
    """Add a bulleted list."""
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_para(text):
    """Add a normal paragraph."""
    doc.add_paragraph(text)


# ============================================================
# Title Page
# ============================================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FedProspect User Guide')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('For Business Development & Capture Managers')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# Table of Contents placeholder
# ============================================================

doc.add_heading('Table of Contents', level=1)
add_para(
    'Chapter 1: Getting Started\n'
    'Chapter 2: Dashboard\n'
    'Chapter 3: Opportunity Search\n'
    'Chapter 4: Opportunity Detail\n'
    'Chapter 5: Award Search & Detail\n'
    'Chapter 6: Entity Search & Detail\n'
    'Chapter 7: Prospect Pipeline\n'
    'Chapter 8: Notifications & Saved Searches\n'
    'Chapter 9: Tips & Best Practices'
)
doc.add_page_break()

# ============================================================
# Chapter 1: Getting Started
# ============================================================

doc.add_heading('Chapter 1: Getting Started', level=1)

doc.add_heading('What Is FedProspect?', level=2)
add_para(
    'FedProspect is a federal contract prospecting system designed for '
    'business development and capture managers. It helps you find, evaluate, '
    'and track WOSB (Women-Owned Small Business) and 8(a) federal contract '
    'opportunities so you can build a strong bid pipeline.'
)
add_para(
    'FedProspect pulls data from SAM.gov, USASpending.gov, and FPDS to give '
    'you a single place to search opportunities, research past awards, analyze '
    'competitors, and manage your prospect pipeline.'
)

doc.add_heading('Logging In', level=2)
add_para(
    'Open your web browser and navigate to the FedProspect URL provided by '
    'your organization administrator. You will see a login page.'
)
add_bullets([
    'Enter your email address and password.',
    'Click the "Sign In" button.',
    'If this is your first time logging in, you may be asked to change your '
    'temporary password.',
])
add_tip(
    'If you forget your password, contact your organization administrator. '
    'They can reset it from the Admin panel.'
)

doc.add_heading('Changing Your Password', level=2)
add_para(
    'To change your password at any time, go to Profile in the sidebar under '
    'Settings. You can also access the Change Password page directly after '
    'login if prompted.'
)

doc.add_heading('Navigating the App', level=2)
add_para(
    'FedProspect uses a sidebar navigation on the left side of the screen. '
    'The sidebar is organized into five sections:'
)

doc.add_heading('Main', level=3)
add_bullets([
    'Dashboard -- Your home base with key metrics and activity at a glance.',
])

doc.add_heading('Pipeline', level=3)
add_bullets([
    'Recommended -- AI-scored opportunities that match your company profile.',
    'Expiring Contracts -- Federal contracts nearing completion that may be '
    're-competed soon.',
    'Prospects -- Your active bid pipeline, shown as a Kanban board or list.',
])

doc.add_heading('Research', level=3)
add_bullets([
    'Opportunities -- Search all federal solicitations from SAM.gov.',
    'Awards -- Search historical contract awards from FPDS and USASpending.',
    'Entities -- Look up contractors, vendors, and competitors registered in SAM.gov.',
    'Teaming -- Find potential teaming partners based on subaward history.',
])

doc.add_heading('Tools', level=3)
add_bullets([
    'Saved Searches -- Manage your saved search criteria and run them on demand.',
])

doc.add_heading('Settings', level=3)
add_bullets([
    'Profile -- Update your display name and password.',
    'Organization -- View organization settings (org admins can manage members).',
    'Admin -- System administration (visible only to system administrators).',
])

add_tip(
    'You can collapse the sidebar to save screen space by clicking the '
    'chevron arrow at the bottom. Hover over icons to see tooltips when '
    'the sidebar is collapsed.'
)

doc.add_page_break()

# ============================================================
# Chapter 2: Dashboard
# ============================================================

doc.add_heading('Chapter 2: Dashboard', level=1)

add_para(
    'The Dashboard is your home page. It gives you a quick overview of your '
    'pipeline health, upcoming deadlines, and new opportunities worth reviewing.'
)

doc.add_heading('Stats Cards', level=2)
add_para(
    'At the top of the Dashboard, you will see four summary cards:'
)
add_bullets([
    'Open Prospects -- The total number of opportunities you are actively '
    'tracking in your pipeline.',
    'Pipeline Value -- The estimated total dollar value of all open prospects.',
    'Due This Week -- Prospects with response deadlines in the next 7 days. '
    'Click this card to jump directly to those items.',
    'Auto-Match Count -- The number of new recommended opportunities found by '
    'the system since your last visit.',
])

doc.add_heading('Due This Week', level=2)
add_para(
    'Below the stats cards, you will see a list of prospects with upcoming '
    'deadlines. Each row shows the opportunity title, deadline, set-aside '
    'type, priority, and the team member assigned to it. Click any row to '
    'open the full prospect detail.'
)

doc.add_heading('Prospects by Status Chart', level=2)
add_para(
    'A bar chart shows how many prospects are in each pipeline stage '
    '(e.g., Identified, Qualifying, Bid/No-Bid, Pursuing, Submitted). '
    'This helps you see at a glance whether your pipeline is healthy or '
    'if too many items are stuck in early stages.'
)

doc.add_heading('Recommended Opportunities', level=2)
add_para(
    'The Dashboard also shows a preview of top recommended opportunities. '
    'These are solicitations that FedProspect has scored as a good fit for '
    'your company based on NAICS codes, set-aside types, and other factors. '
    'Click "View All" to see the full Recommended Opportunities page.'
)

doc.add_heading('Expiring Contracts Preview', level=2)
add_para(
    'A preview of contracts nearing their completion date is also shown. '
    'These represent potential re-compete opportunities. Click "View All" '
    'to navigate to the full Expiring Contracts page.'
)

doc.add_heading('Workload by Assignee', level=2)
add_para(
    'If your organization has multiple users, a chart shows how prospects '
    'are distributed across team members. This helps managers balance '
    'workload and identify who might need help.'
)

add_tip(
    'The Dashboard refreshes automatically every 5 minutes. You can also '
    'refresh by navigating away and back.'
)

doc.add_page_break()

# ============================================================
# Chapter 3: Opportunity Search
# ============================================================

doc.add_heading('Chapter 3: Opportunity Search', level=1)

add_para(
    'The Opportunities page is your primary tool for finding new federal '
    'solicitations to bid on. It searches across all opportunities loaded '
    'from SAM.gov.'
)

doc.add_heading('Search Filters', level=2)
add_para(
    'The filter bar at the top of the page lets you narrow results using '
    'several criteria:'
)
add_bullets([
    'Keyword -- Search within opportunity titles and descriptions.',
    'Solicitation # -- Search by a specific solicitation number if you '
    'already have one.',
    'NAICS Code -- Filter by the 6-digit North American Industry '
    'Classification System code (e.g., 541511 for custom software).',
    'Set-Aside -- Filter by small business set-aside type such as WOSB, '
    '8(a), HUBZone, SDVOSB, or Total Small Business.',
    'Department -- Filter by the federal department or agency name.',
    'State (POP) -- Filter by the Place of Performance state abbreviation '
    '(e.g., VA, DC, MD).',
    'Days Out -- Show only opportunities with a deadline within this many '
    'days.',
    'Open Only -- When set to "Yes" (the default), only shows active, '
    'open solicitations. Switch to "No" or "All" to include closed or '
    'archived opportunities.',
])
add_tip(
    'You can combine multiple filters. For example, set NAICS to "541511", '
    'Set-Aside to "WOSB", and State to "VA" to find women-owned software '
    'opportunities in Virginia.'
)

doc.add_heading('Results Grid', level=2)
add_para(
    'Search results appear in a sortable data grid below the filters. '
    'Key columns include:'
)
add_bullets([
    'Title -- The opportunity title. Click any row to open its detail page.',
    'Solicitation # -- The official solicitation number.',
    'Department -- The issuing federal agency.',
    'Posted Date -- When the opportunity was published on SAM.gov.',
    'Response Deadline -- The due date for responses, with a countdown '
    'showing days remaining.',
    'Set-Aside -- The small business set-aside category, shown as a '
    'color-coded chip.',
    'NAICS -- The NAICS code and sector.',
    'Est. Value -- Estimated contract value, if available.',
    'State -- Place of Performance state.',
    'Status -- Shows "In Pipeline" if you have already added this '
    'opportunity as a prospect.',
])
add_para(
    'Click any column header to sort results. Click again to reverse the '
    'sort order. The grid supports pagination -- use the controls at the '
    'bottom to navigate between pages.'
)

doc.add_heading('Saving a Search', level=2)
add_para(
    'After setting your filters, click the "Save Search" button in the '
    'toolbar. Give your search a name and optional description. Saved '
    'searches appear in the Saved Searches page and on your Dashboard. '
    'You can also enable notifications so FedProspect alerts you when new '
    'matching opportunities appear.'
)

doc.add_heading('Adding to Pipeline', level=2)
add_para(
    'To add an opportunity to your prospect pipeline directly from search '
    'results, select the row and click the "Add to Pipeline" button (the '
    'target icon). This creates a new prospect record and moves the '
    'opportunity into your active tracking workflow.'
)

doc.add_heading('Exporting Results', level=2)
add_para(
    'Click the download icon in the toolbar to export your current search '
    'results as a CSV file. The export respects your active filters.'
)

doc.add_heading('URL Sharing', level=2)
add_para(
    'Your search filters are stored in the browser URL. This means you '
    'can copy the URL from your address bar and share it with a colleague. '
    'When they open the link, they will see the same filters applied.'
)

add_tip(
    'Bookmark frequently-used searches in your browser for quick access, '
    'or use the built-in Saved Searches feature for notification support.'
)

doc.add_page_break()

# ============================================================
# Chapter 4: Opportunity Detail
# ============================================================

doc.add_heading('Chapter 4: Opportunity Detail', level=1)

add_para(
    'When you click on an opportunity from search results, the '
    'Recommended list, or anywhere else in the app, you are taken to '
    'the Opportunity Detail page. This page has six tabs that give you '
    'a comprehensive view of the opportunity.'
)

doc.add_heading('Overview Tab', level=2)
add_para(
    'The Overview tab shows the core information about the opportunity:'
)
add_bullets([
    'Key Facts Grid -- Title, solicitation number, department, sub-tier '
    'agency, office, type of notice, and posted/response dates.',
    'Set-Aside & NAICS -- The set-aside category, NAICS code and '
    'description, and size standard.',
    'Place of Performance -- City, state, ZIP, and country.',
    'Contract Value -- Estimated contract value and award information '
    'if already awarded.',
    'Description -- The full opportunity description text.',
    'Resource Links -- Downloadable attachments and external links '
    'associated with the solicitation (PDFs, Word docs, spreadsheets).',
    'Related Awards -- If this is a re-compete, shows the prior contract '
    'award(s) with vendor names, dates, and dollar amounts.',
    'Amendments -- Lists any amendments or modifications to the original '
    'solicitation, with dates and updated deadlines.',
    'SAM.gov Link -- A direct link to view the opportunity on SAM.gov.',
])
add_tip(
    'The deadline countdown at the top of the page shows how many days '
    'remain until the response is due. Red means less than 7 days -- '
    'act fast!'
)

doc.add_heading('Qualification & pWin Tab', level=2)
add_para(
    'This tab helps you assess whether your company is qualified to bid '
    'and what your probability of winning (pWin) looks like.'
)

doc.add_heading('pWin Gauge', level=3)
add_para(
    'The pWin Gauge is a visual score from 0 to 100 that estimates your '
    'probability of winning. It considers factors such as:'
)
add_bullets([
    'Set-Aside Match -- Does the set-aside type match your certifications?',
    'NAICS Experience -- Do you have past performance in this NAICS code?',
    'Competitive Landscape -- How many competitors typically bid on '
    'similar work?',
    'Incumbent Advantage -- Is there an incumbent with a strong track '
    'record?',
    'Time Remaining -- How much time is left to prepare a quality response?',
])
add_para(
    'The gauge is color-coded: green for High (70+), yellow for Medium '
    '(40-69), orange for Low (20-39), and red for Very Low (under 20). '
    'Below the gauge, you will see a breakdown of each scoring factor '
    'with details and suggestions for improving your score.'
)

doc.add_heading('Qualification Checklist', level=3)
add_para(
    'The Qualification Checklist runs a series of automated checks against '
    'the opportunity requirements and your company profile:'
)
add_bullets([
    'Certification checks -- Do you hold the required SBA certifications '
    '(WOSB, 8(a), HUBZone, etc.)?',
    'Experience checks -- Do you have past contracts in the relevant '
    'NAICS code?',
    'Compliance checks -- Is your SAM.gov registration active? Are you '
    'free of exclusions?',
    'Logistics checks -- Is the Place of Performance in a state where '
    'you operate?',
])
add_para(
    'Each item shows a Pass (green check), Fail (red X), Warning '
    '(yellow), or Unknown status. The overall qualification status is '
    'summarized as "Qualified," "Partially Qualified," or "Not Qualified."'
)

doc.add_heading('Competitive Intelligence Tab', level=2)
add_para(
    'This tab provides intelligence about the competitive landscape for '
    'the opportunity.'
)

doc.add_heading('Incumbent Analysis', level=3)
add_para(
    'If FedProspect identifies an incumbent contractor (the company that '
    'currently holds the contract), you will see:'
)
add_bullets([
    'Incumbent name, UEI, and registration status.',
    'Current contract value, dollars obligated, and burn rate.',
    'Period of performance dates and months remaining.',
    'Vulnerability signals -- Indicators that the incumbent may be '
    'vulnerable, such as SAM.gov registration expiring, active '
    'exclusions, high burn rate, or contract nearing completion.',
    'Consecutive wins -- How many times the incumbent has won this '
    'contract in a row.',
])

doc.add_heading('Competitive Landscape', level=3)
add_para(
    'Shows the broader market for this type of work:'
)
add_bullets([
    'Total contracts and total value in this NAICS/agency combination.',
    'Average award value for similar contracts.',
    'Competition level (Low, Moderate, High) based on the number of '
    'distinct vendors.',
    'Top vendors by market share, shown in a bar chart.',
])

doc.add_heading('Market Share Chart', level=3)
add_para(
    'A visual chart showing the top vendors in this NAICS code and their '
    'share of total contract value. Use this to understand who your main '
    'competitors are and how the market is distributed.'
)

doc.add_heading('Document Intelligence Tab', level=2)
add_para(
    'FedProspect automatically downloads and analyzes solicitation '
    'attachments (PDFs, Word documents, spreadsheets) to extract key '
    'intelligence:'
)
add_bullets([
    'Security Clearance -- Whether the work requires a clearance and '
    'at what level (Secret, Top Secret, etc.).',
    'Evaluation Method -- How proposals will be evaluated (Lowest Price '
    'Technically Acceptable, Best Value Trade-Off, etc.).',
    'Vehicle Type -- The contract vehicle being used.',
    'Recompete Status -- Whether this is a new requirement or a '
    're-competition of an existing contract.',
    'Incumbent Name -- The current contractor, if mentioned in the '
    'documents.',
    'Scope Summary -- A brief summary of the work to be performed.',
    'Period of Performance -- The expected contract duration.',
    'Labor Categories -- Specific labor categories or skill sets '
    'required.',
    'Key Requirements -- Important requirements extracted from the '
    'documents.',
])
add_para(
    'Each extracted field shows a confidence level (High, Medium, or Low) '
    'and the source document and page number where it was found. You can '
    'also see a list of all attachments with their download and text '
    'extraction status.'
)
add_tip(
    'Document Intelligence saves hours of manual reading. Use the extracted '
    'key requirements to quickly assess whether a solicitation is a good '
    'fit before committing time to a full review.'
)

doc.add_heading('History Tab', level=2)
add_para(
    'The History tab shows the timeline of this opportunity:'
)
add_bullets([
    'When it was first loaded into FedProspect.',
    'When it was last updated.',
    'Amendments and modifications over time.',
    'USASpending award data if the contract has been awarded.',
])

doc.add_heading('Actions Tab', level=2)
add_para(
    'From the Actions section (available via buttons at the top of the '
    'detail page), you can:'
)
add_bullets([
    'Add to Pipeline -- Create a prospect record to start tracking this '
    'opportunity through your bid workflow.',
    'Save as Search -- Create a saved search based on this opportunity\'s '
    'NAICS and set-aside to find similar ones.',
    'View on SAM.gov -- Open the original listing on SAM.gov in a new '
    'browser tab.',
])

doc.add_page_break()

# ============================================================
# Chapter 5: Award Search & Detail
# ============================================================

doc.add_heading('Chapter 5: Award Search & Detail', level=1)

add_para(
    'The Awards section helps you research historical federal contract '
    'awards. This is essential for understanding your competition, '
    'estimating contract values, and identifying re-compete opportunities.'
)

doc.add_heading('Searching Awards', level=2)
add_para(
    'Navigate to Awards in the Research section of the sidebar. The search '
    'filters include:'
)
add_bullets([
    'Solicitation # -- Search by the solicitation number that led to the '
    'award.',
    'NAICS Code -- Filter by the industry classification code.',
    'Agency -- The awarding federal agency.',
    'Vendor UEI -- The Unique Entity Identifier of the winning vendor.',
    'Vendor Name -- Search by the contractor name.',
    'Set-Aside -- Filter by small business set-aside type.',
    'Min/Max Value -- Set a dollar range for the award amount.',
    'Date Range -- Filter by the date the award was signed.',
])

doc.add_heading('Award Results', level=2)
add_para(
    'The results grid shows key columns including contract ID, agency, '
    'vendor name, date signed, dollars obligated, base and all options '
    'value, NAICS code, contract type, and competition level. Click any '
    'row to open the full award detail.'
)

doc.add_heading('Award Detail Page', level=2)
add_para(
    'The Award Detail page provides comprehensive information about a '
    'specific contract award:'
)
add_bullets([
    'Contract Information -- Contract ID, type, pricing structure, '
    'description, and solicitation details.',
    'Agency Information -- Awarding agency, contracting office, and '
    'funding agency.',
    'Vendor Information -- Winning contractor name, UEI, and a link to '
    'their entity profile.',
    'Financial Summary -- Dollars obligated, base and all options value, '
    'and full transaction history showing each modification.',
    'Competition Data -- Number of offers received, extent competed, and '
    'set-aside type.',
    'Performance -- Place of performance, effective date, completion date, '
    'and ultimate completion date.',
])

doc.add_heading('Burn Rate Analysis', level=3)
add_para(
    'For awards with transaction history, FedProspect calculates a burn '
    'rate analysis:'
)
add_bullets([
    'Total obligated vs. base and all options -- How much of the ceiling '
    'has been spent.',
    'Percent spent -- A progress bar showing how close the contract is to '
    'being fully funded.',
    'Monthly rate -- The average monthly spend rate.',
    'Monthly Breakdown Chart -- A bar chart showing spending by month, '
    'so you can see trends and acceleration.',
])
add_tip(
    'High burn rates on contracts nearing completion are a strong signal '
    'that a re-compete solicitation may appear soon. Watch these closely '
    'in the Expiring Contracts view.'
)

doc.add_heading('Expiring Contracts', level=2)
add_para(
    'Navigate to Expiring Contracts in the Pipeline section of the '
    'sidebar. This page shows federal contracts that are approaching '
    'their completion date, making them prime candidates for re-compete '
    'solicitations.'
)
add_para('For each expiring contract, you will see:')
add_bullets([
    'Contract ID and description.',
    'Current vendor and their registration status.',
    'Awarding agency and office.',
    'Contract value, dollars obligated, and burn rate.',
    'Completion date and months remaining.',
    'Resolicitation status -- Whether a new solicitation has already '
    'been posted.',
    'Set-aside shift detection -- Whether the set-aside type has changed '
    'from the predecessor contract, which could create new opportunities '
    'for small businesses.',
])
add_tip(
    'Pay special attention to contracts where "Shift Detected" appears. '
    'This means the government may have changed the set-aside category, '
    'potentially opening the door for your business type.'
)

doc.add_page_break()

# ============================================================
# Chapter 6: Entity Search & Detail
# ============================================================

doc.add_heading('Chapter 6: Entity Search & Detail', level=1)

add_para(
    'The Entities section lets you look up contractors, vendors, and '
    'potential competitors registered in SAM.gov. This is valuable for '
    'competitive research and finding teaming partners.'
)

doc.add_heading('Searching Entities', level=2)
add_para('The search filters include:')
add_bullets([
    'Name -- Search by legal business name or DBA name.',
    'UEI -- Search by Unique Entity Identifier.',
    'NAICS Code -- Find entities registered under a specific NAICS code.',
    'State -- Filter by the entity\'s state.',
    'Business Type -- Filter by business type code.',
    'SBA Certification -- Filter by SBA certification type (WOSB, 8(a), '
    'HUBZone, etc.).',
    'Registration Status -- Filter by active or inactive SAM.gov '
    'registration.',
])

doc.add_heading('Entity Detail Page', level=2)
add_para(
    'The Entity Detail page shows everything known about a registered '
    'contractor:'
)
add_bullets([
    'Registration Info -- UEI, CAGE code, registration status, '
    'activation and expiration dates.',
    'Business Profile -- Legal name, DBA, structure (LLC, Corporation, '
    'etc.), and state/country of incorporation.',
    'Addresses -- Physical and mailing addresses with congressional '
    'districts.',
    'NAICS Codes -- All registered NAICS codes, with the primary code '
    'highlighted and SBA small business size status.',
    'PSC Codes -- Product and Service Codes the entity has registered for.',
    'Business Types -- Full list of business type designations (e.g., '
    'Woman-Owned, Minority-Owned, Veteran-Owned).',
    'SBA Certifications -- Active and historical SBA certifications '
    'with entry and exit dates.',
    'Points of Contact -- Government and electronic business points of '
    'contact.',
    'Exclusion Status -- Whether the entity has any active exclusions '
    '(debarment, suspension, etc.).',
])

doc.add_heading('Competitor Profiles', level=2)
add_para(
    'When viewing an entity that has won federal contracts, you can '
    'access their competitor profile which shows:'
)
add_bullets([
    'Total past contracts and total dollars obligated.',
    'Win rate and average contract size.',
    'Most recent award date.',
    'Recent awards list with contract details.',
    'NAICS sector and business type categories.',
])
add_tip(
    'Use competitor profiles to understand your competition before '
    'deciding to bid. If a competitor has won the same contract five '
    'times in a row, they have a strong incumbent advantage.'
)

doc.add_heading('Teaming Partners', level=2)
add_para(
    'Navigate to Teaming in the Research section to find potential '
    'teaming partners. This page analyzes subaward data to identify '
    'prime-sub relationships:'
)
add_bullets([
    'Search by NAICS code to find primes working in your industry.',
    'Filter by minimum number of subawards to find active teaming '
    'partners.',
    'Search by prime or sub UEI to see specific relationships.',
    'View total subaward amounts and number of unique subcontractors.',
])
add_tip(
    'Finding a prime contractor who already teams with small businesses '
    'in your NAICS code is one of the best ways to break into federal '
    'contracting.'
)

doc.add_page_break()

# ============================================================
# Chapter 7: Prospect Pipeline
# ============================================================

doc.add_heading('Chapter 7: Prospect Pipeline', level=1)

add_para(
    'The Prospect Pipeline is where you manage the opportunities you '
    'have decided to pursue. Think of it as your bid tracking system.'
)

doc.add_heading('Adding Opportunities to Your Pipeline', level=2)
add_para(
    'There are several ways to add an opportunity to your pipeline:'
)
add_bullets([
    'From Opportunity Search -- Click the "Add to Pipeline" button on '
    'a search result row.',
    'From Opportunity Detail -- Click "Add to Pipeline" at the top of '
    'the detail page.',
    'From Recommended Opportunities -- Add directly from the '
    'recommendations list.',
])
add_para(
    'When adding a prospect, you can set an initial priority (Critical, '
    'High, Medium, Low) and optionally assign it to a team member.'
)

doc.add_heading('Kanban Board View', level=2)
add_para(
    'The default view of the Prospects page is a Kanban board. Each '
    'column represents a pipeline stage:'
)
add_bullets([
    'Identified -- New opportunities you are aware of but have not yet '
    'evaluated.',
    'Qualifying -- Opportunities you are actively evaluating for fit.',
    'Bid/No-Bid -- Decision point: should you pursue this or pass?',
    'Pursuing -- You have decided to bid and are preparing your response.',
    'Submitted -- Your proposal has been submitted and you are awaiting '
    'the award decision.',
    'Won / Lost / No-Bid -- Final outcomes.',
])
add_para(
    'Each card on the board shows the opportunity title, deadline '
    'countdown, set-aside type, and priority flag. You can drag and drop '
    'cards between columns to change their status.'
)
add_tip(
    'Use the priority flags to highlight your most important opportunities. '
    'Critical items appear with a red flag so they stand out.'
)

doc.add_heading('List View', level=2)
add_para(
    'Toggle to List View using the icon in the toolbar for a traditional '
    'data grid. This view is better when you have many prospects and '
    'need to sort or filter by specific columns. You can filter by '
    'status, priority, assigned user, NAICS code, and set-aside type.'
)

doc.add_heading('Prospect Detail', level=2)
add_para(
    'Click any prospect card or row to open the Prospect Detail page. '
    'Here you can:'
)
add_bullets([
    'Update status -- Move the prospect to a different pipeline stage.',
    'Set priority -- Change the priority level.',
    'Assign team members -- Reassign the prospect to another user.',
    'Add notes -- Record capture notes, call summaries, meeting notes, '
    'go/no-go rationale, and other observations.',
    'View the linked opportunity -- Jump to the full Opportunity Detail '
    'page.',
    'See the Go/No-Go Score -- An automated score based on set-aside '
    'match, time remaining, NAICS experience, and estimated value.',
    'Manage team members -- Add teaming partners or subcontractors with '
    'roles, proposed rates, and commitment percentages.',
    'Create a proposal -- When you move to "Pursuing," you can create '
    'a formal proposal record with milestones and document tracking.',
])

doc.add_heading('Proposal Management', level=2)
add_para(
    'Each prospect can have an associated proposal. The proposal record '
    'tracks:'
)
add_bullets([
    'Proposal status (Draft, In Review, Submitted, etc.).',
    'Submission deadline.',
    'Estimated value and win probability.',
    'Milestones -- Key dates and tasks needed to complete the proposal '
    '(e.g., Draft Due, Review Complete, Final Submission).',
    'Documents -- Track proposal documents with file names, types, '
    'and upload dates.',
    'Lessons Learned -- After the outcome is known, record what went '
    'well and what to improve.',
])

doc.add_page_break()

# ============================================================
# Chapter 8: Notifications & Saved Searches
# ============================================================

doc.add_heading('Chapter 8: Notifications & Saved Searches', level=1)

doc.add_heading('Notifications', level=2)
add_para(
    'FedProspect sends you notifications when important events occur. '
    'Access your notification center by clicking the bell icon in the '
    'top navigation bar, or by going to the Notifications page.'
)
add_para('Notification types include:')
add_bullets([
    'New matching opportunities found by your saved searches.',
    'Deadline reminders for prospects approaching their response date.',
    'Status changes on prospects you are assigned to.',
    'System announcements from administrators.',
])
add_para(
    'Each notification shows a title, message, timestamp, and a link to '
    'the related item. Unread notifications are highlighted. You can mark '
    'individual notifications as read, or mark all as read.'
)
add_para(
    'Filter notifications by type or by read/unread status using the '
    'controls at the top of the notification center.'
)

doc.add_heading('Saved Searches', level=2)
add_para(
    'Saved Searches let you store your search criteria and run them '
    'on demand to check for new results. Navigate to Saved Searches '
    'in the Tools section of the sidebar.'
)
add_para('For each saved search, you can:')
add_bullets([
    'Run Now -- Execute the search and see current matching '
    'opportunities. The system highlights how many are new since the '
    'last run.',
    'Edit -- Update the search name, description, or filter criteria.',
    'Enable/Disable Notifications -- Toggle whether the system '
    'notifies you when new matches are found.',
    'Delete -- Remove the saved search.',
])
add_para(
    'Saved search filters can include set-aside codes, NAICS codes, '
    'states, award amount ranges, opportunity types, and a days-back '
    'window.'
)
add_tip(
    'Create a saved search for each NAICS code and set-aside combination '
    'you want to track. Enable notifications so you never miss a new '
    'opportunity that matches your capabilities.'
)

doc.add_page_break()

# ============================================================
# Chapter 9: Tips & Best Practices
# ============================================================

doc.add_heading('Chapter 9: Tips & Best Practices', level=1)

doc.add_heading('WOSB & 8(a) Search Strategies', level=2)
add_para(
    'If your company holds WOSB or 8(a) certifications, these strategies '
    'will help you find the best opportunities:'
)
add_bullets([
    'Always filter by your set-aside type first. WOSB set-asides are '
    'restricted to women-owned small businesses, dramatically reducing '
    'competition.',
    'Search your primary NAICS codes regularly. Create saved searches '
    'for each NAICS code you are registered under.',
    'Check the "Open Only" filter to focus on active solicitations you '
    'can still respond to.',
    'Use the Recommended Opportunities page -- it already factors in your '
    'certifications and NAICS codes to surface the best matches.',
    'Monitor Expiring Contracts in your NAICS codes. Many expiring '
    'contracts are re-solicited with small business set-asides.',
    'Look for set-aside shifts -- when a full-and-open contract is '
    're-competed as a small business set-aside, it creates a fresh '
    'competitive field.',
])

doc.add_heading('Using Competitive Intelligence Effectively', level=2)
add_bullets([
    'Before deciding to bid, always check the Competitive Intelligence '
    'tab on the Opportunity Detail page.',
    'If there is an incumbent with 3+ consecutive wins and no '
    'vulnerability signals, the opportunity may be hard to win. Consider '
    'whether you can offer a truly differentiated solution.',
    'Look for vulnerability signals such as expiring SAM.gov '
    'registration, active exclusions, or high contract burn rates.',
    'Use the Market Share chart to understand who your real competitors '
    'are. If the top 3 vendors hold 80% of the market, it may be '
    'challenging to break in without a teaming strategy.',
    'Research competitors using the Entity Detail page to understand '
    'their business types, certifications, and past performance.',
])

doc.add_heading('Building a Prospect Pipeline Workflow', level=2)
add_para(
    'A well-managed pipeline is the key to consistent wins. Here is a '
    'recommended workflow:'
)

doc.add_heading('Step 1: Cast a Wide Net', level=3)
add_para(
    'Set up saved searches for all your NAICS codes and set-aside types. '
    'Review the Recommended Opportunities page weekly. Check Expiring '
    'Contracts for re-compete opportunities. Add anything that looks '
    'promising to your pipeline as "Identified."'
)

doc.add_heading('Step 2: Qualify Quickly', level=3)
add_para(
    'Move items to "Qualifying" and use the Qualification Checklist and '
    'pWin score to assess fit. Check Document Intelligence for security '
    'clearance requirements, evaluation criteria, and key requirements. '
    'If you do not meet the basic qualifications, move to "No-Bid" and '
    'add a note explaining why.'
)

doc.add_heading('Step 3: Make a Go/No-Go Decision', level=3)
add_para(
    'Move promising opportunities to "Bid/No-Bid" and review the '
    'competitive landscape. Consider: Is the estimated value worth the '
    'effort? Do you have the past performance? Is there a strong '
    'incumbent? Does the timeline work? Record your rationale in the '
    'prospect notes.'
)

doc.add_heading('Step 4: Pursue with Focus', level=3)
add_para(
    'Move "Go" decisions to "Pursuing." Create a proposal record with '
    'milestones. Assign team members and set commitment percentages. '
    'Use the deadline countdown to stay on track.'
)

doc.add_heading('Step 5: Track Outcomes', level=3)
add_para(
    'After submitting, move to "Submitted." When the result is announced, '
    'record the outcome as Won or Lost. For losses, add lessons learned '
    'to improve future bids. Review your win/loss ratio on the Dashboard.'
)

doc.add_heading('General Tips', level=2)
add_bullets([
    'Log in at least weekly to check your Dashboard for new '
    'recommendations and upcoming deadlines.',
    'Use the priority flags consistently. Critical should mean "must '
    'respond" -- do not overuse it.',
    'Add notes to prospects as you learn new information. This builds '
    'an audit trail and helps teammates who may take over.',
    'Use the Teaming page to find prime contractors when you need a '
    'teaming partner for larger opportunities.',
    'Export search results to CSV when you need to share data with '
    'colleagues who do not have FedProspect access.',
    'Collapse the sidebar when you need more screen space for data '
    'grids and detail pages.',
    'Use keyboard shortcuts: press Tab to move between filter fields '
    'and Enter to search.',
])

add_tip(
    'The most successful capture teams review their pipeline in a weekly '
    'standup meeting. Use the Kanban board view to walk through each '
    'stage and discuss next steps for every active prospect.'
)

# ============================================================
# Save the document
# ============================================================

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'FedProspect-User-Guide.docx')
doc.save(output_path)
print(f'User guide saved to: {output_path}')
